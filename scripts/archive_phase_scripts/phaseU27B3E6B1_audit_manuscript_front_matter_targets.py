#!/usr/bin/env python3
"""
Phase U27B3E6B1
Read-only audit of the accession-corrected UTI HostOmics manuscript before
front-matter insertion and reference finalization.

This phase does not modify the DOCX. It inventories:
- source-document checksum and size;
- paragraphs, styles, headings, and likely front-matter/declaration targets;
- tables and embedded media;
- section/page-layout properties;
- existing dataset accessions and unresolved placeholder-like text.

The outputs are intended to guide a targeted U27B3E6B2 reconstruction without
guessing document structure or disturbing embedded figures and citation fields.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document


VERSION = "U27B3E6B1_v1.0_2026-07-17"
TAG = "phaseU27B3E6B1_manuscript_front_matter_target_audit"

DEFAULT_SOURCE = (
    "09_manuscript_docx/phaseU27B3E22_targeted_accession_correction/"
    "UTI_HostOmics_preZotero_manuscript_v6_3_U27B3E22_accession_corrected.docx"
)

TARGET_PATTERNS = [
    ("title_or_running_title", r"\b(title|running title)\b"),
    ("author_line", r"\b(author|authors)\b"),
    ("affiliation", r"\b(affiliation|department|university|college|company)\b"),
    ("corresponding_author", r"\b(corresponding author|correspondence)\b"),
    ("funding", r"\b(funding|financial support)\b"),
    ("competing_interests", r"\b(competing interests?|conflicts? of interest)\b"),
    ("author_contributions", r"\b(author contributions?|credit|cr[eé]dit)\b"),
    ("acknowledgements", r"\b(acknowledg(e)?ments?)\b"),
    ("data_availability", r"\b(data availability|availability of data)\b"),
    ("code_availability", r"\b(code availability|software availability)\b"),
    ("references", r"^\s*references\s*$"),
    ("supplementary", r"\b(supplementary|supporting information)\b"),
    ("repository", r"\b(github|repository|zenodo|doi)\b"),
    ("dataset_accession", r"\bGSE\d{5,}\b"),
    ("placeholder", r"\b(TBD|TO BE ADDED|PLACEHOLDER|INSERT|PENDING|UNRESOLVED)\b"),
]

EXPECTED_ACCESSIONS = ["GSE112098", "GSE186800", "GSE252321", "GSE280297"]


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def points(length) -> str:
    if length is None:
        return ""
    try:
        return f"{length.pt:.3f}"
    except Exception:
        return str(length)


def inches(length) -> str:
    if length is None:
        return ""
    try:
        return f"{length.inches:.4f}"
    except Exception:
        return str(length)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root",
        default="__UTI_HOSTOMICS_PROJECT_ROOT__",
    )
    parser.add_argument(
        "--source",
        default=DEFAULT_SOURCE,
        help="Source DOCX, absolute or relative to project root.",
    )
    args = parser.parse_args()

    root = Path(args.project_root).resolve()
    source = Path(args.source)
    if not source.is_absolute():
        source = root / source
    source = source.resolve()

    if root != Path("__UTI_HOSTOMICS_PROJECT_ROOT__"):
        raise RuntimeError(f"Wrong project root: {root}")
    if not source.exists():
        raise FileNotFoundError(f"Source manuscript not found: {source}")
    if source.suffix.lower() != ".docx":
        raise RuntimeError(f"Source is not a DOCX: {source}")

    out_tables = root / "06_tables" / TAG
    out_results = root / "05_results" / TAG
    out_metadata = root / "03_metadata" / TAG
    for d in (out_tables, out_results, out_metadata):
        d.mkdir(parents=True, exist_ok=True)

    doc = Document(str(source))

    paragraph_rows = []
    target_rows = []
    all_text = []

    for index, paragraph in enumerate(doc.paragraphs):
        text = clean_text(paragraph.text)
        all_text.append(text)
        style_name = (
            paragraph.style.name
            if paragraph.style is not None
            else ""
        )
        paragraph_rows.append(
            {
                "paragraph_index": index,
                "style": style_name,
                "text_length": len(text),
                "is_blank": text == "",
                "text": text,
            }
        )

        for target_type, pattern in TARGET_PATTERNS:
            if re.search(pattern, text, flags=re.IGNORECASE):
                target_rows.append(
                    {
                        "paragraph_index": index,
                        "style": style_name,
                        "target_type": target_type,
                        "matched_pattern": pattern,
                        "text": text,
                    }
                )

    paragraphs = pd.DataFrame(paragraph_rows)
    targets = pd.DataFrame(target_rows)

    table_rows = []
    for table_index, table in enumerate(doc.tables):
        n_rows = len(table.rows)
        n_cols = max((len(row.cells) for row in table.rows), default=0)
        first_row = " | ".join(
            clean_text(cell.text)
            for cell in table.rows[0].cells
        ) if n_rows else ""
        table_rows.append(
            {
                "table_index": table_index,
                "rows": n_rows,
                "columns": n_cols,
                "first_row_text": first_row,
            }
        )
    tables = pd.DataFrame(table_rows)

    section_rows = []
    for idx, section in enumerate(doc.sections):
        section_rows.append(
            {
                "section_index": idx,
                "start_type": str(section.start_type),
                "page_width_inches": inches(section.page_width),
                "page_height_inches": inches(section.page_height),
                "top_margin_inches": inches(section.top_margin),
                "bottom_margin_inches": inches(section.bottom_margin),
                "left_margin_inches": inches(section.left_margin),
                "right_margin_inches": inches(section.right_margin),
                "header_distance_inches": inches(section.header_distance),
                "footer_distance_inches": inches(section.footer_distance),
                "orientation": str(section.orientation),
            }
        )
    sections = pd.DataFrame(section_rows)

    media_rows = []
    with zipfile.ZipFile(source) as archive:
        for name in sorted(archive.namelist()):
            if name.startswith("word/media/") and not name.endswith("/"):
                info = archive.getinfo(name)
                media_rows.append(
                    {
                        "member": name,
                        "size_bytes": info.file_size,
                        "compressed_size_bytes": info.compress_size,
                        "extension": Path(name).suffix.lower(),
                    }
                )
    media = pd.DataFrame(media_rows)

    body_text = "\n".join(all_text)
    accession_rows = []
    for accession in EXPECTED_ACCESSIONS:
        hits = len(re.findall(rf"\b{re.escape(accession)}\b", body_text))
        accession_rows.append(
            {
                "accession": accession,
                "occurrences_in_paragraph_text": hits,
                "present": hits > 0,
            }
        )
    accession_inventory = pd.DataFrame(accession_rows)

    placeholder_hits = targets[
        targets["target_type"] == "placeholder"
    ].copy() if not targets.empty else pd.DataFrame()

    inventory = pd.DataFrame(
        [
            {
                "version": VERSION,
                "project_root": str(root),
                "source_docx": str(source),
                "source_size_bytes": source.stat().st_size,
                "source_sha256": sha256(source),
                "paragraphs": len(doc.paragraphs),
                "nonblank_paragraphs": int(
                    (~paragraphs["is_blank"]).sum()
                ),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "embedded_media_files": len(media_rows),
                "target_candidate_rows": len(target_rows),
                "placeholder_candidate_rows": len(placeholder_hits),
                "expected_accessions_present": int(
                    accession_inventory["present"].sum()
                ),
                "expected_accessions_total": len(EXPECTED_ACCESSIONS),
                "document_modified": False,
            }
        ]
    )

    decision = (
        "READY_FOR_U27B3E6B2_TARGETED_FRONT_MATTER_INSERTION"
        if len(doc.paragraphs) > 0
        and len(doc.sections) > 0
        and accession_inventory["present"].all()
        else "TARGETED_U27B3E6B1_MANUSCRIPT_STRUCTURE_REVIEW_REQUIRED"
    )

    decision_frame = pd.DataFrame(
        [
            {
                "phase": "U27B3E6B1",
                "decision": decision,
                "source_docx": str(source),
                "source_sha256": sha256(source),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "embedded_media_files": len(media_rows),
                "expected_accessions_present": int(
                    accession_inventory["present"].sum()
                ),
                "expected_accessions_total": len(EXPECTED_ACCESSIONS),
                "document_modified": False,
                "next_phase": (
                    "U27B3E6B2 insert confirmed authorship, affiliations, "
                    "correspondence, declarations and availability statements "
                    "while preserving figures and citation fields"
                ),
            }
        ]
    )

    paragraphs.to_csv(
        out_tables / "UTI_HostOmics_U27B3E6B1_paragraph_inventory.tsv",
        sep="\t",
        index=False,
    )
    targets.to_csv(
        out_tables / "UTI_HostOmics_U27B3E6B1_front_matter_target_candidates.tsv",
        sep="\t",
        index=False,
    )
    tables.to_csv(
        out_tables / "UTI_HostOmics_U27B3E6B1_table_inventory.tsv",
        sep="\t",
        index=False,
    )
    sections.to_csv(
        out_tables / "UTI_HostOmics_U27B3E6B1_section_inventory.tsv",
        sep="\t",
        index=False,
    )
    media.to_csv(
        out_tables / "UTI_HostOmics_U27B3E6B1_embedded_media_inventory.tsv",
        sep="\t",
        index=False,
    )
    accession_inventory.to_csv(
        out_tables / "UTI_HostOmics_U27B3E6B1_accession_inventory.tsv",
        sep="\t",
        index=False,
    )
    inventory.to_csv(
        out_metadata / "UTI_HostOmics_U27B3E6B1_source_document_inventory.tsv",
        sep="\t",
        index=False,
    )
    decision_frame.to_csv(
        out_tables / "UTI_HostOmics_U27B3E6B1_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    report = out_results / "UTI_HostOmics_U27B3E6B1_manuscript_front_matter_target_audit_report.md"
    report.write_text(
        "\n".join(
            [
                "# Phase U27B3E6B1 - Manuscript front-matter target audit",
                "",
                f"- Version: `{VERSION}`",
                f"- Decision: **{decision}**",
                f"- Source: `{source}`",
                f"- Source SHA256: `{sha256(source)}`",
                f"- Paragraphs: **{len(doc.paragraphs)}**",
                f"- Tables: **{len(doc.tables)}**",
                f"- Sections: **{len(doc.sections)}**",
                f"- Embedded media files: **{len(media_rows)}**",
                f"- Front-matter/declaration target candidates: **{len(target_rows)}**",
                f"- Placeholder candidates: **{len(placeholder_hits)}**",
                f"- Expected GEO accessions present: **{int(accession_inventory['present'].sum())}/{len(EXPECTED_ACCESSIONS)}**",
                "",
                "## Integrity boundary",
                "",
                "This phase is read-only. The manuscript, figures, citation fields, supplementary files and source data were not modified.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    run_manifest = {
        "version": VERSION,
        "decision": decision,
        "source_docx": str(source),
        "source_sha256": sha256(source),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "embedded_media_files": len(media_rows),
        "expected_accessions_present": int(
            accession_inventory["present"].sum()
        ),
        "expected_accessions_total": len(EXPECTED_ACCESSIONS),
        "document_modified": False,
    }
    (
        out_results / "UTI_HostOmics_U27B3E6B1_run_manifest.json"
    ).write_text(
        json.dumps(run_manifest, indent=2, sort_keys=True),
        encoding="utf-8",
    )

    print(f"[U27B3E6B1] Decision: {decision}")
    print(f"[U27B3E6B1] Source SHA256: {sha256(source)}")
    print(f"[U27B3E6B1] Paragraphs: {len(doc.paragraphs)}")
    print(f"[U27B3E6B1] Tables: {len(doc.tables)}")
    print(f"[U27B3E6B1] Sections: {len(doc.sections)}")
    print(f"[U27B3E6B1] Embedded media: {len(media_rows)}")
    print(
        "[U27B3E6B1] Expected accessions present: "
        f"{int(accession_inventory['present'].sum())}/{len(EXPECTED_ACCESSIONS)}"
    )
    print(f"[U27B3E6B1] Report: {report}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3E6B1] ERROR: {exc}", file=sys.stderr)
        raise
