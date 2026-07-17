#!/usr/bin/env python3
"""
Phase U27B3E6B2
Insert the confirmed UTI HostOmics authorship, affiliations, correspondence,
declarations, contributions, acknowledgements, and availability statements into
the accession-corrected v6.3 DOCX.

Key safeguards
--------------
- Requires the UTI HostOmics project root and the locked U27B3E6A intake TSV.
- Refuses to proceed if any required intake field is not CONFIRMED/DEFERRED as
  expected.
- Removes the incorrect KIDH/UVA Fogarty D43 funding statement that was carried
  over from a different project and replaces it with the confirmed no-funding
  statement.
- Preserves the source DOCX, embedded figure bytes, and field-code counts.
- Renders the new DOCX to PDF and page PNGs with LibreOffice/pdftoppm.
- Produces a contact sheet for manual visual QA.
- Does not claim final manuscript readiness: references/Zotero, repository URL,
  and target-journal formatting remain deferred.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont


VERSION = "U27B3E6B2_v1.0_2026-07-17"
TAG = "phaseU27B3E6B2_targeted_front_matter_insertion"

EXPECTED_ROOT = Path("__UTI_HOSTOMICS_PROJECT_ROOT__")

SOURCE_REL = Path(
    "09_manuscript_docx/"
    "phaseU27B3E22_targeted_accession_correction/"
    "UTI_HostOmics_preZotero_manuscript_v6_3_"
    "U27B3E22_accession_corrected.docx"
)

INTAKE_REL = Path(
    "00_admin/phaseU27B3E6A_front_matter_intake/"
    "UTI_HostOmics_U27B3E6A_front_matter_intake.tsv"
)

OUTPUT_NAME = (
    "UTI_HostOmics_preZotero_manuscript_v6_4_"
    "U27B3E6B2_front_matter_complete.docx"
)

EXPECTED_ACCESSIONS = [
    "GSE112098",
    "GSE186800",
    "GSE252321",
    "GSE280297",
]

REQUIRED_CONFIRMED_FIELDS = [
    "manuscript_project",
    "author_order",
    "author_affiliation_mapping",
    "corresponding_author_name",
    "corresponding_author_email",
    "corresponding_author_postal_address",
    "funding_statement",
    "competing_interests_statement",
    "author_contributions_CRediT",
    "acknowledgements",
    "data_availability_statement",
]

REQUIRED_DEFERRED_FIELDS = [
    "code_availability_statement",
    "public_repository_URL",
    "target_journal",
]

DECLARATION_HEADINGS = {
    "Data availability": "data_availability_statement",
    "Code availability": "code_availability_statement",
    "Author contributions": "author_contributions_CRediT",
    "Competing interests": "competing_interests_statement",
    "Funding": "funding_statement",
    "Acknowledgements": "acknowledgements",
}

CROSS_PROJECT_FUNDING_MARKERS = [
    "Global Infectious Diseases Research Training Program",
    "Kibong'oto Infectious Diseases Hospital",
    "Kibong’oto Infectious Diseases Hospital",
    "University of Virginia",
    "D43 TW012247",
    "D43TW012247",
]


def log(message: str) -> None:
    print(f"[U27B3E6B2] {message}", flush=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def clean_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def load_intake(path: Path) -> Dict[str, Dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle, delimiter="\t"))

    required_columns = {"field", "status", "value"}
    if not rows:
        raise RuntimeError(f"Front-matter intake is empty: {path}")
    if not required_columns.issubset(rows[0]):
        raise RuntimeError(
            "Front-matter intake lacks required columns: "
            + ", ".join(sorted(required_columns))
        )

    result: Dict[str, Dict[str, str]] = {}
    for row in rows:
        field = clean_text(row.get("field"))
        if not field:
            continue
        if field in result:
            raise RuntimeError(f"Duplicate intake field: {field}")
        result[field] = {
            "status": clean_text(row.get("status")),
            "value": clean_text(row.get("value")),
        }

    missing = sorted(
        set(REQUIRED_CONFIRMED_FIELDS + REQUIRED_DEFERRED_FIELDS)
        - set(result)
    )
    if missing:
        raise RuntimeError(
            "Missing required intake fields: " + ", ".join(missing)
        )

    for field in REQUIRED_CONFIRMED_FIELDS:
        record = result[field]
        if record["status"] != "CONFIRMED" or not record["value"]:
            raise RuntimeError(
                f"Field must be CONFIRMED and nonblank: {field}"
            )

    for field in REQUIRED_DEFERRED_FIELDS:
        record = result[field]
        if record["status"] != "DEFERRED" or not record["value"]:
            raise RuntimeError(
                f"Field must be DEFERRED and nonblank: {field}"
            )

    if result["manuscript_project"]["value"] != "UTI HostOmics":
        raise RuntimeError(
            "Front-matter intake does not identify UTI HostOmics."
        )

    return result


def style_name(paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def is_heading(paragraph) -> bool:
    return style_name(paragraph).lower().startswith("heading")


def find_heading(doc: Document, heading_text: str):
    wanted = clean_text(heading_text).lower()
    matches = [
        paragraph
        for paragraph in doc.paragraphs
        if clean_text(paragraph.text).lower() == wanted
        and is_heading(paragraph)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected exactly one heading {heading_text!r}; "
            f"found {len(matches)}."
        )
    return matches[0]


def next_heading_after(doc: Document, paragraph):
    paragraphs = doc.paragraphs
    try:
        index = paragraphs.index(paragraph)
    except ValueError as exc:
        raise RuntimeError("Paragraph is not part of the document.") from exc

    for candidate in paragraphs[index + 1 :]:
        if is_heading(candidate) and clean_text(candidate.text):
            return candidate
    return None


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def remove_block_between(doc: Document, start_paragraph, end_paragraph) -> int:
    paragraphs = doc.paragraphs
    start_index = paragraphs.index(start_paragraph)
    end_index = paragraphs.index(end_paragraph)
    targets = list(paragraphs[start_index + 1 : end_index])
    for paragraph in targets:
        remove_paragraph(paragraph)
    return len(targets)


def set_default_run_font(run, name: str = "Arial", size: float = 10.0) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    rfonts = run._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def set_paragraph_spacing(
    paragraph,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def insert_before(reference_paragraph, style: str = "Normal"):
    paragraph = reference_paragraph.insert_paragraph_before()
    try:
        paragraph.style = style
    except Exception:
        pass
    return paragraph


def parse_author_order(value: str) -> List[Tuple[str, str]]:
    authors: List[Tuple[str, str]] = []
    for item in [part.strip() for part in value.split(";") if part.strip()]:
        match = re.match(r"^(.*?)(\d+(?:,\d+)*\*?)$", item)
        if not match:
            raise RuntimeError(
                "Could not parse author-affiliation suffix from: "
                f"{item!r}"
            )
        name = clean_text(match.group(1))
        suffix = clean_text(match.group(2))
        authors.append((name, suffix))
    if len(authors) != 5:
        raise RuntimeError(
            f"Expected five authors in author_order; found {len(authors)}."
        )
    return authors


def parse_affiliations(value: str) -> List[Tuple[str, str]]:
    affiliations: List[Tuple[str, str]] = []
    for item in [part.strip() for part in value.split(";") if part.strip()]:
        match = re.match(r"^(\d+)\s+(.+)$", item)
        if not match:
            raise RuntimeError(f"Could not parse affiliation: {item!r}")
        affiliations.append(
            (clean_text(match.group(1)), clean_text(match.group(2)))
        )
    observed_numbers = [number for number, _ in affiliations]
    if observed_numbers != ["1", "2", "3", "4", "5"]:
        raise RuntimeError(
            "Affiliation mapping must contain affiliations 1 through 5 "
            f"in order; observed {observed_numbers}."
        )
    return affiliations


def add_author_line(paragraph, authors: Sequence[Tuple[str, str]]) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=2, after=4, line=1.0)

    for index, (name, suffix) in enumerate(authors):
        if index:
            run = paragraph.add_run(", ")
            set_default_run_font(run, size=10.5)
        name_run = paragraph.add_run(name)
        name_run.bold = True
        set_default_run_font(name_run, size=10.5)

        suffix_run = paragraph.add_run(suffix)
        suffix_run.font.superscript = True
        suffix_run.bold = True
        set_default_run_font(suffix_run, size=8.5)


def add_affiliation_line(paragraph, number: str, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=1, line=1.0)

    number_run = paragraph.add_run(number)
    number_run.font.superscript = True
    number_run.bold = True
    set_default_run_font(number_run, size=8.0)

    text_run = paragraph.add_run(text)
    set_default_run_font(text_run, size=9.0)


def add_correspondence_lines(
    before_paragraph,
    name: str,
    address: str,
    emails: str,
) -> None:
    p1 = insert_before(before_paragraph)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p1, before=5, after=0, line=1.0)
    label = p1.add_run("*Corresponding author: ")
    label.bold = True
    set_default_run_font(label, size=9.0)
    value = p1.add_run(name)
    set_default_run_font(value, size=9.0)

    p2 = insert_before(before_paragraph)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, before=0, after=0, line=1.0)
    run = p2.add_run(address)
    set_default_run_font(run, size=9.0)

    p3 = insert_before(before_paragraph)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p3, before=0, after=6, line=1.0)
    label = p3.add_run("Email: ")
    label.bold = True
    set_default_run_font(label, size=9.0)
    value = p3.add_run(emails.replace(";", ","))
    set_default_run_font(value, size=9.0)


def replace_authors_block(
    doc: Document,
    intake: Dict[str, Dict[str, str]],
) -> Dict[str, object]:
    heading = find_heading(doc, "Authors")
    next_heading = next_heading_after(doc, heading)
    if next_heading is None:
        raise RuntimeError("No heading follows the Authors section.")

    removed = remove_block_between(doc, heading, next_heading)

    authors = parse_author_order(intake["author_order"]["value"])
    affiliations = parse_affiliations(
        intake["author_affiliation_mapping"]["value"]
    )

    author_paragraph = insert_before(next_heading)
    add_author_line(author_paragraph, authors)

    for number, text in affiliations:
        paragraph = insert_before(next_heading)
        add_affiliation_line(paragraph, number, text)

    add_correspondence_lines(
        next_heading,
        intake["corresponding_author_name"]["value"],
        intake["corresponding_author_postal_address"]["value"],
        intake["corresponding_author_email"]["value"],
    )

    return {
        "section": "Authors",
        "paragraphs_removed": removed,
        "authors_inserted": len(authors),
        "affiliations_inserted": len(affiliations),
        "correspondence_inserted": True,
    }


def replace_section_body(
    doc: Document,
    heading_text: str,
    body_text: str,
) -> Dict[str, object]:
    heading = find_heading(doc, heading_text)
    next_heading = next_heading_after(doc, heading)
    if next_heading is None:
        raise RuntimeError(
            f"No following heading found after {heading_text!r}."
        )

    removed = remove_block_between(doc, heading, next_heading)
    paragraph = insert_before(next_heading)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, before=0, after=6, line=1.15)

    run = paragraph.add_run(body_text)
    set_default_run_font(run, size=10.0)

    return {
        "section": heading_text,
        "paragraphs_removed": removed,
        "paragraphs_inserted": 1,
        "inserted_text": body_text,
    }


def zip_member_bytes(docx_path: Path, member: str) -> bytes:
    with zipfile.ZipFile(docx_path) as archive:
        return archive.read(member)


def media_inventory(docx_path: Path) -> pd.DataFrame:
    rows = []
    with zipfile.ZipFile(docx_path) as archive:
        for name in sorted(archive.namelist()):
            if not name.startswith("word/media/") or name.endswith("/"):
                continue
            data = archive.read(name)
            rows.append(
                {
                    "member": name,
                    "size_bytes": len(data),
                    "sha256": hashlib.sha256(data).hexdigest(),
                }
            )
    return pd.DataFrame(rows)


def ooxml_count(docx_path: Path, token: bytes) -> int:
    with zipfile.ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")
    return document_xml.count(token)


def count_occurrences(doc: Document, text: str) -> int:
    body = "\n".join(clean_text(p.text) for p in doc.paragraphs)
    return len(re.findall(re.escape(text), body, flags=re.IGNORECASE))


def find_command(candidates: Sequence[str]) -> Optional[str]:
    for candidate in candidates:
        path = shutil.which(candidate)
        if path:
            return path
    return None


def render_docx(
    docx_path: Path,
    render_dir: Path,
) -> Dict[str, object]:
    render_dir.mkdir(parents=True, exist_ok=True)

    office = find_command(["libreoffice", "soffice"])
    pdftoppm = find_command(["pdftoppm"])

    if not office:
        return {
            "render_pass": False,
            "reason": "LibreOffice/soffice not found",
            "pdf_path": "",
            "page_pngs": 0,
        }
    if not pdftoppm:
        return {
            "render_pass": False,
            "reason": "pdftoppm not found",
            "pdf_path": "",
            "page_pngs": 0,
        }

    profile = render_dir / ".lo_profile"
    home = render_dir / ".lo_home"
    profile.mkdir(parents=True, exist_ok=True)
    home.mkdir(parents=True, exist_ok=True)

    env = os.environ.copy()
    env["HOME"] = str(home)

    convert_command = [
        office,
        "--headless",
        f"-env:UserInstallation=file://{profile}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(render_dir),
        str(docx_path),
    ]
    convert = subprocess.run(
        convert_command,
        text=True,
        capture_output=True,
        env=env,
        timeout=300,
    )

    pdf_path = render_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        return {
            "render_pass": False,
            "reason": (
                "LibreOffice did not produce a nonempty PDF. "
                f"stdout={convert.stdout!r}; stderr={convert.stderr!r}"
            ),
            "pdf_path": str(pdf_path),
            "page_pngs": 0,
        }

    prefix = render_dir / "page"
    raster_command = [
        pdftoppm,
        "-png",
        "-r",
        "120",
        str(pdf_path),
        str(prefix),
    ]
    raster = subprocess.run(
        raster_command,
        text=True,
        capture_output=True,
        timeout=300,
    )
    page_paths = sorted(
        render_dir.glob("page-*.png"),
        key=lambda path: int(re.search(r"(\d+)$", path.stem).group(1)),
    )

    if raster.returncode != 0 or not page_paths:
        return {
            "render_pass": False,
            "reason": (
                "pdftoppm did not produce page PNGs. "
                f"stdout={raster.stdout!r}; stderr={raster.stderr!r}"
            ),
            "pdf_path": str(pdf_path),
            "page_pngs": len(page_paths),
        }

    return {
        "render_pass": True,
        "reason": "",
        "pdf_path": str(pdf_path),
        "page_pngs": len(page_paths),
        "page_paths": [str(path) for path in page_paths],
    }


def make_contact_sheet(
    page_paths: Sequence[Path],
    output_path: Path,
    columns: int = 4,
    thumb_width: int = 300,
) -> None:
    if not page_paths:
        raise RuntimeError("No page PNGs available for contact sheet.")

    thumbs: List[Tuple[Image.Image, str]] = []
    font = ImageFont.load_default()

    for index, path in enumerate(page_paths, start=1):
        with Image.open(path) as image:
            rgb = image.convert("RGB")
            ratio = thumb_width / rgb.width
            thumb = rgb.resize(
                (thumb_width, max(1, int(rgb.height * ratio))),
                Image.Resampling.LANCZOS,
            )
            thumbs.append((thumb, f"Page {index}"))

    label_height = 24
    cell_height = max(thumb.height for thumb, _ in thumbs) + label_height
    rows = (len(thumbs) + columns - 1) // columns
    margin = 12
    sheet = Image.new(
        "RGB",
        (
            columns * thumb_width + (columns + 1) * margin,
            rows * cell_height + (rows + 1) * margin,
        ),
        "white",
    )
    draw = ImageDraw.Draw(sheet)

    for index, (thumb, label) in enumerate(thumbs):
        row = index // columns
        column = index % columns
        x = margin + column * (thumb_width + margin)
        y = margin + row * (cell_height + margin)
        sheet.paste(thumb, (x, y + label_height))
        draw.text((x + 4, y + 4), label, fill="black", font=font)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root",
        default=str(EXPECTED_ROOT),
    )
    args = parser.parse_args()

    root = Path(args.project_root).resolve()
    if root != EXPECTED_ROOT:
        raise RuntimeError(
            f"Wrong project root: {root}; expected {EXPECTED_ROOT}"
        )

    source = root / SOURCE_REL
    intake_path = root / INTAKE_REL
    if not source.exists():
        raise FileNotFoundError(f"Source DOCX not found: {source}")
    if not intake_path.exists():
        raise FileNotFoundError(
            f"Front-matter intake not found: {intake_path}"
        )

    output_dir = root / "09_manuscript_docx" / TAG
    render_dir = output_dir / "render_qa"
    table_dir = root / "06_tables" / TAG
    metadata_dir = root / "03_metadata" / TAG
    results_dir = root / "05_results" / TAG
    log_dir = root / "08_logs" / TAG

    for directory in (
        output_dir,
        render_dir,
        table_dir,
        metadata_dir,
        results_dir,
        log_dir,
    ):
        directory.mkdir(parents=True, exist_ok=True)

    output_docx = output_dir / OUTPUT_NAME
    source_hash_before = sha256(source)
    intake = load_intake(intake_path)

    source_media = media_inventory(source)
    source_field_chars = ooxml_count(source, b"<w:fldChar")
    source_instr_text = ooxml_count(source, b"<w:instrText")

    doc = Document(str(source))

    edit_rows: List[Dict[str, object]] = []
    edit_rows.append(replace_authors_block(doc, intake))

    for heading, intake_field in DECLARATION_HEADINGS.items():
        edit_rows.append(
            replace_section_body(
                doc,
                heading,
                intake[intake_field]["value"],
            )
        )

    # Save to a temporary file first, then atomically promote.
    temp_output = output_docx.with_suffix(".tmp.docx")
    if temp_output.exists():
        temp_output.unlink()
    doc.save(str(temp_output))
    temp_output.replace(output_docx)

    if sha256(source) != source_hash_before:
        raise RuntimeError("The source DOCX changed during the run.")

    output_media = media_inventory(output_docx)
    output_field_chars = ooxml_count(output_docx, b"<w:fldChar")
    output_instr_text = ooxml_count(output_docx, b"<w:instrText")

    media_merge = source_media.merge(
        output_media,
        on="member",
        how="outer",
        suffixes=("_source", "_output"),
        indicator=True,
    )
    media_merge["byte_identical"] = (
        (media_merge["_merge"] == "both")
        & (
            media_merge["size_bytes_source"]
            == media_merge["size_bytes_output"]
        )
        & (
            media_merge["sha256_source"]
            == media_merge["sha256_output"]
        )
    )
    media_pass = bool(
        len(source_media) == len(output_media)
        and not media_merge.empty
        and media_merge["byte_identical"].all()
    )

    output_doc = Document(str(output_docx))
    body_text = "\n".join(
        clean_text(paragraph.text)
        for paragraph in output_doc.paragraphs
    )

    required_text_rows = []
    for field in (
        "corresponding_author_name",
        "funding_statement",
        "competing_interests_statement",
        "author_contributions_CRediT",
        "acknowledgements",
        "code_availability_statement",
        "data_availability_statement",
    ):
        value = intake[field]["value"]
        present = value in body_text
        required_text_rows.append(
            {
                "field": field,
                "required_value": value,
                "present_exactly": present,
                "occurrences": body_text.count(value),
            }
        )

    for author_name in [
        "Reuben S. Maghembe",
        "Samweli Bahati",
        "Abdalah Makaranga",
        "Maximilian A.K. Magulye",
        "Benson R. Kidenya",
    ]:
        required_text_rows.append(
            {
                "field": f"author::{author_name}",
                "required_value": author_name,
                "present_exactly": author_name in body_text,
                "occurrences": body_text.count(author_name),
            }
        )

    text_audit = pd.DataFrame(required_text_rows)

    accession_rows = []
    for accession in EXPECTED_ACCESSIONS:
        count = len(
            re.findall(
                rf"\b{re.escape(accession)}\b",
                body_text,
            )
        )
        accession_rows.append(
            {
                "accession": accession,
                "occurrences": count,
                "present": count > 0,
            }
        )
    accession_audit = pd.DataFrame(accession_rows)

    cross_project_rows = []
    for marker in CROSS_PROJECT_FUNDING_MARKERS:
        count = body_text.lower().count(marker.lower())
        cross_project_rows.append(
            {
                "marker": marker,
                "occurrences": count,
                "absent": count == 0,
            }
        )
    cross_project_audit = pd.DataFrame(cross_project_rows)

    structural_rows = [
        {
            "audit_id": "source_hash_unchanged",
            "pass": sha256(source) == source_hash_before,
            "observed": sha256(source),
            "expected": source_hash_before,
        },
        {
            "audit_id": "embedded_media_byte_identical",
            "pass": media_pass,
            "observed": (
                f"source={len(source_media)}; output={len(output_media)}; "
                f"identical={int(media_merge['byte_identical'].sum())}"
            ),
            "expected": (
                f"{len(source_media)} output media files, all byte-identical"
            ),
        },
        {
            "audit_id": "field_char_count_preserved",
            "pass": source_field_chars == output_field_chars,
            "observed": output_field_chars,
            "expected": source_field_chars,
        },
        {
            "audit_id": "instr_text_count_preserved",
            "pass": source_instr_text == output_instr_text,
            "observed": output_instr_text,
            "expected": source_instr_text,
        },
        {
            "audit_id": "confirmed_text_present",
            "pass": bool(text_audit["present_exactly"].all()),
            "observed": int(text_audit["present_exactly"].sum()),
            "expected": len(text_audit),
        },
        {
            "audit_id": "all_four_accessions_present",
            "pass": bool(accession_audit["present"].all()),
            "observed": int(accession_audit["present"].sum()),
            "expected": len(EXPECTED_ACCESSIONS),
        },
        {
            "audit_id": "cross_project_funding_markers_absent",
            "pass": bool(cross_project_audit["absent"].all()),
            "observed": int(cross_project_audit["absent"].sum()),
            "expected": len(cross_project_audit),
        },
        {
            "audit_id": "source_docx_not_modified",
            "pass": sha256(source) == source_hash_before,
            "observed": False,
            "expected": False,
        },
    ]
    structural_audit = pd.DataFrame(structural_rows)

    render_result = render_docx(output_docx, render_dir)
    contact_sheet = (
        render_dir
        / "UTI_HostOmics_U27B3E6B2_render_contact_sheet.png"
    )
    contact_sheet_created = False
    if render_result.get("render_pass"):
        page_paths = [
            Path(path)
            for path in render_result.get("page_paths", [])
        ]
        make_contact_sheet(page_paths, contact_sheet)
        contact_sheet_created = contact_sheet.exists()

    structural_pass = bool(structural_audit["pass"].all())
    render_pass = bool(render_result.get("render_pass"))
    decision = (
        "READY_FOR_U27B3E6B21_MANUAL_VISUAL_QA_AND_REFERENCE_FINALIZATION"
        if structural_pass and render_pass and contact_sheet_created
        else "TARGETED_U27B3E6B2_FRONT_MATTER_OR_RENDER_REPAIR_REQUIRED"
    )

    edit_audit = pd.DataFrame(edit_rows)
    edit_audit.to_csv(
        table_dir / "UTI_HostOmics_U27B3E6B2_edit_audit.tsv",
        sep="\t",
        index=False,
    )
    text_audit.to_csv(
        table_dir / "UTI_HostOmics_U27B3E6B2_required_text_audit.tsv",
        sep="\t",
        index=False,
    )
    accession_audit.to_csv(
        table_dir / "UTI_HostOmics_U27B3E6B2_accession_audit.tsv",
        sep="\t",
        index=False,
    )
    cross_project_audit.to_csv(
        table_dir
        / "UTI_HostOmics_U27B3E6B2_cross_project_funding_marker_audit.tsv",
        sep="\t",
        index=False,
    )
    media_merge.to_csv(
        table_dir / "UTI_HostOmics_U27B3E6B2_embedded_media_audit.tsv",
        sep="\t",
        index=False,
    )
    structural_audit.to_csv(
        table_dir / "UTI_HostOmics_U27B3E6B2_structural_audit.tsv",
        sep="\t",
        index=False,
    )

    render_audit = pd.DataFrame(
        [
            {
                "render_pass": render_pass,
                "reason": render_result.get("reason", ""),
                "pdf_path": render_result.get("pdf_path", ""),
                "page_pngs": render_result.get("page_pngs", 0),
                "contact_sheet_path": (
                    str(contact_sheet)
                    if contact_sheet_created
                    else ""
                ),
                "contact_sheet_created": contact_sheet_created,
                "manual_visual_qa_required": True,
            }
        ]
    )
    render_audit.to_csv(
        table_dir / "UTI_HostOmics_U27B3E6B2_render_audit.tsv",
        sep="\t",
        index=False,
    )

    decision_frame = pd.DataFrame(
        [
            {
                "phase": "U27B3E6B2",
                "decision": decision,
                "source_docx": str(source),
                "output_docx": str(output_docx),
                "source_sha256": source_hash_before,
                "output_sha256": sha256(output_docx),
                "structural_checks": len(structural_audit),
                "structural_checks_passed": int(
                    structural_audit["pass"].sum()
                ),
                "embedded_media_files": len(output_media),
                "embedded_media_byte_identical": media_pass,
                "field_char_count_preserved": (
                    source_field_chars == output_field_chars
                ),
                "expected_accessions_present": int(
                    accession_audit["present"].sum()
                ),
                "expected_accessions_total": len(EXPECTED_ACCESSIONS),
                "cross_project_funding_markers_absent": bool(
                    cross_project_audit["absent"].all()
                ),
                "render_pass": render_pass,
                "rendered_pages": render_result.get("page_pngs", 0),
                "contact_sheet_created": contact_sheet_created,
                "source_docx_modified": False,
                "scientific_values_recalculated": False,
                "references_finalized": False,
                "repository_url_finalized": False,
                "target_journal_finalized": False,
                "next_phase": (
                    "U27B3E6B2.1 manually inspect every rendered page; "
                    "then proceed to reference-gap resolution and Zotero finalization"
                ),
            }
        ]
    )
    decision_frame.to_csv(
        table_dir / "UTI_HostOmics_U27B3E6B2_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    metadata = pd.DataFrame(
        [
            {
                "version": VERSION,
                "project_root": str(root),
                "intake_file": str(intake_path),
                "source_docx": str(source),
                "source_sha256": source_hash_before,
                "output_docx": str(output_docx),
                "output_sha256": sha256(output_docx),
                "source_size_bytes": source.stat().st_size,
                "output_size_bytes": output_docx.stat().st_size,
                "source_media_files": len(source_media),
                "output_media_files": len(output_media),
                "source_field_chars": source_field_chars,
                "output_field_chars": output_field_chars,
                "source_instr_text": source_instr_text,
                "output_instr_text": output_instr_text,
                "source_docx_modified": False,
            }
        ]
    )
    metadata.to_csv(
        metadata_dir
        / "UTI_HostOmics_U27B3E6B2_document_provenance.tsv",
        sep="\t",
        index=False,
    )

    report_path = (
        results_dir
        / "UTI_HostOmics_U27B3E6B2_front_matter_insertion_report.md"
    )
    report_lines = [
        "# Phase U27B3E6B2 - Targeted front-matter insertion",
        "",
        f"- Version: `{VERSION}`",
        f"- Decision: **{decision}**",
        f"- Source: `{source}`",
        f"- Output: `{output_docx}`",
        f"- Source SHA256: `{source_hash_before}`",
        f"- Output SHA256: `{sha256(output_docx)}`",
        f"- Structural checks passed: **{int(structural_audit['pass'].sum())}/{len(structural_audit)}**",
        f"- Embedded media byte-identical: **{media_pass}**",
        f"- Expected GEO accessions retained: **{int(accession_audit['present'].sum())}/{len(EXPECTED_ACCESSIONS)}**",
        f"- Cross-project funding markers absent: **{bool(cross_project_audit['absent'].all())}**",
        f"- Render pass: **{render_pass}**",
        f"- Rendered pages: **{render_result.get('page_pngs', 0)}**",
        f"- Contact sheet created: **{contact_sheet_created}**",
        "",
        "## Inserted and replaced content",
        "",
        "- Five-author line with affiliation superscripts.",
        "- Five institutional affiliations.",
        "- Corresponding-author name, institutional address and two email addresses.",
        "- Data availability statement with GSE112098, GSE186800, GSE252321 and GSE280297.",
        "- Provisional code-availability statement pending the final GitHub URL.",
        "- Confirmed CRediT-style author contributions.",
        "- Competing-interests, funding and acknowledgements statements.",
        "",
        "## Cross-project correction",
        "",
        "The prior KIDH/UVA Fogarty D43 funding statement belonged to a different project and has been removed from the UTI HostOmics manuscript. The confirmed UTI statement is that no specific funding was received for this study.",
        "",
        "## Remaining blockers",
        "",
        "- Manual visual inspection of every rendered page.",
        "- Reference-gap resolution and Zotero finalization.",
        "- Public GitHub repository creation and insertion of its final URL.",
        "- Target-journal selection and journal-specific formatting.",
        "",
        "## Integrity boundary",
        "",
        "No scientific values were recalculated. The source DOCX was not overwritten. Embedded figure bytes and field-code counts were audited for preservation.",
        "",
    ]
    report_path.write_text(
        "\n".join(report_lines),
        encoding="utf-8",
    )

    run_manifest = {
        "version": VERSION,
        "decision": decision,
        "source_docx": str(source),
        "output_docx": str(output_docx),
        "source_sha256": source_hash_before,
        "output_sha256": sha256(output_docx),
        "structural_checks": int(len(structural_audit)),
        "structural_checks_passed": int(
            structural_audit["pass"].sum()
        ),
        "embedded_media_files": int(len(output_media)),
        "embedded_media_byte_identical": bool(media_pass),
        "field_char_count_preserved": bool(
            source_field_chars == output_field_chars
        ),
        "instr_text_count_preserved": bool(
            source_instr_text == output_instr_text
        ),
        "expected_accessions_present": int(
            accession_audit["present"].sum()
        ),
        "expected_accessions_total": int(len(EXPECTED_ACCESSIONS)),
        "cross_project_funding_markers_absent": bool(
            cross_project_audit["absent"].all()
        ),
        "render_pass": bool(render_pass),
        "rendered_pages": int(render_result.get("page_pngs", 0)),
        "contact_sheet_created": bool(contact_sheet_created),
        "contact_sheet_path": (
            str(contact_sheet) if contact_sheet_created else ""
        ),
        "source_docx_modified": False,
        "scientific_values_recalculated": False,
        "references_finalized": False,
        "repository_url_finalized": False,
        "target_journal_finalized": False,
    }
    (
        results_dir / "UTI_HostOmics_U27B3E6B2_run_manifest.json"
    ).write_text(
        json.dumps(run_manifest, indent=2, sort_keys=True),
        encoding="utf-8",
    )

    log(f"Decision: {decision}")
    log(
        "Structural checks passed: "
        f"{int(structural_audit['pass'].sum())}/{len(structural_audit)}"
    )
    log(f"Embedded media byte-identical: {media_pass}")
    log(
        "Cross-project funding markers absent: "
        f"{bool(cross_project_audit['absent'].all())}"
    )
    log(f"Render pass: {render_pass}")
    log(f"Rendered pages: {render_result.get('page_pngs', 0)}")
    log(f"Output DOCX: {output_docx}")
    log(f"Contact sheet: {contact_sheet if contact_sheet_created else 'NOT CREATED'}")
    log(f"Report: {report_path}")

    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3E6B2] ERROR: {exc}", file=sys.stderr)
        raise
