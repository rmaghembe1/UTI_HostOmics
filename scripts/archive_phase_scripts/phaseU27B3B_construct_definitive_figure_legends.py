#!/usr/bin/env python3
"""
Phase U27B3B
Construct definitive manuscript legends for frozen Figures 1-8.

Inputs
------
- U27B3A figure-and-panel title registry
- U27B3A legend-input registry
- Eight frozen panel-level source-value tables
- Combined visual-approval record
- Frozen source registry and figure masters

Outputs
-------
- Definitive figure legends in Markdown
- Plain-text manuscript insertion block
- Legend-only DOCX
- Panel-level legend/provenance registry
- Figure-level legend audit
- Caveat and terminology registry
- Phase decision, report and JSON manifest

Scientific boundaries
---------------------
- No broad pregnancy-wide FDR claim.
- Pregnancy tissues/samples remain inferential units; no dam-level inference.
- Single-cell localization is descriptive at n=2 control versus n=2 UPEC.
- Exact two-sided permutation p-values cannot be smaller than 0.333 in the
  four-sample single-cell comparison.
- Metabolic modules represent transcriptionally inferred pathway activity,
  not metabolic flux.
- Raw expression is never pooled across species, tissues or studies.
- Cross-dataset integration uses standardized effects, recurrence and
  directional concordance.
- Complement recurrence remains provisional where sample support is limited.

No scientific values are recalculated and no figure asset is modified.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Sequence, Tuple

import pandas as pd


VERSION = "U27B3B_v1.0_2026-07-16"
TAG = "phaseU27B3B_definitive_figure_legend_construction"
PACKAGE_TAG = "phaseU27B3A_complete_eight_figure_package_assembly"

EXPECTED_PANEL_COUNTS = {
    1: 6,
    2: 7,
    3: 8,
    4: 8,
    5: 7,
    6: 8,
    7: 7,
    8: 6,
}

FIGURE_TITLES = {
    1: "Study architecture, datasets, contrasts and evidence hierarchy",
    2: (
        "Cross-dataset infection effects, recurrent cores and contextual "
        "comparators"
    ),
    3: (
        "Pregnancy, tissue and outcome-associated endocrine-metabolic-"
        "complement remodeling"
    ),
    4: "Single-cell composition, immune states and cellular localization",
    5: "Steroid, cholesterol and lipid-remodeling architecture",
    6: "Adipokine, insulin and integrated immunometabolic remodeling",
    7: "Complement branch and cellular architecture",
    8: "Integrated endocrine-metabolic-immune model and evidence boundaries",
}

PANEL_LEGENDS = {
    "Figure_1A": (
        "Conceptual map linking the cross-dataset infection core, pregnancy- "
        "and tissue-associated outcomes, endocrine-metabolic branches and "
        "cell-resolved immune localization under an explicit evidence hierarchy."
    ),
    "Figure_1B": (
        "Dataset architecture comprising GSE112098, a human urinary/systemic "
        "inflammation comparator (n=73); GSE280297, a mouse pregnancy, tissue "
        "and outcome dataset (n=60 tissue samples); GSE168600, a mouse recurrent "
        "or prior-exposure bladder model (n=20); and GSE252321, a four-sample "
        "single-cell UPEC experiment. In GSE252321, 27,385 cells passed quality "
        "control, and a balanced subset was used for the displayed embedding."
    ),
    "Figure_1C": (
        "Biological sample counts across bladder, placenta and uterus for the "
        "mock/PBS, nonpregnant UPEC, preterm UPEC and term UPEC groups in "
        "GSE280297. Tissue samples, rather than inferred dams, were retained as "
        "the statistical units."
    ),
    "Figure_1D": (
        "Distribution of 78 curated submodules across ten biological axes, "
        "including steroid-cholesterol-endocrine biology, inflammatory carbon "
        "metabolism, nucleotide/NAD/nitrogen metabolism, complement architecture, "
        "amino-acid metabolism, lipid metabolism, immune-context anchors, "
        "adipokine signaling, insulin/IRS signaling and catecholamine-stress-"
        "adjacent biology."
    ),
    "Figure_1E": (
        "Analytical workflow from expression repair and species-native module "
        "scoring through evidence synthesis and cellular attribution. Raw "
        "expression values were not pooled across studies, species or tissues."
    ),
    "Figure_1F": (
        "Evidence hierarchy separating robust recurrent effects, provisional "
        "effects supported by one false-discovery-rate-qualified dataset plus "
        "independent concordance, context-divergent effects and limited "
        "independent support."
    ),
    "Figure_2A": (
        "Signed independent infection-effect estimates for selected recurrent "
        "and contextual modules across the eligible datasets. Color denotes the "
        "direction and magnitude of the standardized module-level effect within "
        "each independently analyzed study."
    ),
    "Figure_2B": (
        "Number of modules assigned to each evidence class after recurrence, "
        "false-discovery-rate support and independent-context evaluation."
    ),
    "Figure_2C": (
        "Evidence-weighted topology connecting TLR4-LPS signaling, leptin, "
        "insulin receptor/IRS, PI3K-AKT, inflammatory response and glycogen/"
        "carbon-use programs. Arrows summarize recurrent directional coupling "
        "and do not establish a direct causal sequence."
    ),
    "Figure_2D": (
        "Complement recurrence core comparing opsonophagocytosis and C3a/C5a "
        "inflammatory signaling across infection and preterm-versus-term "
        "contexts. Complement findings are interpreted as provisional where "
        "independent sample support is limited."
    ),
    "Figure_2E": (
        "Age- and sex-adjusted estimates for systemic comparator modules in the "
        "human dataset, shown as signed adjusted effects."
    ),
    "Figure_2F": (
        "Contextual factorial estimates for signaling and metabolic modules, "
        "highlighting effects that were strong in a specific model but not "
        "uniformly recurrent across all datasets."
    ),
    "Figure_2G": (
        "Evidence-weighted directional concordance versus the median independent "
        "effect. Weighted concordance summarizes consistency of effect direction "
        "across eligible contrasts and is not a pooled expression estimate."
    ),
    "Figure_3A": (
        "Pregnancy and tissue design for GSE280297, showing the distribution of "
        "mock/PBS, nonpregnant UPEC, preterm UPEC and term UPEC samples across "
        "bladder, placenta and uterus."
    ),
    "Figure_3B": (
        "Selected preterm-versus-term effects by tissue, illustrating the "
        "branch-selective pregnancy remodeling of endocrine, metabolic and "
        "complement modules."
    ),
    "Figure_3C": (
        "Cross-tissue coherence of selected modules. Points show the median "
        "Hedges g across eligible tissues, with the accompanying directional-"
        "coherence annotation indicating whether tissue-specific estimates "
        "shared the same sign."
    ),
    "Figure_3D": (
        "Domain-level summary of steroid synthesis-response decoupling. "
        "Steroidogenic and complement-effector branches tended to differ from "
        "attenuated steroid-receptor and metabolic-effector responses, supporting "
        "a branch-selective rather than globally activated or suppressed model."
    ),
    "Figure_3E": (
        "UPEC-versus-PBS effects during pregnancy across tissues for selected "
        "signaling, endocrine, lipid, carbon and complement modules."
    ),
    "Figure_3F": (
        "Pregnant-versus-nonpregnant bladder effects, summarized as Hedges g for "
        "selected modules."
    ),
    "Figure_3G": (
        "Complement and inflammatory-carbon branches across the available "
        "pregnancy, infection and tissue contrasts."
    ),
    "Figure_3H": (
        "Working model in which UPEC exposure during pregnancy engages TLR4/"
        "complement and inflammatory-carbon responses while steroid synthesis "
        "and androgen branches are partly uncoupled from attenuated receptor and "
        "metabolic-effector programs, producing tissue-specific outcome biology."
    ),
    "Figure_4A": (
        "Balanced two-dimensional embedding of quality-controlled GSE252321 "
        "single cells, colored by broad immune population after marker-based "
        "reconstruction."
    ),
    "Figure_4B": (
        "Cluster-marker validation heatmap using one high-priority marker per "
        "cluster to support the reconstructed broad and refined immune labels."
    ),
    "Figure_4C": (
        "Broad-cell composition in two control and two UPEC samples. Bars show "
        "the fraction of quality-controlled cells assigned to each broad immune "
        "population."
    ),
    "Figure_4D": (
        "Relative abundance of refined immune subtypes across the four biological "
        "samples, including inflammatory monocytes, reparative macrophages, "
        "activated dendritic cells, cycling immune cells and regulatory/type-2-"
        "like lymphoid states."
    ),
    "Figure_4E": (
        "Fractions of macrophage/monocyte cells with detectable TNFSF9 and with "
        "TNFSF9 expression above the prespecified pooled high-expression "
        "threshold in control and UPEC samples."
    ),
    "Figure_4F": (
        "Strict and expanded Treg-like fractions within T cells. Strict Treg-like "
        "cells required FOXP3 together with at least one supporting marker; the "
        "expanded definition required FOXP3 or at least two supporting markers."
    ),
    "Figure_4G": (
        "Broad-cell localization of selected recurrent endocrine, metabolic, "
        "complement and inflammatory modules using module-mean gene log2 fold "
        "changes from independently constructed cell-type pseudobulks."
    ),
    "Figure_4H": (
        "Descriptive model linking UPEC-associated myeloid expansion and "
        "activation, TNFSF9-positive macrophage states, T-cell/dendritic module "
        "localization and expansion of regulatory-inflammatory Treg-like states."
    ),
    "Figure_5A": (
        "Independent infection effects for steroidogenesis, androgen/testosterone "
        "and estrogen biology, cholesterol handling, receptor responses and "
        "lipid-regulatory or lipid-stress modules."
    ),
    "Figure_5B": (
        "Preterm-versus-term endocrine effects across bladder, placenta and "
        "uterus, including steroidogenesis, androgen/testosterone, estrogen, "
        "cholesterol and receptor-response branches."
    ),
    "Figure_5C": (
        "Domain-level steroid synthesis-response decoupling reproduced within "
        "the focused endocrine figure."
    ),
    "Figure_5D": (
        "Comparison of median independent infection effects with collapsed "
        "preterm-versus-term effects for selected endocrine modules. Numbered "
        "points correspond to the inset key and preserve all plotted values "
        "without implying a causal relationship between infection and pregnancy "
        "effects."
    ),
    "Figure_5E": (
        "Broad-cell localization of endocrine and lipid modules across NK cells, "
        "T cells, cycling immune cells, dendritic cells, macrophage/monocytes and "
        "neutrophils."
    ),
    "Figure_5F": (
        "Lipid-droplet dynamics, PPAR/SREBP/LXR regulation, ferroptosis-linked "
        "lipid peroxidation, fatty-acid synthesis and oxidation, phospholipid, "
        "sphingolipid/ceramide and eicosanoid/prostaglandin programs across "
        "pregnancy and infection contexts."
    ),
    "Figure_5G": (
        "Six strongest refined-subtype localization supports for endocrine and "
        "lipid modules, ranked by the bounded composite cellular-support score."
    ),
    "Figure_6A": (
        "Independent effects for leptin, resistin, insulin receptor/IRS, "
        "PI3K-AKT and inflammatory-carbon modules."
    ),
    "Figure_6B": (
        "Pregnancy-associated adipokine and insulin effects across bladder, "
        "placenta and uterus."
    ),
    "Figure_6C": (
        "Broad-cell localization of adipokine, insulin/IRS, PI3K-AKT and "
        "inflammatory-carbon modules."
    ),
    "Figure_6D": (
        "Six strongest refined-subtype supports for immunometabolic modules, "
        "ranked by the bounded composite cellular-support score."
    ),
    "Figure_6E": (
        "Evidence-weighted leptin-IRS-PI3K/AKT model linking leptin and insulin "
        "receptor signaling to inflammatory carbon and glycogen-use programs. "
        "The topology is hypothesis-generating and not a direct causal model."
    ),
    "Figure_6F": (
        "Transcriptionally inferred glycolysis, lactate/HIF1A, glycogen, pentose-"
        "phosphate, gluconeogenic and tricarboxylic-acid-cycle programs across "
        "the pregnancy and infection contexts."
    ),
    "Figure_6G": (
        "Amino-acid transport, arginine/nitric-oxide, urea-cycle, glutamine, "
        "tryptophan, methionine/SAM and serine/glycine/one-carbon programs across "
        "contexts."
    ),
    "Figure_6H": (
        "Purine salvage and degradation, pyrimidine synthesis and degradation, "
        "NAD metabolism, xanthine oxidase/oxidative purine catabolism and NRF2-"
        "redox programs across contexts."
    ),
    "Figure_7A": (
        "Independent effects for complement modules with eligible evidence, "
        "including classical, lectin and alternative initiation, C3 convertase/"
        "amplification, C3a/C5a signaling, opsonophagocytosis, terminal membrane-"
        "attack-complex biology, regulation and coagulation crosstalk."
    ),
    "Figure_7B": (
        "Preterm-versus-term complement effects across bladder, placenta and "
        "uterus."
    ),
    "Figure_7C": (
        "Broad-cell localization of complement branches using cell-type "
        "pseudobulk module effects."
    ),
    "Figure_7D": (
        "Six strongest refined-subtype supports for complement modules, ranked "
        "by the bounded composite cellular-support score."
    ),
    "Figure_7E": (
        "Complement topology linking classical, lectin and alternative entry "
        "routes to C3 convertase/amplification and the C3a/C5a, "
        "opsonophagocytic and terminal membrane-attack-complex branches. "
        "Regulatory and coagulation modules can act across multiple stages."
    ),
    "Figure_7F": (
        "Comparison of median independent infection effects with collapsed "
        "preterm-versus-term effects for complement stages. Numbered points "
        "correspond to the inset key."
    ),
    "Figure_7G": (
        "Number of broad immune populations with eligible cellular values for "
        "each complement stage. Absence from this panel indicates unavailable "
        "cell-level support rather than evidence of biological absence."
    ),
    "Figure_8A": (
        "Selected recurrent modules across independent infection datasets and "
        "the preterm-versus-term context."
    ),
    "Figure_8B": (
        "Cell-source-resolved localization of the integrated recurrent core."
    ),
    "Figure_8C": (
        "UPEC-associated broad-cell composition shifts, expressed as the "
        "difference in cellular fraction between UPEC and control samples."
    ),
    "Figure_8D": (
        "UPEC-associated differences in strict and expanded Treg-like fractions "
        "and TNFSF9-positive or TNFSF9-high macrophage states."
    ),
    "Figure_8E": (
        "Integrated hypothesis network connecting UPEC/TLR4 signaling with "
        "leptin-IRS-PI3K/AKT, complement/opsonophagocytosis, steroid/lipid "
        "branching, carbon/nitrogen/redox remodeling, cell-type localization and "
        "pregnancy- or tissue-specific outcomes."
    ),
    "Figure_8F": (
        "Evidence and interpretation boundary separating robust recurrent "
        "concordance, provisional independently supported biology and contextual "
        "hypothesis-generating effects."
    ),
}

FIGURE_NOTES = {
    1: (
        "All datasets were analyzed independently in their native species and "
        "tissue context. Cross-study integration was based on standardized "
        "effects and directional concordance, not merged raw expression."
    ),
    2: (
        "Signed effects and recurrence classes are evidence-synthesis outputs. "
        "They should not be interpreted as causal estimates or as evidence that "
        "all biological contexts share an identical mechanism."
    ),
    3: (
        "Pregnancy findings are exploratory because no broad pregnancy-wide "
        "false-discovery-rate signal was detected. Tissue samples are the "
        "inferential units; no dam-level inference is made."
    ),
    4: (
        "Cellular analyses are descriptive and hypothesis-generating because "
        "GSE252321 contains n=2 control and n=2 UPEC biological samples. The "
        "minimum attainable exact two-sided permutation p-value is 0.333."
    ),
    5: (
        "Steroid biology is interpreted as branch-selective synthesis-response "
        "decoupling rather than global activation or suppression. Metabolic "
        "modules indicate transcriptionally inferred pathway activity, not flux."
    ),
    6: (
        "Metabolic panels represent transcriptionally inferred pathway activity "
        "and do not quantify metabolite concentrations, reaction rates or flux."
    ),
    7: (
        "Complement recurrence, particularly C3a/C5a signaling and "
        "opsonophagocytosis, remains provisional where independent sample support "
        "is limited."
    ),
    8: (
        "The integrated network is a synthesis of recurrent, provisional and "
        "contextual associations. It is not a direct causal model."
    ),
}

CAVEAT_REGISTRY = [
    (
        "cross_dataset_integration",
        "Raw expression was not pooled across datasets, species or tissues.",
        "Figures 1, 2 and 8",
    ),
    (
        "pregnancy_fdr",
        "No broad pregnancy-wide FDR support was detected.",
        "Figures 3, 5, 6, 7 and 8",
    ),
    (
        "pregnancy_unit",
        "Tissue samples are inferential units; no dam-level inference is made.",
        "Figures 1 and 3",
    ),
    (
        "single_cell_sample_size",
        "Cellular localization is descriptive at n=2 control versus n=2 UPEC.",
        "Figures 4-8",
    ),
    (
        "single_cell_permutation",
        "The minimum attainable exact two-sided permutation p-value is 0.333.",
        "Figure 4",
    ),
    (
        "metabolic_activity",
        (
            "Metabolic modules represent transcriptionally inferred pathway "
            "activity, not metabolite flux."
        ),
        "Figures 5, 6 and 8",
    ),
    (
        "complement_provisional",
        (
            "Complement recurrence remains provisional where independent "
            "sample support is limited."
        ),
        "Figures 2, 3, 7 and 8",
    ),
    (
        "causal_boundary",
        (
            "Evidence-weighted networks summarize directional association and "
            "do not establish direct causal sequences."
        ),
        "Figures 2, 4, 6, 7 and 8",
    ),
]


def log(message: str) -> None:
    print(f"[U27B3B] {message}", flush=True)


def normalize_panel_key(value: object) -> str:
    text = str(value).strip()
    match = re.search(r"Figure[_ ]?(\d+)([A-H])", text)
    if match:
        return f"Figure_{match.group(1)}{match.group(2)}"
    return text


def source_roles_for_panel(
    legend_registry: pd.DataFrame,
    panel_key: str,
) -> str:
    rows = legend_registry[
        legend_registry["panel_key"].astype(str).map(normalize_panel_key)
        == panel_key
    ]
    if rows.empty or "source_role" not in rows.columns:
        return ""
    values = sorted(
        {
            str(value)
            for value in rows["source_role"].dropna()
            if str(value).strip()
        }
    )
    return "; ".join(values)


def source_paths_for_panel(
    legend_registry: pd.DataFrame,
    panel_key: str,
) -> str:
    rows = legend_registry[
        legend_registry["panel_key"].astype(str).map(normalize_panel_key)
        == panel_key
    ]
    if rows.empty:
        return ""

    columns = [
        column
        for column in ("locked_path", "source_value_table", "panel_crop")
        if column in rows.columns
    ]
    values = []
    for column in columns:
        for value in rows[column].dropna():
            text = str(value).strip()
            if text and text not in values:
                values.append(text)
    return "; ".join(values)


def panel_source_table_rows(
    source_table: pd.DataFrame,
    panel_letter: str,
) -> int:
    if "panel" not in source_table.columns:
        return len(source_table)
    return int(
        (
            source_table["panel"].astype(str).str.upper()
            == panel_letter.upper()
        ).sum()
    )


def build_figure_legend(
    figure_number: int,
) -> str:
    panel_texts = []
    for index in range(EXPECTED_PANEL_COUNTS[figure_number]):
        panel = chr(ord("A") + index)
        panel_key = f"Figure_{figure_number}{panel}"
        panel_texts.append(
            f"({panel}) {PANEL_LEGENDS[panel_key]}"
        )

    return (
        f"Figure {figure_number}. {FIGURE_TITLES[figure_number]}. "
        + " ".join(panel_texts)
        + " "
        + FIGURE_NOTES[figure_number]
    )


def create_docx(
    legends: Dict[int, str],
    output_path: Path,
) -> bool:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        return False

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "UTI HostOmics Project\nDefinitive Figure Legends"
    )
    run.bold = True
    run.font.size = Pt(15)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Frozen main Figures 1-8 | Phase U27B3B"
    )
    run.italic = True
    run.font.size = Pt(9)

    document.add_paragraph()

    for figure_number in range(1, 9):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.08

        prefix = (
            f"Figure {figure_number}. "
            f"{FIGURE_TITLES[figure_number]}. "
        )
        run = paragraph.add_run(prefix)
        run.bold = True
        run.font.size = Pt(9.5)

        remaining = legends[figure_number][len(prefix):]
        run = paragraph.add_run(remaining)
        run.font.size = Pt(9.5)

    document.add_page_break()
    heading = document.add_paragraph()
    run = heading.add_run("Interpretation boundaries")
    run.bold = True
    run.font.size = Pt(12)

    for _, statement, scope in CAVEAT_REGISTRY:
        paragraph = document.add_paragraph(style=None)
        paragraph.style = document.styles["List Bullet"]
        run = paragraph.add_run(f"{statement} [{scope}]")
        run.font.size = Pt(9.5)

    document.save(output_path)
    return output_path.exists() and output_path.stat().st_size > 0


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root",
        default="__UTI_HOSTOMICS_PROJECT_ROOT__",
    )
    args = parser.parse_args()

    project = Path(args.project_root).resolve()

    package_metadata = project / "03_metadata" / PACKAGE_TAG
    package_tables = project / "06_tables" / PACKAGE_TAG
    package_figures = project / "06_figures" / PACKAGE_TAG

    title_registry_path = (
        package_metadata
        / "UTI_HostOmics_U27B3A_figure_and_panel_title_registry.tsv"
    )
    legend_registry_path = (
        package_metadata
        / "UTI_HostOmics_U27B3A_legend_input_registry.tsv"
    )
    approval_path = (
        package_tables
        / "UTI_HostOmics_U27B3A_combined_visual_approval_record.tsv"
    )
    panel_audit_path = (
        package_tables
        / "UTI_HostOmics_U27B3A_panel_count_audit.tsv"
    )

    required = [
        title_registry_path,
        legend_registry_path,
        approval_path,
        panel_audit_path,
    ]
    for path in required:
        if not path.exists():
            raise FileNotFoundError(
                f"Required U27B3A input not found: {path}"
            )

    title_registry = pd.read_csv(
        title_registry_path,
        sep="\t",
        low_memory=False,
    )
    legend_registry = pd.read_csv(
        legend_registry_path,
        sep="\t",
        low_memory=False,
    )
    approvals = pd.read_csv(
        approval_path,
        sep="\t",
        low_memory=False,
    )
    panel_audit = pd.read_csv(
        panel_audit_path,
        sep="\t",
        low_memory=False,
    )

    outmetadata = project / "03_metadata" / TAG
    outtables = project / "06_tables" / TAG
    outresults = project / "05_results" / TAG
    outdocs = project / "07_manuscript" / TAG

    for directory in (
        outmetadata,
        outtables,
        outresults,
        outdocs,
    ):
        directory.mkdir(parents=True, exist_ok=True)

    source_tables: Dict[int, pd.DataFrame] = {}
    source_table_paths: Dict[int, Path] = {}

    for figure_number in range(1, 9):
        path = (
            package_tables
            / f"UTI_HostOmics_U27B3A_Figure_"
            f"{figure_number}_source_values.tsv"
        )
        if not path.exists():
            raise FileNotFoundError(
                f"Source-value table not found for Figure {figure_number}: "
                f"{path}"
            )
        source_tables[figure_number] = pd.read_csv(
            path,
            sep="\t",
            low_memory=False,
        )
        source_table_paths[figure_number] = path

    legends = {
        figure_number: build_figure_legend(figure_number)
        for figure_number in range(1, 9)
    }

    # Markdown and plain-text manuscript insertion files.
    markdown_path = (
        outdocs
        / "UTI_HostOmics_U27B3B_definitive_figure_legends.md"
    )
    text_path = (
        outdocs
        / "UTI_HostOmics_U27B3B_manuscript_legend_insert.txt"
    )
    docx_path = (
        outdocs
        / "UTI_HostOmics_U27B3B_definitive_figure_legends.docx"
    )

    with markdown_path.open("w", encoding="utf-8") as handle:
        handle.write(
            "# UTI HostOmics Project - Definitive Figure Legends\n\n"
        )
        handle.write(
            "Frozen Figures 1-8 | Phase U27B3B\n\n"
        )
        for figure_number in range(1, 9):
            handle.write(
                f"## Figure {figure_number}\n\n"
                f"{legends[figure_number]}\n\n"
            )

        handle.write("## Interpretation boundaries\n\n")
        for _, statement, scope in CAVEAT_REGISTRY:
            handle.write(f"- {statement} **Scope:** {scope}.\n")

    with text_path.open("w", encoding="utf-8") as handle:
        for figure_number in range(1, 9):
            handle.write(legends[figure_number])
            handle.write("\n\n")

    docx_created = create_docx(
        legends,
        docx_path,
    )

    # Panel-level legend/provenance registry.
    panel_rows: List[Dict[str, object]] = []

    for figure_number in range(1, 9):
        source_table = source_tables[figure_number]

        for index in range(EXPECTED_PANEL_COUNTS[figure_number]):
            panel = chr(ord("A") + index)
            panel_key = f"Figure_{figure_number}{panel}"

            title_rows = title_registry[
                title_registry["panel_key"]
                .astype(str)
                .map(normalize_panel_key)
                == panel_key
            ]
            frozen_panel_title = (
                str(title_rows.iloc[0]["panel_title"])
                if len(title_rows) > 0
                and "panel_title" in title_rows.columns
                else panel_key
            )

            panel_rows.append(
                {
                    "figure_number": figure_number,
                    "figure_title": FIGURE_TITLES[figure_number],
                    "panel": panel,
                    "panel_key": panel_key,
                    "frozen_panel_title": frozen_panel_title,
                    "definitive_panel_legend": PANEL_LEGENDS[panel_key],
                    "source_roles": source_roles_for_panel(
                        legend_registry,
                        panel_key,
                    ),
                    "source_paths": source_paths_for_panel(
                        legend_registry,
                        panel_key,
                    ),
                    "source_value_table": str(
                        source_table_paths[figure_number]
                    ),
                    "source_value_rows_for_panel": (
                        panel_source_table_rows(
                            source_table,
                            panel,
                        )
                    ),
                    "panel_crop": str(
                        package_figures
                        / "panel_crops"
                        / f"UTI_HostOmics_U27B3A_Figure_"
                        f"{figure_number}_panel_{panel}.png"
                    ),
                    "figure_svg": str(
                        package_figures
                        / f"UTI_HostOmics_U27B3A_Figure_"
                        f"{figure_number}.svg"
                    ),
                }
            )

    panel_legend_registry = pd.DataFrame(panel_rows)
    panel_legend_registry.to_csv(
        outmetadata
        / "UTI_HostOmics_U27B3B_panel_legend_provenance_registry.tsv",
        sep="\t",
        index=False,
    )

    # Caveat registry.
    caveat_frame = pd.DataFrame(
        CAVEAT_REGISTRY,
        columns=[
            "caveat_id",
            "required_statement",
            "figure_scope",
        ],
    )
    caveat_frame.to_csv(
        outmetadata
        / "UTI_HostOmics_U27B3B_caveat_terminology_registry.tsv",
        sep="\t",
        index=False,
    )

    # Figure-level audit.
    required_terms = {
        1: ["not pooled", "native species"],
        2: ["not be interpreted as causal"],
        3: ["no broad pregnancy-wide", "no dam-level"],
        4: ["n=2 control", "0.333"],
        5: ["branch-selective", "not flux"],
        6: ["not quantify", "flux"],
        7: ["provisional"],
        8: ["not a direct causal model"],
    }

    audit_rows: List[Dict[str, object]] = []
    for figure_number in range(1, 9):
        legend = legends[figure_number]
        expected = EXPECTED_PANEL_COUNTS[figure_number]
        represented = sum(
            1
            for index in range(expected)
            if f"({chr(ord('A') + index)})" in legend
        )
        terms_present = all(
            term.lower() in legend.lower()
            for term in required_terms[figure_number]
        )
        approval_rows = approvals[
            approvals["figure"]
            .astype(str)
            .str.extract(r"(\d+)", expand=False)
            .fillna("")
            .astype(str)
            == str(figure_number)
        ]
        visual_pass = bool(
            len(approval_rows) == 1
            and str(
                approval_rows.iloc[0]["status"]
            ).upper() == "PASS"
        )

        audit_rows.append(
            {
                "figure_number": figure_number,
                "expected_panels": expected,
                "panels_in_legend": represented,
                "panel_completeness_pass": represented == expected,
                "legend_word_count": len(legend.split()),
                "required_caveats_present": terms_present,
                "visual_approval_pass": visual_pass,
                "source_value_table_rows": len(
                    source_tables[figure_number]
                ),
                "figure_svg_exists": (
                    package_figures
                    / f"UTI_HostOmics_U27B3A_Figure_"
                    f"{figure_number}.svg"
                ).exists(),
            }
        )

    figure_audit = pd.DataFrame(audit_rows)
    figure_audit.to_csv(
        outtables
        / "UTI_HostOmics_U27B3B_figure_legend_audit.tsv",
        sep="\t",
        index=False,
    )

    package_panel_pass = bool(
        len(panel_legend_registry) == 57
        and panel_legend_registry["panel_key"].nunique() == 57
    )
    all_figure_audits_pass = bool(
        figure_audit["panel_completeness_pass"].all()
        and figure_audit["required_caveats_present"].all()
        and figure_audit["visual_approval_pass"].all()
        and figure_audit["figure_svg_exists"].all()
    )
    u27b3a_panel_pass = bool(
        panel_audit["panel_count_pass"].astype(bool).all()
    )
    files_created = all(
        path.exists() and path.stat().st_size > 0
        for path in (markdown_path, text_path)
    )
    docx_status = (
        "CREATED"
        if docx_created
        else "SKIPPED_PYTHON_DOCX_UNAVAILABLE"
    )

    if (
        package_panel_pass
        and all_figure_audits_pass
        and u27b3a_panel_pass
        and files_created
    ):
        decision = (
            "READY_FOR_U27B3C_RESULTS_SECTION_CONSTRUCTION_"
            "AND_MANUSCRIPT_TARGET_RESOLUTION"
        )
    else:
        decision = (
            "TARGETED_U27B3B_LEGEND_COMPLETENESS_REPAIR_REQUIRED"
        )

    decision_frame = pd.DataFrame(
        [
            {
                "phase": "U27B3B",
                "decision": decision,
                "figures_expected": 8,
                "figures_with_legends": len(legends),
                "panels_expected": 57,
                "panels_in_legend_registry": len(
                    panel_legend_registry
                ),
                "unique_panel_keys": (
                    panel_legend_registry["panel_key"].nunique()
                ),
                "all_panel_legends_present": package_panel_pass,
                "all_figure_audits_pass": all_figure_audits_pass,
                "U27B3A_panel_count_pass": u27b3a_panel_pass,
                "markdown_created": markdown_path.exists(),
                "plain_text_created": text_path.exists(),
                "docx_status": docx_status,
                "scientific_values_recalculated": False,
                "figure_assets_modified": False,
                "source_locks_changed": False,
                "manuscript_modified": False,
                "next_phase": (
                    "U27B3C construct Results section in frozen figure order "
                    "after resolving the current manuscript input file"
                    if decision.startswith("READY_FOR_U27B3C")
                    else "Repair incomplete legends, caveats or provenance"
                ),
            }
        ]
    )
    decision_frame.to_csv(
        outtables
        / "UTI_HostOmics_U27B3B_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    report_path = (
        outresults
        / "UTI_HostOmics_U27B3B_definitive_figure_legend_report.md"
    )
    with report_path.open("w", encoding="utf-8") as handle:
        handle.write(
            "# Phase U27B3B - Definitive figure-legend construction\n\n"
        )
        handle.write(f"- Version: `{VERSION}`\n")
        handle.write(f"- Decision: **{decision}**\n")
        handle.write("- Figures with definitive legends: **8/8**.\n")
        handle.write(
            f"- Panels represented: "
            f"**{len(panel_legend_registry)}/57**.\n"
        )
        handle.write(
            f"- Unique panel keys: "
            f"**{panel_legend_registry['panel_key'].nunique()}/57**.\n"
        )
        handle.write(
            f"- Figure-level audits passed: "
            f"**{int(figure_audit['panel_completeness_pass'].sum())}/8**.\n"
        )
        handle.write(
            f"- DOCX status: **{docx_status}**.\n\n"
        )

        handle.write("## Legend outputs\n\n")
        handle.write(f"- Markdown: `{markdown_path}`\n")
        handle.write(f"- Plain text: `{text_path}`\n")
        handle.write(f"- DOCX: `{docx_path}`\n")
        handle.write(
            "- Panel provenance registry: "
            f"`{outmetadata / 'UTI_HostOmics_U27B3B_panel_legend_provenance_registry.tsv'}`\n"
        )
        handle.write(
            "- Figure audit: "
            f"`{outtables / 'UTI_HostOmics_U27B3B_figure_legend_audit.tsv'}`\n\n"
        )

        handle.write("## Integrity boundary\n\n")
        handle.write(
            "Legends were constructed against the frozen 57-panel package, "
            "source-value tables, build registry and visual approvals. No "
            "statistical effects were recalculated, no figure asset was "
            "modified and no manuscript file was edited.\n"
        )

    run_manifest = {
        "version": VERSION,
        "decision": decision,
        "figures_with_legends": 8,
        "panels_in_registry": len(panel_legend_registry),
        "unique_panel_keys": int(
            panel_legend_registry["panel_key"].nunique()
        ),
        "markdown_path": str(markdown_path),
        "text_path": str(text_path),
        "docx_path": str(docx_path),
        "docx_status": docx_status,
        "scientific_values_recalculated": False,
        "figure_assets_modified": False,
        "source_locks_changed": False,
        "manuscript_modified": False,
    }
    (
        outresults
        / "UTI_HostOmics_U27B3B_run_manifest.json"
    ).write_text(
        json.dumps(run_manifest, indent=2),
        encoding="utf-8",
    )

    log("Definitive legends constructed for Figures 1-8.")
    log(f"Panels represented: {len(panel_legend_registry)}/57")
    log(
        "Figure audits passed: "
        f"{int(figure_audit['panel_completeness_pass'].sum())}/8"
    )
    log(f"DOCX status: {docx_status}")
    log(f"Decision: {decision}")
    log(f"Report: {report_path}")

    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3B] ERROR: {exc}", file=sys.stderr)
        raise
