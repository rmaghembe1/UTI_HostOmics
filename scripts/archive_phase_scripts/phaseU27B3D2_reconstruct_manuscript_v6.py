#!/usr/bin/env python3
"""
Phase U27B3D2
Reconstruct the manuscript-wide scientific narrative as a new v6.0 derivative.

Frozen components preserved exactly:
- Results section from U27B3C2
- Figures 1-8 and definitive legends from U27B3C4.1
- Existing tables and reference-key table unless explicitly replaced

Reconstructed components:
- Front matter/version label/title/keywords
- Abstract
- Introduction
- Methods
- Discussion, Limitations, Future directions, Concluding model
- Data availability
- Code availability
- Funding
- Supplementary-table architecture

The v5.2 manuscript is never overwritten.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image

VERSION = "U27B3D2_v1.0_2026-07-16"
TAG = "phaseU27B3D2_manuscript_wide_v6_reconstruction"

DEFAULT_SOURCE = (
    "__UTI_HOSTOMICS_PROJECT_ROOT__/"
    "09_manuscript_docx/phaseU27B3C41_integrated_figure_section_cleanup/"
    "UTI_HostOmics_preZotero_manuscript_v5_2_U27B3C41_figure_section_cleaned.docx"
)

OUTPUT_NAME = (
    "UTI_HostOmics_preZotero_manuscript_"
    "v6_0_U27B3D2_scientifically_harmonized.docx"
)

TITLE = (
    "Endocrine-metabolic-immune remodeling across urinary tract infection, "
    "pregnancy-associated outcomes and cell-resolved host responses"
)

KEYWORDS = (
    "urinary tract infection; UPEC; pregnancy; recurrent UTI; TLR4; leptin; "
    "PI3K-AKT; steroid metabolism; cholesterol; complement; single-cell "
    "transcriptomics; immunometabolism; host omics"
)

ABSTRACT_PARAGRAPHS = [
    (
        "Background: Urinary tract infection arises within heterogeneous host "
        "contexts in which innate sensing, endocrine state, cellular composition "
        "and metabolic adaptation can jointly shape inflammation and outcome. "
        "These processes are usually examined in isolated datasets, limiting "
        "mechanistic comparison across recurrence-related, pregnancy-associated, "
        "human inflammatory and cell-resolved settings."
    ),
    (
        "Methods: We reconstructed a multi-context host-response atlas from four "
        "public datasets: GSE112098, GSE280297, GSE168600 and GSE252321. Seventy-"
        "eight curated submodules spanning ten biological axes were scored within "
        "each native species and tissue context. Cross-dataset synthesis used "
        "standardized effects, recurrence and directional concordance rather than "
        "pooled raw expression. The GSE252321 single-cell experiment was rebuilt "
        "into broad and refined immune populations and used for descriptive "
        "cellular localization and pseudobulk module scoring."
    ),
    (
        "Results: TLR4-LPS, leptin and PI3K-AKT formed the strongest recurrent "
        "infection-associated core. Pregnancy-associated UTI showed tissue- and "
        "outcome-specific remodeling rather than a uniform response, including "
        "branch-selective steroid synthesis-response decoupling and heterogeneous "
        "complement and inflammatory-carbon effects. Single-cell reconstruction "
        "of 27,385 quality-controlled cells resolved 18 clusters, six broad immune "
        "populations and 14 refined subtypes. UPEC was associated with expansion "
        "of neutrophil and macrophage/monocyte compartments, higher strict Treg-"
        "like and TNFSF9-positive macrophage fractions, and cell-context-specific "
        "localization of complement, endocrine and metabolic modules. C3a/C5a "
        "signaling and opsonophagocytosis were recurrent but remained provisional "
        "where independent sample support was limited."
    ),
    (
        "Conclusions: UTI host biology is organized around a recurrent innate-"
        "metabolic core that is reshaped by pregnancy, tissue compartment and "
        "immune-cell state. The resulting framework separates robust recurrent "
        "signals from provisional and context-divergent biology and provides a "
        "mechanistic basis for prospective endocrine, metabolic, complement and "
        "cell-resolved validation."
    ),
]

INTRODUCTION_PARAGRAPHS = [
    (
        "Urinary tract infection is a common infectious syndrome whose clinical "
        "expression ranges from self-limited lower-tract disease to recurrent "
        "episodes, pyelonephritis and pregnancy-associated complications. Although "
        "uropathogenic Escherichia coli and other uropathogens initiate infection, "
        "host outcome depends on more than microbial burden alone. Epithelial "
        "integrity, innate sensing, myeloid recruitment, endocrine physiology and "
        "metabolic adaptation all contribute to the inflammatory state in which "
        "infection is contained, persists or becomes clinically consequential "
        "[Hannan_2023_immune_defenses_urinary_tract; "
        "Schwartz_2023_uropathogen_host_pyelonephritis]."
    ),
    (
        "Pregnancy adds a particularly complex physiological context. Steroid and "
        "cholesterol synthesis, hormone-receptor responsiveness, maternal-fetal "
        "immune tolerance, placental metabolism and complement regulation are "
        "dynamically remodeled across gestation. UTI-related inflammation may "
        "therefore intersect with pregnancy biology through branch-specific rather "
        "than globally concordant pathways. Distinguishing steroid synthesis from "
        "downstream receptor response, and complement initiation from inflammatory "
        "or opsonophagocytic effector branches, is important for understanding how "
        "infection-associated signals might differ between term and preterm "
        "outcomes [GSE280297_maternofetal_UTI_preterm; "
        "Solano_Arck_2020_steroids_pregnancy_fetal_development]."
    ),
    (
        "Metabolic transcriptional programs are similarly heterogeneous. Leptin, "
        "insulin-receptor/IRS and PI3K-AKT signaling can connect immune activation "
        "to carbon use, glycogen handling, lipid synthesis and redox control, but "
        "the direction and magnitude of these programs can vary across tissues and "
        "cell types. Transcriptomic pathway scores cannot measure metabolite flux, "
        "yet they can reveal coordinated pathway activity that motivates direct "
        "metabolomic or perturbational testing."
    ),
    (
        "Single-cell analysis provides a route to localize recurrent bulk signals "
        "to specific immune compartments. However, cell-level abundance does not "
        "replace biological replication, and small sample numbers require careful "
        "separation of descriptive cellular localization from inferential claims. "
        "A cross-context analysis must therefore preserve species- and tissue-"
        "native modeling while integrating evidence through standardized effects, "
        "directional recurrence and explicit interpretation boundaries."
    ),
    (
        "Here, we reconstructed a multi-context UTI host-response atlas using "
        "GSE112098, GSE280297, GSE168600 and GSE252321. We expanded the analysis to "
        "78 submodules across ten endocrine, metabolic, complement and immune axes; "
        "resolved recurrent, provisional and context-divergent effects; rebuilt the "
        "single-cell experiment into broad and refined immune states; and integrated "
        "the results into an eight-figure mechanistic framework."
    ),
]

METHODS_SECTIONS: List[Tuple[str, List[str]]] = [
    ("Study design and evidence architecture", [
        "This study used an integrative reanalysis of publicly available bulk and single-cell transcriptomic datasets. Each dataset was analyzed independently in its native species, tissue and experimental design. Raw expression values were not pooled across studies. Cross-dataset synthesis was restricted to standardized module effects, recurrence, directional concordance and evidence-class assignment.",
        "The final framework comprised GSE112098 as a human urinary/systemic inflammatory comparator, GSE280297 as a mouse pregnancy-associated UTI model across bladder, placenta and uterus, GSE168600 as a mouse recurrent or prior-exposure bladder model, and GSE252321 as a four-sample single-cell UPEC experiment."
    ]),
    ("Expanded submodule library", [
        "A curated library of 78 submodules was organized into ten biological axes: steroid-cholesterol-endocrine biology, lipid metabolism, adipokine signaling, insulin/IRS signaling, inflammatory carbon metabolism, amino-acid metabolism, nucleotide/NAD/nitrogen metabolism, complement architecture, immune-context anchors and catecholamine-stress-adjacent signaling. Modules were retained only when sufficient genes were represented in the relevant matrix.",
        "Metabolic modules were interpreted as transcriptionally inferred pathway activity and not as direct measurements of metabolite concentration, reaction rate or flux."
    ]),
    ("Bulk matrix preparation and species-native scoring", [
        "Dataset-specific matrices were inspected for feature identity, orientation, sample annotation and expression scale. Mouse and human matrices were scored in their native gene-symbol space. Dataset-specific contrasts were estimated within each study and tissue, and no gene-level meta-analysis was performed across species.",
        "For GSE280297, tissue samples were retained as the inferential units because dam identifiers were unavailable. Pregnancy contrasts included preterm versus term, UPEC versus mock/PBS during pregnancy, and pregnant versus nonpregnant bladder comparisons where design support was available."
    ]),
    ("Effect-size synthesis and evidence classes", [
        "Module-level contrasts were summarized using signed effects, Hedges g where applicable, false-discovery-rate information, tissue coherence and recurrence across eligible contexts. An evidence hierarchy distinguished robust recurrent effects, provisional effects supported by one false-discovery-rate-qualified dataset plus independent directional concordance, context-divergent effects and limited independent support.",
        "Weighted concordance summarized consistency of effect direction across eligible independent contrasts; it was not treated as a pooled expression estimate or a causal effect."
    ]),
    ("Single-cell quality control and marker-based reconstruction", [
        "Four flat gene-by-cell matrices from GSE252321 were processed independently. Of 28,313 cells, 27,385 passed quality control. Marker-based clustering resolved 18 clusters, which were consolidated into six broad immune populations and 14 refined subtypes. Refined labels included LY6C2-VCAN inflammatory monocytes, RETNLA-MRC1 reparative macrophages, CD83-CLEC10A activated dendritic cells and MKI67-TOP2A cycling immune cells.",
        "Strict Treg-like cells required FOXP3 plus at least one of IL2RA, CTLA4 or TNFRSF18. The expanded Treg-like definition required FOXP3 or at least two supporting markers. TNFSF9-positive macrophages were defined by detectable expression, and TNFSF9-high cells used the prespecified pooled threshold of 2.208856."
    ]),
    ("Cell-type pseudobulk scoring and cellular attribution", [
        "Broad and refined cell populations were aggregated into sample-aware pseudobulks. Seventy-three modules were score-eligible in the single-cell layer. Cellular attribution integrated module-mean gene log2 fold change, directional coherence and bounded support scores rather than relying on unstable large standardized effects from the n=2 versus n=2 design.",
        "Because the single-cell experiment contained two control and two UPEC biological samples, composition and localization analyses were descriptive and hypothesis-generating. The minimum attainable exact two-sided permutation p-value was 0.333."
    ]),
    ("Figure assembly and reproducibility", [
        "All analyses, source locks, figure builds, visual audits and manuscript transformations were implemented reproducibly under the project directory __UTI_HOSTOMICS_PROJECT_ROOT__. Executable scripts are stored under 10_scripts/. Figures 1-8 were frozen after source-value, checksum, panel-count and visual audits."
    ]),
]

DISCUSSION_SECTIONS: List[Tuple[str, List[str]]] = [
    ("Principal findings", [
        "The integrated atlas identifies a recurrent UTI host-response core centered on TLR4-LPS sensing, leptin signaling and PI3K-AKT activity. Around this core, insulin/IRS, glycogen and inflammatory-carbon programs varied by model, indicating that recurrent direction does not imply uniform magnitude or identical mechanism across tissues and species.",
        "Pregnancy-associated UTI produced a distinct branch-selective architecture. Steroidogenesis and androgen/testosterone synthesis could diverge from estrogen, cholesterol-biosynthetic, receptor-response and metabolic-effector programs, while complement and inflammatory-carbon branches varied across tissues and outcomes. This pattern is more consistent with synthesis-response decoupling than with a single globally activated endocrine state."
    ]),
    ("Recurrent innate-metabolic coupling", [
        "The recurrence of TLR4, leptin and PI3K-AKT across independently analyzed contexts suggests a coordinated interface between microbial sensing and metabolic adaptation. Leptin and insulin-response pathways can influence immune-cell activation, survival and substrate handling, while PI3K-AKT integrates receptor-proximal signals with downstream metabolic and inflammatory programs. The present data support coordinated directional association but do not establish a direct causal sequence."
    ]),
    ("Pregnancy steroid and lipid remodeling", [
        "The pregnancy results emphasize that steroid biology should be decomposed into synthesis, interconversion, transport and receptor-response layers. A positive shift in steroidogenic or androgen-associated transcription does not necessarily imply increased downstream hormone action, and reduced estrogen or cholesterol-biosynthetic activity in one tissue does not define the entire maternal-fetal system. These findings motivate direct measurement of steroid intermediates, aromatization, receptor activity and tissue-specific hormone availability in future studies."
    ]),
    ("Complement branching and cellular context", [
        "Complement initiation, amplification, C3a/C5a signaling, opsonophagocytosis, terminal membrane-attack-complex biology and regulatory/coagulation crosstalk did not move as a single unit. C3a/C5a signaling localized broadly, with strong dendritic and cDC1-like support, whereas opsonophagocytosis was more myeloid weighted and strongest in activated dendritic cells. These recurrent observations remain provisional where independent sample support is limited."
    ]),
    ("Cell-resolved immune remodeling", [
        "UPEC-associated expansion of neutrophil and macrophage/monocyte fractions, together with higher strict Treg-like and TNFSF9-positive macrophage fractions, suggests that infection alters both effector composition and regulatory-inflammatory states. The localization of IRS signaling to T-cell and regulatory-like compartments, androgen-receptor activity to dendritic states and TLR4-LPS signaling to activated lymphoid states further argues against assigning bulk pathway effects to a single dominant cell type."
    ]),
    ("Immunometabolic interpretation", [
        "Carbon, amino-acid, nucleotide, NAD and redox modules displayed context-specific remodeling. Glycogen synthesis showed coherent cellular support, whereas fatty-acid synthesis, amino-acid transport and several nitrogen or redox pathways were branch divergent. These patterns identify candidate metabolic programs but remain transcriptional inferences rather than biochemical flux measurements."
    ]),
]

LIMITATIONS_PARAGRAPHS = [
    "The study is limited by heterogeneity in species, tissue, platform and experimental design. Cross-dataset integration therefore used standardized module effects and directional concordance rather than merged expression or direct gene-level meta-analysis.",
    "Pregnancy analyses did not show broad pregnancy-wide false-discovery-rate support. Tissue samples were the inferential units because dam identifiers were unavailable, preventing maternal-level clustering or causal attribution to preterm outcome.",
    "The GSE252321 single-cell layer contained two control and two UPEC biological samples. Cellular composition and localization are descriptive, and the minimum exact two-sided permutation p-value was 0.333. Large standardized effects in this setting can be variance artifacts and were not used as stand-alone evidence.",
    "Transcriptomic metabolic modules indicate inferred pathway activity and cannot establish metabolite abundance, enzymatic rate or metabolic flux. Complement recurrence also remains provisional where independent sample support is limited."
]

FUTURE_PARAGRAPHS = [
    "Prospective studies should combine urine, blood, exfoliated urothelial cells and tissue-relevant sampling with pathogen metadata, recurrence follow-up and pregnancy outcome information. Matched steroid, lipid, cytokine, complement, metabolomic and proteomic measurements would directly test the branch-selective mechanisms identified here.",
    "Larger single-cell and spatial datasets are needed to validate TNFSF9-positive macrophage states, regulatory-inflammatory T-cell programs, dendritic complement localization and tissue-specific endocrine responses. Perturbational models targeting TLR4, leptin/IRS/PI3K-AKT, complement branches, aromatization and lipid-redox programs would help distinguish correlation from mechanism."
]

CONCLUDING_PARAGRAPHS = [
    "The final model separates a robust recurrent infection core from provisional and context-divergent biology. TLR4-LPS, leptin and PI3K-AKT provide the strongest recurrent framework; pregnancy-associated steroid biology is best described as branch-selective synthesis-response decoupling; complement inflammatory and opsonophagocytic branches are recurrent but provisional; and single-cell analysis localizes these programs to distinct immune compartments without overcoming limited biological replication.",
    "This endocrine-metabolic-immune atlas provides a reproducible foundation for mechanistic validation of UTI susceptibility, recurrence and pregnancy-associated outcomes."
]

DATA_AVAILABILITY = (
    "All datasets analyzed in this study are publicly available through the Gene "
    "Expression Omnibus under accession numbers GSE112098, GSE280297, GSE168600 "
    "and GSE252321. Dataset-specific roles, inclusion decisions, processed "
    "matrices, source locks and analysis manifests are documented within the "
    "project outputs. No new participant-level data were generated."
)

CODE_AVAILABILITY = (
    "All analysis and manuscript-generation steps were implemented as reproducible "
    "scripts under `10_scripts/` in the project directory. Frozen figure assets, "
    "source-value tables, audit manifests and checksum records are retained in the "
    "corresponding phase-specific output directories. A public repository archive "
    "will be finalized before journal submission."
)

FUNDING = (
    "This work was supported by the Global Infectious Diseases Research Training "
    "Program, a collaboration between Kibong'oto Infectious Diseases Hospital and "
    "the University of Virginia, under the Fogarty International Center, National "
    "Institutes of Health grant D43 TW012247."
)

SUPPLEMENTARY_ITEMS = [
    "Table S1. Dataset architecture, sample design and inclusion roles for GSE112098, GSE280297, GSE168600 and GSE252321.",
    "Table S2. Expanded 78-submodule library organized across ten biological axes.",
    "Table S3. Dataset-specific module effects and factorial or adjusted contrasts.",
    "Table S4. Cross-dataset recurrence, directional concordance and evidence-class assignments.",
    "Table S5. GSE280297 pregnancy, tissue and outcome-specific module effects.",
    "Table S6. GSE252321 quality control, cluster markers, broad populations and refined subtypes.",
    "Table S7. Broad-cell and refined-subtype pseudobulk module localization results.",
    "Table S8. Complement-stage and endocrine-metabolic cellular attribution tables.",
    "Table S9. Figure 1-8 source-value manifest and panel-level provenance registry.",
    "Table S10. Interpretation-boundary, sensitivity and manuscript claim-traceability register."
]


def log(message: str) -> None:
    print(f"[U27B3D2] {message}", flush=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def heading_indices(document: Document) -> Dict[str, List[int]]:
    result: Dict[str, List[int]] = {}
    for idx, paragraph in enumerate(document.paragraphs):
        text = normalize(paragraph.text).lower().rstrip(".: ")
        if text:
            result.setdefault(text, []).append(idx)
    return result


def find_heading(document: Document, name: str) -> int:
    target = name.lower()
    matches = []
    for idx, paragraph in enumerate(document.paragraphs):
        cleaned = normalize(paragraph.text).lower().rstrip(".: ")
        cleaned = re.sub(r"^\d+(?:\.\d+)*\s*", "", cleaned)
        if cleaned == target:
            matches.append(idx)
    if len(matches) != 1:
        raise RuntimeError(f"Expected one heading {name!r}; observed {matches}")
    return matches[0]


def clear_between(document: Document, start_idx: int, end_idx: int) -> None:
    paragraphs = document.paragraphs
    start_p = paragraphs[start_idx]._p
    end_p = paragraphs[end_idx]._p
    element = start_p.getnext()
    while element is not None and element is not end_p:
        next_element = element.getnext()
        element.getparent().remove(element)
        element = next_element


def insert_paragraph_before(anchor, text: str = "", style: Optional[str] = None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    return paragraph


def replace_section(document: Document, heading: str, next_heading: str, paragraphs: Sequence[str]) -> None:
    start = find_heading(document, heading)
    end = find_heading(document, next_heading)
    clear_between(document, start, end)
    anchor = document.paragraphs[end]
    for text in paragraphs:
        p = insert_paragraph_before(anchor, text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08


def replace_section_with_subheadings(document: Document, heading: str, next_heading: str, sections: Sequence[Tuple[str, Sequence[str]]]) -> None:
    start = find_heading(document, heading)
    end = find_heading(document, next_heading)
    clear_between(document, start, end)
    anchor = document.paragraphs[end]
    for subheading, paragraphs in sections:
        hp = insert_paragraph_before(anchor, subheading, "Heading 2")
        hp.paragraph_format.space_before = Pt(8)
        hp.paragraph_format.space_after = Pt(3)
        for text in paragraphs:
            p = insert_paragraph_before(anchor, text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.08


def replace_single_paragraph_section(document: Document, heading: str, next_heading: str, text: str) -> None:
    replace_section(document, heading, next_heading, [text])


def update_front_matter(document: Document) -> None:
    # Replace visible version/generation labels and working title wherever present.
    replacements = {
        "UTI HostOmics Project - Draft manuscript v4.1": "UTI HostOmics Project - Scientifically harmonized manuscript v6.0",
        "Pre-Zotero manuscript v4 figure-aligned draft": "Pre-Zotero manuscript v6.0 | Frozen Results and Figures 1-8",
        "Generated: 2026-07-09": "Generated: 2026-07-16",
        "Cellular and immunometabolic architecture of urinary tract infection susceptibility, recurrence, and pregnancy-associated inflammation: an integrative host-omics and single-cell validation study": TITLE,
        "urinary tract infection; recurrent UTI; pregnancy-associated UTI; UPEC; host omics; single-cell transcriptomics; innate immunity; NLRP3; TLR4; TLR5; glucocorticoid signaling; oxytocin signaling; immunometabolism": KEYWORDS,
        "Draft generated 2026-07-09 | Pre-Zotero draft": "Draft generated 2026-07-16 | Pre-Zotero scientifically harmonized v6.0",
    }
    for paragraph in document.paragraphs:
        text = normalize(paragraph.text)
        if text in replacements:
            paragraph.text = replacements[text]


def replace_discussion_block(document: Document) -> None:
    start = find_heading(document, "Discussion")
    end = find_heading(document, "Data availability")
    clear_between(document, start, end)
    anchor = document.paragraphs[end]

    for subheading, paras in DISCUSSION_SECTIONS:
        hp = insert_paragraph_before(anchor, subheading, "Heading 2")
        for text in paras:
            p = insert_paragraph_before(anchor, text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.08

    hp = insert_paragraph_before(anchor, "Limitations", "Heading 2")
    for text in LIMITATIONS_PARAGRAPHS:
        p = insert_paragraph_before(anchor, text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

    hp = insert_paragraph_before(anchor, "Future directions", "Heading 2")
    for text in FUTURE_PARAGRAPHS:
        p = insert_paragraph_before(anchor, text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

    hp = insert_paragraph_before(anchor, "Concluding model", "Heading 2")
    for text in CONCLUDING_PARAGRAPHS:
        p = insert_paragraph_before(anchor, text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)


def replace_supplementary_section(document: Document) -> None:
    start = find_heading(document, "Supplementary tables")
    end = find_heading(document, "Remaining reference gaps after citation-key cleanup")
    clear_between(document, start, end)
    anchor = document.paragraphs[end]
    for item in SUPPLEMENTARY_ITEMS:
        p = insert_paragraph_before(anchor, item)
        p.paragraph_format.space_after = Pt(3)


def section_text(document: Document, start_heading: str, end_heading: str) -> List[str]:
    start = find_heading(document, start_heading)
    end = find_heading(document, end_heading)
    return [normalize(p.text) for p in document.paragraphs[start + 1:end]]


def render_docx(docx_path: Path, outdir: Path) -> Dict[str, object]:
    outdir.mkdir(parents=True, exist_ok=True)
    renderer = Path("/home/oai/skills/docx/render_docx.py")
    if not renderer.exists():
        return {"render_pass": False, "reason": "render_docx.py not found", "page_count": 0, "contact_sheet": ""}
    cmd = [sys.executable, str(renderer), str(docx_path), "--output_dir", str(outdir), "--emit_pdf"]
    result = subprocess.run(cmd, text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    pages = sorted(outdir.glob("page-*.png"))
    if result.returncode != 0 or not pages:
        return {"render_pass": False, "reason": result.stderr.strip() or result.stdout.strip(), "page_count": len(pages), "contact_sheet": ""}

    # Contact sheet
    images = []
    width = 520
    for page in pages:
        img = Image.open(page).convert("RGB")
        ratio = width / img.width
        images.append(img.resize((width, max(1, int(img.height * ratio)))))
    cols = 3
    pad = 20
    rows = (len(images) + cols - 1) // cols
    row_heights = [max(img.height for img in images[r*cols:(r+1)*cols]) for r in range(rows)]
    canvas = Image.new("RGB", (cols*width + (cols+1)*pad, sum(row_heights)+(rows+1)*pad), "white")
    y = pad
    for r in range(rows):
        x = pad
        for img in images[r*cols:(r+1)*cols]:
            canvas.paste(img, (x, y))
            x += width + pad
        y += row_heights[r] + pad
    contact = outdir / "UTI_HostOmics_U27B3D2_render_contact_sheet.png"
    canvas.save(contact)
    return {"render_pass": True, "reason": "rendered", "page_count": len(pages), "contact_sheet": str(contact)}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-root", default="__UTI_HOSTOMICS_PROJECT_ROOT__")
    parser.add_argument("--source", default=DEFAULT_SOURCE)
    args = parser.parse_args()

    project = Path(args.project_root).resolve()
    source = Path(args.source).resolve()
    if not source.exists():
        raise FileNotFoundError(source)

    outdoc = project / "09_manuscript_docx" / TAG
    outtables = project / "06_tables" / TAG
    outmeta = project / "03_metadata" / TAG
    outresults = project / "05_results" / TAG
    renderdir = outdoc / "render_qa"
    for d in (outdoc, outtables, outmeta, outresults, renderdir):
        d.mkdir(parents=True, exist_ok=True)

    source_hash_before = sha256(source)
    document = Document(source)
    frozen_results_before = section_text(document, "Results", "Discussion")
    frozen_figure_text_before = [normalize(p.text) for p in document.paragraphs if re.match(r"^Figure\s+[1-8]\.\s", normalize(p.text))]
    embedded_images_before = len(document.inline_shapes)

    update_front_matter(document)
    replace_section(document, "Abstract", "Introduction", ABSTRACT_PARAGRAPHS)
    replace_section(document, "Introduction", "Methods", INTRODUCTION_PARAGRAPHS)
    replace_section_with_subheadings(document, "Methods", "Results", METHODS_SECTIONS)
    replace_discussion_block(document)
    replace_single_paragraph_section(document, "Data availability", "Code availability", DATA_AVAILABILITY)
    replace_single_paragraph_section(document, "Code availability", "Ethics statement", CODE_AVAILABILITY)
    replace_single_paragraph_section(document, "Funding", "Acknowledgements", FUNDING)
    replace_supplementary_section(document)

    output = outdoc / OUTPUT_NAME
    if output.exists():
        output.unlink()
    document.save(output)

    source_hash_after = sha256(source)
    output_hash = sha256(output)
    out_doc = Document(output)

    frozen_results_after = section_text(out_doc, "Results", "Discussion")
    frozen_figure_text_after = [normalize(p.text) for p in out_doc.paragraphs if re.match(r"^Figure\s+[1-8]\.\s", normalize(p.text))]
    embedded_images_after = len(out_doc.inline_shapes)

    results_preserved = frozen_results_before == frozen_results_after
    figure_legends_preserved = frozen_figure_text_before == frozen_figure_text_after and len(frozen_figure_text_after) == 8
    images_preserved = embedded_images_before == embedded_images_after == 8
    source_unchanged = source_hash_before == source_hash_after

    full_text = "\n".join(normalize(p.text) for p in out_doc.paragraphs)
    obsolete_terms = ["GSE186800", "GSE261018", "04_scripts/", "17 modules", "17 validated modules"]
    obsolete_audit = pd.DataFrame([{"term": t, "occurrence_count": len(re.findall(re.escape(t), full_text, flags=re.I)), "absent": re.search(re.escape(t), full_text, flags=re.I) is None} for t in obsolete_terms])
    obsolete_audit.to_csv(outtables / "UTI_HostOmics_U27B3D2_obsolete_term_audit.tsv", sep="\t", index=False)

    required = {
        "datasets": ["GSE112098", "GSE280297", "GSE168600", "GSE252321"],
        "architecture": ["78", "ten biological axes"],
        "infection_core": ["TLR4", "leptin", "PI3K-AKT"],
        "pregnancy": ["branch-selective", "steroid", "preterm"],
        "single_cell": ["27,385", "18 clusters", "14 refined subtypes", "0.333"],
        "complement": ["C3a/C5a", "opsonophagocytosis", "provisional"],
        "code_path": ["10_scripts/"],
    }
    content_rows = []
    for audit_id, terms in required.items():
        present = all(term.lower() in full_text.lower() for term in terms)
        content_rows.append({"audit_id": audit_id, "required_terms": "; ".join(terms), "pass": present})
    content_audit = pd.DataFrame(content_rows)
    content_audit.to_csv(outtables / "UTI_HostOmics_U27B3D2_content_audit.tsv", sep="\t", index=False)

    render = render_docx(output, renderdir)
    pd.DataFrame([render]).to_csv(outtables / "UTI_HostOmics_U27B3D2_render_audit.tsv", sep="\t", index=False)

    structural_pass = all([results_preserved, figure_legends_preserved, images_preserved, source_unchanged])
    content_pass = bool(content_audit["pass"].all() and obsolete_audit["absent"].all())
    if structural_pass and content_pass and render["render_pass"]:
        decision = "READY_FOR_U27B3D3_V6_MANUSCRIPT_VISUAL_AND_SCIENTIFIC_AUDIT"
    elif structural_pass and content_pass:
        decision = "V6_RECONSTRUCTION_COMPLETE_RENDER_QA_PENDING"
    else:
        decision = "TARGETED_U27B3D2_V6_RECONSTRUCTION_REPAIR_REQUIRED"

    preservation = pd.DataFrame([{
        "source_path": str(source),
        "source_sha256_before": source_hash_before,
        "source_sha256_after": source_hash_after,
        "source_unchanged": source_unchanged,
        "output_path": str(output),
        "output_sha256": output_hash,
        "results_preserved": results_preserved,
        "figure_legends_preserved": figure_legends_preserved,
        "embedded_images_before": embedded_images_before,
        "embedded_images_after": embedded_images_after,
        "embedded_images_preserved": images_preserved,
    }])
    preservation.to_csv(outtables / "UTI_HostOmics_U27B3D2_preservation_audit.tsv", sep="\t", index=False)

    pd.DataFrame([{
        "phase": "U27B3D2",
        "decision": decision,
        "source_unchanged": source_unchanged,
        "results_preserved": results_preserved,
        "figures_and_legends_preserved": figure_legends_preserved,
        "embedded_images_preserved": images_preserved,
        "all_required_content_present": content_pass,
        "obsolete_terms_absent": bool(obsolete_audit["absent"].all()),
        "render_pass": render["render_pass"],
        "page_count": render["page_count"],
        "manuscript_source_modified": False,
        "new_derivative_created": True,
        "next_phase": "U27B3D3 inspect v6.0 render, verify scientific coherence and approve for supplementary-table reconstruction" if decision.startswith("READY_FOR_U27B3D3") else "Inspect failed audits",
    }]).to_csv(outtables / "UTI_HostOmics_U27B3D2_phase_decision.tsv", sep="\t", index=False)

    pd.DataFrame([
        {"field": "source_path", "value": str(source)},
        {"field": "source_sha256", "value": source_hash_before},
        {"field": "output_path", "value": str(output)},
        {"field": "output_sha256", "value": output_hash},
        {"field": "version", "value": VERSION},
    ]).to_csv(outmeta / "UTI_HostOmics_U27B3D2_derivative_record.tsv", sep="\t", index=False)

    report = outresults / "UTI_HostOmics_U27B3D2_v6_reconstruction_report.md"
    report.write_text(
        f"# Phase U27B3D2 - Manuscript-wide v6.0 reconstruction\n\n"
        f"- Version: `{VERSION}`\n"
        f"- Decision: **{decision}**\n"
        f"- Source: `{source}`\n"
        f"- Output: `{output}`\n"
        f"- Results preserved: **{results_preserved}**.\n"
        f"- Figures and legends preserved: **{figure_legends_preserved}**.\n"
        f"- Embedded images preserved: **{images_preserved}**.\n"
        f"- Required content audit: **{content_pass}**.\n"
        f"- Obsolete terms absent: **{bool(obsolete_audit['absent'].all())}**.\n"
        f"- Render pass: **{render['render_pass']}**.\n"
        f"- Render pages: **{render['page_count']}**.\n"
        f"- Contact sheet: `{render.get('contact_sheet', '')}`.\n\n"
        "## Reconstruction boundary\n\n"
        "The U27B3C2 Results section and U27B3A/U27B3B frozen figure package were preserved. Abstract, Introduction, Methods, Discussion, limitations, conclusions, availability statements, funding and supplementary-table architecture were reconstructed in a new derivative.\n",
        encoding="utf-8",
    )

    manifest = {
        "version": VERSION,
        "decision": decision,
        "source": str(source),
        "source_sha256": source_hash_before,
        "output": str(output),
        "output_sha256": output_hash,
        "results_preserved": results_preserved,
        "figures_legends_preserved": figure_legends_preserved,
        "images_preserved": images_preserved,
        "content_pass": content_pass,
        "render_pass": render["render_pass"],
        "page_count": render["page_count"],
    }
    (outresults / "UTI_HostOmics_U27B3D2_run_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    log(f"Results preserved: {results_preserved}")
    log(f"Figures/legends preserved: {figure_legends_preserved}")
    log(f"Embedded images preserved: {images_preserved}")
    log(f"Required content pass: {content_pass}")
    log(f"Render pass: {render['render_pass']}")
    log(f"Decision: {decision}")
    log(f"Output: {output}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3D2] ERROR: {exc}", file=sys.stderr)
        raise
