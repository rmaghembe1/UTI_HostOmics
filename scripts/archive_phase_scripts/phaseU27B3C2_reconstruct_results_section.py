#!/usr/bin/env python3
"""
Phase U27B3C2
Create a new v5.0 manuscript derivative with the Results section rebuilt in
frozen Figure 1-8 order.

The authoritative v4.1 U23 review-handoff manuscript is read-only. This phase
replaces only the block-level content between the existing Results and
Discussion headings. The original file is never overwritten.

The new Results section is aligned to the frozen eight-figure package and
preserves the established inferential boundaries:
- no broad pregnancy-wide FDR claim;
- tissue samples, not inferred dams, are the pregnancy units;
- single-cell localization is descriptive at n=2 control versus n=2 UPEC;
- the minimum attainable exact two-sided permutation p-value is 0.333;
- metabolic modules indicate transcriptionally inferred pathway activity,
  not metabolic flux;
- raw expression is not pooled across species, tissues or studies;
- complement recurrence remains provisional where support is limited.

Outputs include the derivative DOCX, Results text, paragraph manifest,
preservation audit, render pages/contact sheet when LibreOffice and pdftoppm
are available, and a phase decision.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

try:
    from PIL import Image
except ImportError:
    Image = None


VERSION = "U27B3C2_v1.0_2026-07-16"
TAG = "phaseU27B3C2_results_section_reconstruction"
EXPECTED_SOURCE_SHA256 = (
    "148b4f7b6cb0c60e620913229e0f490fb"
    "621778b5a3ae8d6b8ddab2145d15e90"
)
DEFAULT_SOURCE = (
    "__UTI_HOSTOMICS_PROJECT_ROOT__/"
    "09_manuscript_docx/phaseU23_review_handoff_package/"
    "01_review_main_files/"
    "UTI_HostOmics_preZotero_manuscript_v4_1_draft_with_figures.docx"
)
OUTPUT_FILENAME = (
    "UTI_HostOmics_preZotero_manuscript_"
    "v5_0_U27B3C2_results_reconstructed.docx"
)

RESULTS_SECTIONS: List[Tuple[str, List[str]]] = [
    (
        "An expanded multi-context atlas establishes the analytical and evidence framework",
        [
            (
                "We assembled a multi-context host-response atlas spanning a human urinary/systemic inflammatory comparator (GSE112098; n=73), a mouse pregnancy-associated UTI model sampled across bladder, placenta and uterus (GSE280297; n=60 tissue samples), a mouse recurrent or prior-exposure bladder model (GSE168600; n=20), and a four-sample single-cell UPEC experiment (GSE252321). The analytical framework comprised 78 curated submodules distributed across ten biological axes, including steroid-cholesterol-endocrine biology, lipid metabolism, adipokine and insulin/IRS signaling, inflammatory carbon metabolism, amino-acid metabolism, nucleotide/NAD/nitrogen metabolism, complement architecture, immune-context anchors and stress-adjacent signaling. Each dataset was analyzed independently in its native species and tissue context, and cross-study integration was performed using signed standardized effects, recurrence and directional concordance rather than pooled raw expression (Figure 1A-E)."
            ),
            (
                "The evidence framework separated modules with recurrent cross-dataset concordance from provisional effects supported by one false-discovery-rate-qualified dataset plus independent directional support, context-divergent effects, and modules with limited independent evidence. This hierarchy was used throughout the atlas to distinguish reproducible host-response architecture from biologically plausible but context-sensitive signals. The cellular scoring layer retained 73 score-eligible modules and was interpreted as localization evidence rather than as an additional independent cohort (Figure 1F)."
            ),
        ],
    ),
    (
        "Independent datasets identify a recurrent TLR4-leptin-PI3K-AKT infection core",
        [
            (
                "Independent analysis of the eligible infection contrasts identified a recurrent signaling core centered on TLR4-LPS sensing, leptin signaling and PI3K-AKT activity. These modules showed the most consistent positive infection-associated direction across the independently analyzed datasets. Insulin receptor/IRS signaling, glycogen synthesis and inflammatory-carbon programs contributed additional support, but their magnitudes and recurrence varied across models, consistent with context-dependent metabolic engagement rather than a uniform response across all tissues and species (Figure 2A-C)."
            ),
            (
                "The adjusted human comparator and the mouse factorial contrasts further separated recurrent from context-specific biology. TLR4-LPS, leptin and PI3K-AKT retained the clearest cross-dataset support, whereas several endocrine, lipid, carbon-use and immune-effector modules were prominent only in particular models. Evidence-weighted concordance therefore captured shared direction without assuming equivalent effect magnitude or identical mechanism across datasets. Complement C3a/C5a signaling and opsonophagocytosis also recurred across contexts, but these were retained as provisional core features because independent sample support was more limited (Figure 2D-G)."
            ),
        ],
    ),
    (
        "Pregnancy-associated UTI shows tissue-specific and branch-selective remodeling",
        [
            (
                "The GSE280297 design resolved mock/PBS, nonpregnant UPEC, preterm UPEC and term UPEC states across bladder, placenta and uterus. This architecture enabled separate evaluation of preterm-versus-term outcome effects, UPEC-versus-PBS effects during pregnancy and pregnant-versus-nonpregnant bladder effects. No broad pregnancy-wide false-discovery-rate signal was detected; accordingly, the pregnancy analyses were interpreted through effect size, tissue coherence and branch-level biological organization rather than through a claim of global significance. Tissue samples were retained as the inferential units because dam identifiers were unavailable (Figure 3A-C)."
            ),
            (
                "Within the preterm-versus-term comparison, steroid biology was characterized by synthesis-response decoupling rather than uniform activation or suppression. Core steroidogenesis and androgen/testosterone biosynthesis tended to shift positively, whereas estrogen biosynthesis/interconversion and cholesterol-biosynthetic programs tended to shift negatively. Broad receptor-response and metabolic-effector domains were comparatively attenuated. The resulting pattern supports a branch-selective pregnancy-associated reorganization in which steroid synthesis or transformation can diverge from downstream cellular response programs (Figure 3D-F)."
            ),
            (
                "Complement and inflammatory-carbon biology was likewise branch-selective. C3a/C5a signaling, opsonophagocytosis, complement regulation, coagulation crosstalk and carbon-use programs differed across tissues and contrasts rather than moving as a single coordinated pathway. These findings define a tissue- and outcome-dependent pregnancy model in which UPEC-associated innate sensing and complement activation intersect with endocrine and metabolic remodeling, while remaining exploratory in the absence of broad pregnancy-wide FDR support (Figure 3G,H)."
            ),
        ],
    ),
    (
        "Single-cell reconstruction localizes recurrent modules to specific immune compartments",
        [
            (
                "The four GSE252321 samples contained 28,313 cells, of which 27,385 passed quality control. Marker-based reconstruction resolved 18 clusters that were consolidated into six broad immune populations and 14 refined subtypes, including LY6C2-VCAN inflammatory monocytes, RETNLA-MRC1 reparative macrophages, CD83-CLEC10A activated dendritic cells and MKI67-TOP2A cycling immune cells. Because the biological design comprised two control and two UPEC samples, all composition and localization results were treated as descriptive biological-replicate summaries rather than cell-level independent tests (Figure 4A-D)."
            ),
            (
                "UPEC exposure was associated with a larger neutrophil fraction (difference, +0.210) and macrophage/monocyte fraction (+0.097), accompanied by lower T-cell (-0.132), dendritic-cell (-0.103), cycling-immune (-0.036), NK-cell (-0.025) and mast-cell (-0.010) fractions. Within T cells, the strict Treg-like fraction increased from 0.017 in controls to 0.083 in UPEC samples. Within macrophage/monocytes, the TNFSF9-positive fraction increased from 0.171 to 0.304, and the prespecified TNFSF9-high fraction increased from 0.037 to 0.080 (Figure 4C-F)."
            ),
            (
                "Cell-type pseudobulk scoring localized the recurrent modules to distinct immune compartments. C3a/C5a signaling showed coherent pan-immune support with strongest localization to dendritic and cDC1-like states; glycogen synthesis was broadly coherent and strongest in neutrophils or inflammatory monocytes; IRS signaling was strongest in T-cell and regulatory-like T-cell compartments; androgen-receptor signaling was strongest in dendritic and cDC2-like states; and TLR4-LPS signaling showed a lymphoid-weighted pattern with strongest refined support in activated T cells. Opsonophagocytosis was predominantly myeloid and strongest in activated dendritic cells. These localizations are hypothesis-generating: with n=2 samples per group, the minimum attainable exact two-sided permutation p-value was 0.333 (Figure 4G,H)."
            ),
        ],
    ),
    (
        "Steroid, cholesterol and lipid programs separate synthesis from response",
        [
            (
                "Focused endocrine analysis confirmed that the pregnancy-associated steroid pattern was organized across distinct branches rather than along a single activation axis. Steroidogenesis, androgen/testosterone synthesis, estrogen synthesis and interconversion, cholesterol biosynthesis and handling, and androgen, estrogen and glucocorticoid response programs displayed different directions across infection, pregnancy and tissue contexts. The infection-versus-pregnancy comparison consequently separated synthesis-associated effects from receptor-response effects and reinforced the synthesis-response decoupling model (Figure 5A-D)."
            ),
            (
                "Cellular localization added a compartmental dimension to this endocrine architecture. Androgen-receptor signaling showed broad immune support with strongest refined localization in cDC2-like cells, whereas cholesterol and lipid-regulatory programs varied among dendritic, macrophage/monocyte, neutrophil, lymphoid and cycling compartments. Lipid-droplet dynamics, PPAR/SREBP/LXR regulation, ferroptosis-linked lipid peroxidation, fatty-acid synthesis and beta-oxidation, phospholipid metabolism, sphingolipid/ceramide metabolism and eicosanoid/prostaglandin metabolism were context dependent. Fatty-acid synthesis was especially branch divergent, with strongest broad support in cycling immune cells and strongest refined support in regulatory/type-2-like T cells (Figure 5E-G)."
            ),
        ],
    ),
    (
        "Adipokine, insulin and metabolic programs define a broader immunometabolic response",
        [
            (
                "The focused immunometabolic analysis positioned leptin, insulin receptor/IRS and PI3K-AKT signaling at the interface between innate sensing and metabolic adaptation. Leptin and PI3K-AKT were recurrently positive in infection contexts but attenuated in the preterm-versus-term comparison, whereas IRS signaling showed broad cellular coherence with strongest localization in T-cell and regulatory-like T-cell compartments. The integrated topology therefore supports coordinated, context-dependent coupling among adipokine signaling, insulin-response machinery, TLR4-LPS sensing and downstream carbon-use programs, without establishing a direct causal sequence (Figure 6A-E)."
            ),
            (
                "Transcriptionally inferred carbon-use programs included glycolysis, lactate/HIF1A signaling, glycogen synthesis, pentose-phosphate activity, gluconeogenic and tricarboxylic-acid-cycle modules. Glycogen synthesis showed coherent pan-immune support, with strongest broad localization in neutrophils and strongest refined localization in inflammatory monocytes. The remaining carbon programs varied across pregnancy, tissue and infection contrasts, indicating flexible pathway deployment rather than a single universal metabolic state (Figure 6F)."
            ),
            (
                "Amino-acid, nitrogen, nucleotide and redox modules extended this metabolic architecture. Amino-acid transport was context dependent, with strongest broad localization in dendritic cells and strongest refined localization in cytotoxic NK cells. Arginine/nitric-oxide and urea-cycle biology, glutamine and tryptophan metabolism, methionine/SAM and serine/glycine/one-carbon pathways, purine and pyrimidine metabolism, NAD biology, xanthine-oxidase-linked oxidative purine catabolism and NRF2-redox programs all showed contrast-specific remodeling. These modules represent transcriptionally inferred pathway activity and do not measure metabolite concentrations, reaction rates or metabolic flux (Figure 6G,H)."
            ),
        ],
    ),
    (
        "Complement remodeling is branch-selective and cell-context dependent",
        [
            (
                "Complement-specific reconstruction resolved classical, lectin and alternative initiation, C3 convertase and amplification, C3a/C5a signaling, opsonophagocytosis, terminal membrane-attack-complex biology, complement regulation and coagulation crosstalk. The branches differed across independent infection and pregnancy contrasts, arguing against a simple global complement activation model. C3a/C5a signaling and opsonophagocytosis provided the most recurrent cross-context support, although both remained provisional where independent sample support was limited (Figure 7A,B,E,F)."
            ),
            (
                "Cellular attribution further separated inflammatory and effector branches. C3a/C5a signaling was broadly distributed and strongest in dendritic or cDC1-like states, whereas opsonophagocytosis was myeloid weighted and strongest in activated dendritic cells. Seven of nine complement modules had eligible cell-level support; lectin-pathway and terminal membrane-attack-complex modules were not represented in the cell-level matrix and should therefore be considered unavailable rather than biologically absent. The resulting complement architecture links branch-specific tissue effects to distinct cellular compartments (Figure 7C,D,G)."
            ),
        ],
    ),
    (
        "Integrated evidence separates recurrent, provisional and contextual UTI biology",
        [
            (
                "Across the complete atlas, UPEC/TLR4 signaling converged with leptin-IRS-PI3K/AKT activity, complement and opsonophagocytosis, steroid and lipid branching, and carbon, nitrogen and redox remodeling. These programs were distributed across expanded neutrophil and macrophage/monocyte compartments, activated dendritic states, lymphoid signaling niches, TNFSF9-positive macrophages and regulatory-inflammatory Treg-like states. Pregnancy and tissue context altered the relative balance among these branches rather than creating a single uniformly activated host state (Figure 8A-E)."
            ),
            (
                "The final evidence model therefore distinguishes a robust recurrent core from provisional and context-divergent biology. TLR4-LPS, leptin and PI3K-AKT constitute the strongest recurrent infection-associated framework; complement C3a/C5a signaling and opsonophagocytosis remain provisional recurrent features; pregnancy-associated steroid biology is best described as branch-selective synthesis-response decoupling; and the single-cell layer localizes these programs without overcoming the limited biological replication. The integrated network is an evidence-weighted synthesis rather than a direct causal model, and all metabolic interpretations refer to transcriptionally inferred pathway activity rather than flux (Figure 8F)."
            ),
        ],
    ),
]


def log(message: str) -> None:
    print(f"[U27B3C2] {message}", flush=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            block = handle.read(1024 * 1024)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def normalized_heading(text: str) -> str:
    value = re.sub(r"\s+", " ", text).strip().lower()
    value = re.sub(r"^\d+(?:\.\d+)*\s*", "", value)
    return value.rstrip(".:")


def find_heading_index(paragraphs: Sequence[Paragraph], heading: str) -> int:
    target = heading.lower()
    matches = [
        index
        for index, paragraph in enumerate(paragraphs)
        if normalized_heading(paragraph.text) == target
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected exactly one {heading} heading; observed {len(matches)}."
        )
    return matches[0]


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    style_name: Optional[str],
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_paragraph.style = style_name
        except KeyError:
            pass
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def choose_styles(old_results: Sequence[Paragraph]) -> Tuple[str, str]:
    subsection_style = "Heading 2"
    body_style = "Normal"

    nonempty = [paragraph for paragraph in old_results if paragraph.text.strip()]
    if nonempty:
        first = nonempty[0]
        if first.style is not None and first.style.name:
            subsection_style = first.style.name

    if len(nonempty) >= 2:
        second = nonempty[1]
        if second.style is not None and second.style.name:
            body_style = second.style.name

    return subsection_style, body_style


def remove_body_elements_between(
    document: Document,
    start_paragraph: Paragraph,
    end_paragraph: Paragraph,
) -> int:
    body = document._element.body
    children = list(body)
    start_element = start_paragraph._p
    end_element = end_paragraph._p

    try:
        start_index = children.index(start_element)
        end_index = children.index(end_element)
    except ValueError as exc:
        raise RuntimeError("Results/Discussion XML boundary could not be resolved.") from exc

    if start_index >= end_index:
        raise RuntimeError("Results heading does not precede Discussion heading.")

    removable = children[start_index + 1:end_index]
    for element in removable:
        body.remove(element)
    return len(removable)


def results_plain_text() -> str:
    blocks: List[str] = []
    for heading, paragraphs in RESULTS_SECTIONS:
        blocks.append(heading)
        blocks.extend(paragraphs)
    return "\n\n".join(blocks)


def segment_text(document: Document, start: int, end: int) -> str:
    return "\n".join(
        paragraph.text
        for paragraph in document.paragraphs[start:end]
    )


def document_structure(document: Document) -> Dict[str, object]:
    paragraphs = document.paragraphs
    results_index = find_heading_index(paragraphs, "results")
    discussion_index = find_heading_index(paragraphs, "discussion")
    return {
        "results_index": results_index,
        "discussion_index": discussion_index,
        "paragraph_count": len(paragraphs),
        "pre_results_text": segment_text(document, 0, results_index + 1),
        "discussion_onward_text": segment_text(document, discussion_index, len(paragraphs)),
        "results_text": segment_text(document, results_index + 1, discussion_index),
    }


def results_content_audit(text: str) -> Dict[str, object]:
    lower = text.lower()
    figure_positions = []
    for number in range(1, 9):
        match = re.search(rf"figure\s+{number}\b", lower)
        figure_positions.append(match.start() if match else -1)

    figure_order_pass = all(
        position >= 0 for position in figure_positions
    ) and figure_positions == sorted(figure_positions)

    required_phrases = {
        "no_broad_pregnancy_fdr": "no broad pregnancy-wide false-discovery-rate",
        "no_dam_inference": "dam identifiers were unavailable",
        "single_cell_n2": "two control and two upec samples",
        "minimum_p": "0.333",
        "not_flux": "not measure metabolite concentrations, reaction rates or metabolic flux",
        "no_raw_pooling": "rather than pooled raw expression",
        "complement_provisional": "remained provisional",
        "noncausal_model": "rather than a direct causal model",
    }

    phrase_results = {
        key: phrase in lower
        for key, phrase in required_phrases.items()
    }

    return {
        "word_count": len(re.findall(r"\b[\w'-]+\b", text)),
        "subsection_count": len(RESULTS_SECTIONS),
        "figure_order_pass": figure_order_pass,
        **phrase_results,
        "old_gse186800_absent": "gse186800" not in lower,
        "current_gse168600_present": "gse168600" in lower,
        "old_traceability_note_absent": "phase u10" not in lower and "phase u11" not in lower,
    }


def count_docx_members(path: Path, prefix: str) -> int:
    import zipfile

    with zipfile.ZipFile(path) as archive:
        return sum(
            1
            for name in archive.namelist()
            if name.startswith(prefix) and not name.endswith("/")
        )


def find_executable(names: Sequence[str]) -> Optional[str]:
    for name in names:
        path = shutil.which(name)
        if path:
            return path
    return None


def render_docx(
    docx_path: Path,
    render_dir: Path,
) -> Dict[str, object]:
    render_dir.mkdir(parents=True, exist_ok=True)
    office = find_executable(["libreoffice", "soffice"])
    pdftoppm = find_executable(["pdftoppm"])

    if office is None:
        return {
            "render_status": "SKIPPED_LIBREOFFICE_NOT_FOUND",
            "pdf_path": "",
            "page_png_count": 0,
            "contact_sheet": "",
        }

    profile_dir = render_dir / "libreoffice_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    profile_uri = profile_dir.resolve().as_uri()

    command = [
        office,
        f"-env:UserInstallation={profile_uri}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(render_dir),
        str(docx_path),
    ]
    completed = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=180,
    )

    pdf_path = render_dir / f"{docx_path.stem}.pdf"
    if completed.returncode != 0 or not pdf_path.exists():
        return {
            "render_status": "FAILED_PDF_CONVERSION",
            "pdf_path": str(pdf_path),
            "page_png_count": 0,
            "contact_sheet": "",
            "stdout": completed.stdout,
            "stderr": completed.stderr,
        }

    if pdftoppm is None:
        return {
            "render_status": "PDF_CREATED_PDFTOPPM_NOT_FOUND",
            "pdf_path": str(pdf_path),
            "page_png_count": 0,
            "contact_sheet": "",
        }

    page_prefix = render_dir / "page"
    raster = subprocess.run(
        [
            pdftoppm,
            "-png",
            "-r",
            "144",
            str(pdf_path),
            str(page_prefix),
        ],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=240,
    )

    page_paths = sorted(render_dir.glob("page-*.png"))
    if raster.returncode != 0 or not page_paths:
        return {
            "render_status": "FAILED_PAGE_RASTERIZATION",
            "pdf_path": str(pdf_path),
            "page_png_count": len(page_paths),
            "contact_sheet": "",
            "stdout": raster.stdout,
            "stderr": raster.stderr,
        }

    contact_sheet = render_dir / "UTI_HostOmics_U27B3C2_render_contact_sheet.png"
    if Image is not None:
        images = [Image.open(path).convert("RGB") for path in page_paths]
        thumb_width = 420
        resized = []
        for image in images:
            ratio = thumb_width / image.width
            resized.append(
                image.resize((thumb_width, max(1, int(image.height * ratio))))
            )
        columns = 3
        padding = 24
        rows = math.ceil(len(resized) / columns)
        row_heights = []
        for row in range(rows):
            subset = resized[row * columns:(row + 1) * columns]
            row_heights.append(max(image.height for image in subset))
        canvas = Image.new(
            "RGB",
            (
                columns * thumb_width + (columns + 1) * padding,
                sum(row_heights) + (rows + 1) * padding,
            ),
            "white",
        )
        y = padding
        for row in range(rows):
            x = padding
            subset = resized[row * columns:(row + 1) * columns]
            for image in subset:
                canvas.paste(image, (x, y))
                x += thumb_width + padding
            y += row_heights[row] + padding
        canvas.save(contact_sheet)

    return {
        "render_status": "RENDERED_PENDING_MANUAL_VISUAL_AUDIT",
        "pdf_path": str(pdf_path),
        "page_png_count": len(page_paths),
        "contact_sheet": str(contact_sheet) if contact_sheet.exists() else "",
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root",
        default="__UTI_HOSTOMICS_PROJECT_ROOT__",
    )
    parser.add_argument(
        "--source",
        default=DEFAULT_SOURCE,
    )
    args = parser.parse_args()

    project = Path(args.project_root).resolve()
    source = Path(args.source).resolve()

    if not project.exists():
        raise FileNotFoundError(f"Project root not found: {project}")
    if not source.exists():
        raise FileNotFoundError(f"Source manuscript not found: {source}")

    source_hash_before = sha256(source)
    if source_hash_before != EXPECTED_SOURCE_SHA256:
        raise RuntimeError(
            "Authoritative manuscript SHA256 does not match the confirmed v4.1 source."
        )

    decision_input = (
        project
        / "06_tables"
        / "phaseU27B3C11_authoritative_target_and_structure_audit"
        / "UTI_HostOmics_U27B3C11_phase_decision.tsv"
    )
    if not decision_input.exists():
        raise FileNotFoundError(
            f"U27B3C1.1 decision not found: {decision_input}"
        )
    decision_frame = pd.read_csv(decision_input, sep="\t", low_memory=False)
    if len(decision_frame) != 1 or not str(
        decision_frame.iloc[0]["decision"]
    ).startswith("READY_FOR_U27B3C2"):
        raise RuntimeError("U27B3C1.1 did not release the manuscript to U27B3C2.")

    outdir = (
        project
        / "09_manuscript_docx"
        / TAG
    )
    outtables = project / "06_tables" / TAG
    outmetadata = project / "03_metadata" / TAG
    outresults = project / "05_results" / TAG
    render_dir = outdir / "render_qa"

    for directory in (
        outdir,
        outtables,
        outmetadata,
        outresults,
        render_dir,
    ):
        directory.mkdir(parents=True, exist_ok=True)

    output_docx = outdir / OUTPUT_FILENAME

    source_document = Document(source)
    source_structure = document_structure(source_document)
    source_paragraphs = source_document.paragraphs
    results_index = int(source_structure["results_index"])
    discussion_index = int(source_structure["discussion_index"])
    old_results = source_paragraphs[results_index + 1:discussion_index]
    subsection_style, body_style = choose_styles(old_results)

    document = Document(source)
    paragraphs = document.paragraphs
    results_index = find_heading_index(paragraphs, "results")
    discussion_index = find_heading_index(paragraphs, "discussion")
    results_heading = paragraphs[results_index]
    discussion_heading = paragraphs[discussion_index]

    removed_blocks = remove_body_elements_between(
        document,
        results_heading,
        discussion_heading,
    )

    manifest_rows: List[Dict[str, object]] = []
    cursor = results_heading
    paragraph_order = 0

    for section_number, (heading, body_paragraphs) in enumerate(
        RESULTS_SECTIONS,
        start=1,
    ):
        cursor = insert_paragraph_after(
            cursor,
            heading,
            subsection_style,
        )
        cursor.paragraph_format.keep_with_next = True
        manifest_rows.append(
            {
                "order": paragraph_order,
                "section_number": section_number,
                "paragraph_type": "subheading",
                "style": cursor.style.name if cursor.style else "",
                "text": heading,
                "word_count": len(re.findall(r"\b[\w'-]+\b", heading)),
            }
        )
        paragraph_order += 1

        for body_text in body_paragraphs:
            cursor = insert_paragraph_after(
                cursor,
                body_text,
                body_style,
            )
            cursor.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cursor.paragraph_format.widow_control = True
            manifest_rows.append(
                {
                    "order": paragraph_order,
                    "section_number": section_number,
                    "paragraph_type": "body",
                    "style": cursor.style.name if cursor.style else "",
                    "text": body_text,
                    "word_count": len(
                        re.findall(r"\b[\w'-]+\b", body_text)
                    ),
                }
            )
            paragraph_order += 1

    document.save(output_docx)

    source_hash_after = sha256(source)
    if source_hash_after != source_hash_before:
        raise RuntimeError("The read-only v4.1 source manuscript changed unexpectedly.")

    derivative_hash = sha256(output_docx)
    derivative_document = Document(output_docx)
    derivative_structure = document_structure(derivative_document)
    derivative_results_text = str(derivative_structure["results_text"])

    results_text_path = (
        outdir
        / "UTI_HostOmics_U27B3C2_reconstructed_results_section.txt"
    )
    results_text_path.write_text(
        derivative_results_text,
        encoding="utf-8",
    )

    manifest = pd.DataFrame(manifest_rows)
    manifest.to_csv(
        outtables
        / "UTI_HostOmics_U27B3C2_results_paragraph_manifest.tsv",
        sep="\t",
        index=False,
    )

    content_audit = results_content_audit(derivative_results_text)

    source_media_count = count_docx_members(source, "word/media/")
    derivative_media_count = count_docx_members(output_docx, "word/media/")

    pre_results_preserved = (
        str(source_structure["pre_results_text"])
        == str(derivative_structure["pre_results_text"])
    )
    discussion_preserved = (
        str(source_structure["discussion_onward_text"])
        == str(derivative_structure["discussion_onward_text"])
    )

    structural_audit = pd.DataFrame(
        [
            {
                "source_path": str(source),
                "source_sha256_before": source_hash_before,
                "source_sha256_after": source_hash_after,
                "source_unchanged": source_hash_before == source_hash_after,
                "derivative_path": str(output_docx),
                "derivative_sha256": derivative_hash,
                "removed_block_elements": removed_blocks,
                "new_results_subsections": len(RESULTS_SECTIONS),
                "new_results_paragraphs": len(manifest),
                "results_heading_count": 1,
                "discussion_heading_count": 1,
                "pre_results_text_preserved": pre_results_preserved,
                "discussion_onward_text_preserved": discussion_preserved,
                "source_media_files": source_media_count,
                "derivative_media_files": derivative_media_count,
                "source_overwritten": False,
            }
        ]
    )
    structural_audit.to_csv(
        outtables
        / "UTI_HostOmics_U27B3C2_docx_preservation_audit.tsv",
        sep="\t",
        index=False,
    )

    content_audit_frame = pd.DataFrame([content_audit])
    content_audit_frame.to_csv(
        outtables
        / "UTI_HostOmics_U27B3C2_results_content_audit.tsv",
        sep="\t",
        index=False,
    )

    render_record = render_docx(output_docx, render_dir)
    pd.DataFrame([render_record]).to_csv(
        outtables
        / "UTI_HostOmics_U27B3C2_render_audit.tsv",
        sep="\t",
        index=False,
    )

    content_pass = bool(
        content_audit["subsection_count"] == 8
        and content_audit["figure_order_pass"]
        and content_audit["no_broad_pregnancy_fdr"]
        and content_audit["no_dam_inference"]
        and content_audit["single_cell_n2"]
        and content_audit["minimum_p"]
        and content_audit["not_flux"]
        and content_audit["no_raw_pooling"]
        and content_audit["complement_provisional"]
        and content_audit["noncausal_model"]
        and content_audit["old_gse186800_absent"]
        and content_audit["current_gse168600_present"]
        and content_audit["old_traceability_note_absent"]
    )
    structure_pass = bool(
        source_hash_before == source_hash_after
        and pre_results_preserved
        and discussion_preserved
        and output_docx.exists()
        and output_docx.stat().st_size > 0
    )
    render_status = str(render_record.get("render_status", ""))
    render_pass = render_status == "RENDERED_PENDING_MANUAL_VISUAL_AUDIT"

    if content_pass and structure_pass and render_pass:
        decision = (
            "READY_FOR_U27B3C3_RESULTS_SECTION_VISUAL_AND_SCIENTIFIC_AUDIT"
        )
    elif content_pass and structure_pass:
        decision = (
            "RESULTS_DERIVATIVE_CREATED_PENDING_MANUAL_RENDER_QA"
        )
    else:
        decision = (
            "TARGETED_U27B3C2_RESULTS_OR_DOCX_PRESERVATION_REPAIR_REQUIRED"
        )

    decision_output = pd.DataFrame(
        [
            {
                "phase": "U27B3C2",
                "decision": decision,
                "source_read_only": True,
                "source_sha256_unchanged": source_hash_before == source_hash_after,
                "derivative_created": output_docx.exists(),
                "results_subsections": len(RESULTS_SECTIONS),
                "results_paragraphs": len(manifest),
                "results_word_count": content_audit["word_count"],
                "figure_1_to_8_order_pass": content_audit["figure_order_pass"],
                "content_boundary_audit_pass": content_pass,
                "pre_results_text_preserved": pre_results_preserved,
                "discussion_onward_text_preserved": discussion_preserved,
                "render_status": render_status,
                "rendered_pages": int(render_record.get("page_png_count", 0)),
                "manuscript_modified": True,
                "original_manuscript_modified": False,
                "figure_assets_modified": False,
                "next_phase": (
                    "U27B3C3 manually inspect rendered pages and scientifically audit the reconstructed Results section"
                    if decision.startswith("READY_FOR_U27B3C3")
                    else "Complete render QA or repair failed content/preservation checks"
                ),
            }
        ]
    )
    decision_output.to_csv(
        outtables
        / "UTI_HostOmics_U27B3C2_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    derivative_record = pd.DataFrame(
        [
            {
                "field": "authoritative_source_path",
                "value": str(source),
            },
            {
                "field": "authoritative_source_sha256",
                "value": source_hash_before,
            },
            {
                "field": "derivative_path",
                "value": str(output_docx),
            },
            {
                "field": "derivative_sha256",
                "value": derivative_hash,
            },
            {
                "field": "edit_scope",
                "value": "Results block between existing Results and Discussion headings",
            },
            {
                "field": "figure_integration_status",
                "value": "Frozen Figures 1-8 referenced but not inserted in this phase",
            },
            {
                "field": "next_phase",
                "value": "U27B3C3 visual and scientific audit",
            },
        ]
    )
    derivative_record.to_csv(
        outmetadata
        / "UTI_HostOmics_U27B3C2_derivative_record.tsv",
        sep="\t",
        index=False,
    )

    report_path = (
        outresults
        / "UTI_HostOmics_U27B3C2_results_reconstruction_report.md"
    )
    with report_path.open("w", encoding="utf-8") as handle:
        handle.write("# Phase U27B3C2 - Results-section reconstruction\n\n")
        handle.write(f"- Version: `{VERSION}`\n")
        handle.write(f"- Decision: **{decision}**\n")
        handle.write(f"- Read-only source: `{source}`\n")
        handle.write(f"- New derivative: `{output_docx}`\n")
        handle.write(f"- Results subsections: **{len(RESULTS_SECTIONS)}**.\n")
        handle.write(f"- Results paragraphs: **{len(manifest)}**.\n")
        handle.write(f"- Results word count: **{content_audit['word_count']}**.\n")
        handle.write(
            f"- Figure 1-8 order audit: **{content_audit['figure_order_pass']}**.\n"
        )
        handle.write(f"- Content-boundary audit: **{content_pass}**.\n")
        handle.write(f"- Pre-Results text preserved: **{pre_results_preserved}**.\n")
        handle.write(f"- Discussion onward preserved: **{discussion_preserved}**.\n")
        handle.write(f"- Render status: **{render_status}**.\n")
        handle.write(
            f"- Rendered pages: **{int(render_record.get('page_png_count', 0))}**.\n\n"
        )
        handle.write("## Scientific architecture\n\n")
        handle.write(
            "The Results section now follows the frozen figure sequence: atlas and evidence framework; recurrent infection core; pregnancy and tissue remodeling; single-cell localization; steroid/cholesterol/lipid biology; adipokine/insulin and metabolic remodeling; complement architecture; and integrated evidence synthesis.\n\n"
        )
        handle.write("## Integrity boundary\n\n")
        handle.write(
            "The v4.1 source was not overwritten. Only the Results block in the new derivative was replaced. Text before Results and from Discussion onward was preserved, and frozen figure assets were not modified or inserted during this phase.\n"
        )

    run_manifest = {
        "version": VERSION,
        "decision": decision,
        "source_path": str(source),
        "source_sha256": source_hash_before,
        "derivative_path": str(output_docx),
        "derivative_sha256": derivative_hash,
        "results_subsections": len(RESULTS_SECTIONS),
        "results_paragraphs": len(manifest),
        "results_word_count": content_audit["word_count"],
        "content_audit_pass": content_pass,
        "structure_audit_pass": structure_pass,
        "render_status": render_status,
        "rendered_pages": int(render_record.get("page_png_count", 0)),
        "source_overwritten": False,
        "figure_assets_modified": False,
    }
    (
        outresults
        / "UTI_HostOmics_U27B3C2_run_manifest.json"
    ).write_text(
        json.dumps(run_manifest, indent=2),
        encoding="utf-8",
    )

    log(f"Source unchanged: {source_hash_before == source_hash_after}")
    log(f"Derivative: {output_docx}")
    log(f"Results subsections: {len(RESULTS_SECTIONS)}")
    log(f"Results paragraphs: {len(manifest)}")
    log(f"Results word count: {content_audit['word_count']}")
    log(f"Render status: {render_status}")
    log(f"Decision: {decision}")
    log(f"Report: {report_path}")

    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3C2] ERROR: {exc}", file=sys.stderr)
        raise
