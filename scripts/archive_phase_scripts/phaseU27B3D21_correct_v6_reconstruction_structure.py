#!/usr/bin/env python3
"""
Phase U27B3D2.1
Correct structural placement defects in the U27B3D2 v6 manuscript reconstruction.

The initial U27B3D2 script used paragraph indices captured before deleting section
contents. After deletion, those indices became stale, causing reconstructed
Abstract/Introduction/Methods/Discussion and availability text to be inserted at
incorrect downstream positions. Its renderer also depended on a container-only
render_docx.py path that is not present in the user's WSL environment.

This corrected phase:
1. loads the validated scientific text constants from the U27B3D2 script;
2. uses stable paragraph objects, not stale indices, as insertion anchors;
3. reconstructs Abstract, Introduction, Methods and Discussion in their correct
   manuscript sections;
4. repairs Data availability, Code availability and Funding placement;
5. removes any duplicate misplaced Discussion block between Figures 3 and 4;
6. removes legacy figure-selection and open-issue tables while preserving the
   Zotero reference table;
7. updates visible title/header/footer labels to v6.1;
8. preserves the frozen Results, Figures 1-8, legends and embedded image bytes;
9. renders through LibreOffice plus pdftoppm and creates a contact sheet.

The source DOCX is read-only and is never overwritten.
"""

from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image

VERSION = "U27B3D21_v1.0_2026-07-16"
TAG = "phaseU27B3D21_corrected_v6_reconstruction"

DEFAULT_SOURCE = (
    "__UTI_HOSTOMICS_PROJECT_ROOT__/"
    "09_manuscript_docx/phaseU27B3D2_manuscript_wide_v6_reconstruction/"
    "UTI_HostOmics_preZotero_manuscript_v6_0_U27B3D2_scientifically_harmonized.docx"
)

OUTPUT_NAME = (
    "UTI_HostOmics_preZotero_manuscript_"
    "v6_1_U27B3D21_scientifically_harmonized_corrected.docx"
)

DISCUSSION_SUBHEADINGS = [
    "Principal findings",
    "Recurrent innate-metabolic coupling",
    "Pregnancy steroid and lipid remodeling",
    "Complement branching and cellular context",
    "Cell-resolved immune remodeling",
    "Immunometabolic interpretation",
    "Limitations",
    "Future directions",
    "Concluding model",
]

FIGURE_CAPTION_RE = re.compile(r"^Figure\s+([1-8])\.\s+", re.I)


def log(message: str) -> None:
    print(f"[U27B3D2.1] {message}", flush=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            block = handle.read(1024 * 1024)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def load_u27b3d2_module(project: Path):
    candidates = [
        project / "10_scripts" / "phaseU27B3D2_reconstruct_manuscript_v6.py",
        Path("/mnt/data/phaseU27B3D2_reconstruct_manuscript_v6.py"),
    ]
    module_path = next((path for path in candidates if path.exists()), None)
    if module_path is None:
        raise FileNotFoundError(
            "The U27B3D2 reconstruction script was not found in 10_scripts/."
        )

    spec = importlib.util.spec_from_file_location("u27b3d2_source", module_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load U27B3D2 module: {module_path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module, module_path


def clean_heading_text(text: str) -> str:
    value = normalize(text).lower().rstrip(".: ")
    return re.sub(r"^\d+(?:\.\d+)*\s*", "", value)


def find_heading_paragraph(document: Document, heading: str):
    target = heading.lower()
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if clean_heading_text(paragraph.text) == target
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one heading {heading!r}; observed {len(matches)}."
        )
    return matches[0]


def clear_between_paragraphs(start_paragraph, end_paragraph) -> int:
    removed = 0
    element = start_paragraph._p.getnext()
    while element is not None and element is not end_paragraph._p:
        next_element = element.getnext()
        element.getparent().remove(element)
        removed += 1
        element = next_element
    return removed


def insert_body_paragraph(anchor, text: str, style: Optional[str] = None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    return paragraph


def format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08


def replace_section(
    document: Document,
    heading: str,
    next_heading: str,
    paragraphs: Sequence[str],
) -> int:
    start = find_heading_paragraph(document, heading)
    anchor = find_heading_paragraph(document, next_heading)
    removed = clear_between_paragraphs(start, anchor)
    for text in paragraphs:
        paragraph = insert_body_paragraph(anchor, text)
        format_body_paragraph(paragraph)
    return removed


def replace_section_with_subheadings(
    document: Document,
    heading: str,
    next_heading: str,
    sections: Sequence[Tuple[str, Sequence[str]]],
) -> int:
    start = find_heading_paragraph(document, heading)
    anchor = find_heading_paragraph(document, next_heading)
    removed = clear_between_paragraphs(start, anchor)
    for subheading, paragraphs in sections:
        heading_paragraph = insert_body_paragraph(anchor, subheading, "Heading 2")
        heading_paragraph.paragraph_format.space_before = Pt(8)
        heading_paragraph.paragraph_format.space_after = Pt(3)
        for text in paragraphs:
            paragraph = insert_body_paragraph(anchor, text)
            format_body_paragraph(paragraph)
    return removed


def replace_discussion_block(document: Document, source_module) -> int:
    start = find_heading_paragraph(document, "Discussion")
    anchor = find_heading_paragraph(document, "Data availability")
    removed = clear_between_paragraphs(start, anchor)

    for subheading, paragraphs in source_module.DISCUSSION_SECTIONS:
        heading_paragraph = insert_body_paragraph(anchor, subheading, "Heading 2")
        heading_paragraph.paragraph_format.space_before = Pt(8)
        heading_paragraph.paragraph_format.space_after = Pt(3)
        for text in paragraphs:
            paragraph = insert_body_paragraph(anchor, text)
            format_body_paragraph(paragraph)

    for subheading, paragraphs in (
        ("Limitations", source_module.LIMITATIONS_PARAGRAPHS),
        ("Future directions", source_module.FUTURE_PARAGRAPHS),
        ("Concluding model", source_module.CONCLUDING_PARAGRAPHS),
    ):
        heading_paragraph = insert_body_paragraph(anchor, subheading, "Heading 2")
        heading_paragraph.paragraph_format.space_before = Pt(8)
        heading_paragraph.paragraph_format.space_after = Pt(3)
        for text in paragraphs:
            paragraph = insert_body_paragraph(anchor, text)
            format_body_paragraph(paragraph)

    return removed


def replace_supplementary_section(document: Document, source_module) -> int:
    start = find_heading_paragraph(document, "Supplementary tables")
    anchor = find_heading_paragraph(
        document, "Remaining reference gaps after citation-key cleanup"
    )
    removed = clear_between_paragraphs(start, anchor)
    for item in source_module.SUPPLEMENTARY_ITEMS:
        paragraph = insert_body_paragraph(anchor, item)
        paragraph.paragraph_format.space_after = Pt(3)
    return removed


def body_element_text(element) -> str:
    texts = []
    for node in element.iter():
        if node.tag.rsplit("}", 1)[-1] == "t":
            texts.append(node.text or "")
    return normalize("".join(texts))


def paragraph_has_drawing(element) -> bool:
    return bool(
        element.findall(f".//{qn('w:drawing')}")
        or element.findall(f".//{qn('w:pict')}")
    )


def remove_misplaced_discussion_between_figures(document: Document) -> int:
    """Remove the known U27B3D2 stale-anchor duplicate between Figures 3 and 4."""
    body = document._element.body
    children = list(body)

    figure3_caption_index = None
    figure4_drawing_index = None

    for index, element in enumerate(children):
        if element.tag != qn("w:p"):
            continue
        text = body_element_text(element)
        if text.startswith("Figure 3."):
            figure3_caption_index = index
        if (
            figure3_caption_index is not None
            and index > figure3_caption_index
            and paragraph_has_drawing(element)
        ):
            figure4_drawing_index = index
            break

    if figure3_caption_index is None or figure4_drawing_index is None:
        return 0

    interval = children[figure3_caption_index + 1:figure4_drawing_index]
    contains_duplicate = any(
        body_element_text(element) == "Principal findings"
        for element in interval
        if element.tag == qn("w:p")
    )
    if not contains_duplicate:
        return 0

    removed = 0
    for element in interval:
        body.remove(element)
        removed += 1
    return removed


def remove_legacy_tracking_tables(document: Document) -> Dict[str, int]:
    body = document._element.body
    removed_figure_selection_tables = 0
    removed_open_issue_elements = 0

    # Remove the legacy selected-asset table.
    for element in list(body):
        if element.tag == qn("w:tbl"):
            text = body_element_text(element).lower()
            if (
                "primary_asset" in text
                and "selection_status" in text
                and "recommended_action" in text
            ):
                body.remove(element)
                removed_figure_selection_tables += 1

    # Remove the Open issues heading, its table and blank separators up to the
    # reference-table heading.
    children = list(body)
    start_index = None
    end_index = None
    for index, element in enumerate(children):
        text = body_element_text(element)
        if text == "Open issues retained for final submission":
            start_index = index
        if text == "Reference table for Zotero finalization":
            end_index = index
            break

    if start_index is not None and end_index is not None and start_index < end_index:
        for element in children[start_index:end_index]:
            body.remove(element)
            removed_open_issue_elements += 1

    return {
        "legacy_figure_selection_tables_removed": removed_figure_selection_tables,
        "legacy_open_issue_elements_removed": removed_open_issue_elements,
    }



def cleanup_shifted_section_duplicates(document: Document, source_module) -> Dict[str, int]:
    """Remove residual paragraphs shifted by the original stale-anchor build."""
    removed_code_from_ethics = 0
    removed_funding_from_acknowledgements = 0
    removed_duplicate_supplementary_items = 0
    removed_redundant_page_breaks = 0

    # Remove the duplicated Code availability paragraph from the Ethics section.
    ethics = find_heading_paragraph(document, "Ethics statement")
    author_contributions = find_heading_paragraph(document, "Author contributions")
    element = ethics._p.getnext()
    while element is not None and element is not author_contributions._p:
        next_element = element.getnext()
        text_value = body_element_text(element)
        if text_value.startswith("All analysis and manuscript-generation steps were implemented"):
            element.getparent().remove(element)
            removed_code_from_ethics += 1
        element = next_element

    # Remove the duplicated Funding paragraph from Acknowledgements.
    acknowledgements = find_heading_paragraph(document, "Acknowledgements")
    supplementary = find_heading_paragraph(document, "Supplementary tables")
    element = acknowledgements._p.getnext()
    while element is not None and element is not supplementary._p:
        next_element = element.getnext()
        if normalize(body_element_text(element)) == normalize(source_module.FUNDING):
            element.getparent().remove(element)
            removed_funding_from_acknowledgements += 1
        element = next_element

    # Remove the old duplicate Table S1-S10 descriptions that occur after the
    # reference-gap heading and before the frozen figure block.
    gap_heading = find_heading_paragraph(
        document, "Remaining reference gaps after citation-key cleanup"
    )
    first_figure_caption = next(
        paragraph
        for paragraph in document.paragraphs
        if normalize(paragraph.text).startswith("Figure 1.")
    )
    element = gap_heading._p.getnext()
    while element is not None and element is not first_figure_caption._p:
        next_element = element.getnext()
        text_value = body_element_text(element)
        if re.match(r"^Table S(?:[1-9]|10)\.\s", text_value, flags=re.I):
            element.getparent().remove(element)
            removed_duplicate_supplementary_items += 1
        element = next_element

    # Collapse multiple inherited page-break paragraphs before the Zotero
    # reference table into one deterministic page break on the heading itself.
    reference_heading = find_heading_paragraph(
        document, "Reference table for Zotero finalization"
    )
    previous = reference_heading._p.getprevious()
    while previous is not None and previous.tag == qn("w:p"):
        text_value = body_element_text(previous)
        has_page_break = any(
            br.get(qn("w:type")) == "page"
            for br in previous.findall(".//" + qn("w:br"))
        )
        ppr = previous.find(qn("w:pPr"))
        page_before = (
            ppr is not None
            and ppr.find(qn("w:pageBreakBefore")) is not None
        )
        if text_value == "" and (has_page_break or page_before):
            to_remove = previous
            previous = previous.getprevious()
            to_remove.getparent().remove(to_remove)
            removed_redundant_page_breaks += 1
        else:
            break
    reference_heading.paragraph_format.page_break_before = True

    return {
        "duplicate_code_paragraphs_removed_from_ethics": removed_code_from_ethics,
        "duplicate_funding_paragraphs_removed_from_acknowledgements": removed_funding_from_acknowledgements,
        "duplicate_supplementary_items_removed": removed_duplicate_supplementary_items,
        "redundant_page_break_paragraphs_removed": removed_redundant_page_breaks,
    }

def replace_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def update_front_matter_headers_footers(document: Document, source_module) -> None:
    exact_replacements = {
        "Pre-Zotero manuscript v5.2 | Results reconstructed; frozen Figures 1-8 integrated": (
            "Pre-Zotero manuscript v6.1 | Scientifically harmonized; frozen Results and Figures 1-8"
        ),
        "Pre-Zotero manuscript v4 figure-aligned draft": (
            "Pre-Zotero manuscript v6.1 | Scientifically harmonized; frozen Results and Figures 1-8"
        ),
        "Updated: 2026-07-16": "Updated: 2026-07-16",
        "Generated: 2026-07-09": "Updated: 2026-07-16",
        (
            "Cellular and immunometabolic architecture of urinary tract infection "
            "susceptibility, recurrence, and pregnancy-associated inflammation: "
            "an integrative host-omics and single-cell validation study"
        ): source_module.TITLE,
        (
            "urinary tract infection; recurrent UTI; pregnancy-associated UTI; UPEC; "
            "host omics; single-cell transcriptomics; innate immunity; NLRP3; TLR4; "
            "TLR5; glucocorticoid signaling; oxytocin signaling; immunometabolism"
        ): source_module.KEYWORDS,
    }

    for paragraph in document.paragraphs:
        text = normalize(paragraph.text)
        if text in exact_replacements:
            replace_paragraph_text(paragraph, exact_replacements[text])

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if "UTI HostOmics Project" in paragraph.text:
                replace_paragraph_text(
                    paragraph,
                    "UTI HostOmics Project - Scientifically harmonized manuscript v6.1",
                )
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                replace_paragraph_text(
                    paragraph,
                    "Updated 2026-07-16 | Pre-Zotero scientifically harmonized draft",
                )

    document.core_properties.title = source_module.TITLE
    document.core_properties.subject = (
        "Integrated endocrine-metabolic-immune UTI host-response atlas"
    )


def section_text(document: Document, start_heading: str, end_heading: str) -> List[str]:
    start = find_heading_paragraph(document, start_heading)
    end = find_heading_paragraph(document, end_heading)
    paragraphs = document.paragraphs
    start_index = next(
        index for index, paragraph in enumerate(paragraphs)
        if paragraph._p is start._p
    )
    end_index = next(
        index for index, paragraph in enumerate(paragraphs)
        if paragraph._p is end._p
    )
    return [normalize(paragraph.text) for paragraph in paragraphs[start_index + 1:end_index]]


def figure_caption_texts(document: Document) -> List[str]:
    return [
        normalize(paragraph.text)
        for paragraph in document.paragraphs
        if FIGURE_CAPTION_RE.match(normalize(paragraph.text))
    ]


def embedded_media_hashes(docx_path: Path) -> List[str]:
    with zipfile.ZipFile(docx_path) as archive:
        media = sorted(
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        )
        return [hashlib.sha256(archive.read(name)).hexdigest() for name in media]


def section_word_count(document: Document, start_heading: str, end_heading: str) -> int:
    text = " ".join(section_text(document, start_heading, end_heading))
    return len(re.findall(r"\b[\w'-]+\b", text))


def run_command(
    command: Sequence[str],
    env: Optional[Dict[str, str]] = None,
    timeout_seconds: int = 180,
):
    return subprocess.run(
        list(command),
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        env=env,
        check=False,
        timeout=timeout_seconds,
    )


def make_contact_sheet(
    pages: Sequence[Path],
    output: Path,
    columns: int = 3,
    cell_width: int = 600,
    padding: int = 20,
) -> None:
    images: List[Image.Image] = []
    for page in pages:
        image = Image.open(page).convert("RGB")
        ratio = cell_width / image.width
        images.append(
            image.resize((cell_width, max(1, int(image.height * ratio))))
        )

    rows = (len(images) + columns - 1) // columns
    row_heights = [
        max(image.height for image in images[row * columns:(row + 1) * columns])
        for row in range(rows)
    ]
    canvas = Image.new(
        "RGB",
        (
            columns * cell_width + (columns + 1) * padding,
            sum(row_heights) + (rows + 1) * padding,
        ),
        "white",
    )

    y = padding
    for row in range(rows):
        x = padding
        for image in images[row * columns:(row + 1) * columns]:
            canvas.paste(image, (x, y))
            x += cell_width + padding
        y += row_heights[row] + padding

    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output)


def render_docx(docx_path: Path, outdir: Path) -> Dict[str, object]:
    outdir.mkdir(parents=True, exist_ok=True)
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    pdftocairo = shutil.which("pdftocairo")
    pdftoppm = shutil.which("pdftoppm")

    if not libreoffice:
        return {
            "render_pass": False,
            "reason": "LibreOffice/soffice not found",
            "page_count": 0,
            "contact_sheet": "",
        }
    if not pdftocairo and not pdftoppm:
        return {
            "render_pass": False,
            "reason": "Neither pdftocairo nor pdftoppm was found",
            "page_count": 0,
            "contact_sheet": "",
        }

    for old in outdir.glob("page-*.png"):
        old.unlink()

    with tempfile.TemporaryDirectory(prefix="u27b3d21_lo_") as temp_home:
        env = os.environ.copy()
        env["HOME"] = temp_home
        profile_uri = Path(temp_home).resolve().as_uri()
        try:
            conversion = run_command(
                [
                    libreoffice,
                    "--headless",
                    f"-env:UserInstallation={profile_uri}",
                    "--convert-to",
                    "pdf:writer_pdf_Export",
                    "--outdir",
                    str(outdir),
                    str(docx_path),
                ],
                env=env,
                timeout_seconds=180,
            )
        except subprocess.TimeoutExpired:
            return {
                "render_pass": False,
                "reason": "LibreOffice PDF conversion timed out after 180 seconds",
                "page_count": 0,
                "contact_sheet": "",
            }

    pdf_path = outdir / f"{docx_path.stem}.pdf"
    if conversion.returncode != 0 or not pdf_path.exists():
        return {
            "render_pass": False,
            "reason": conversion.stderr.strip() or conversion.stdout.strip(),
            "page_count": 0,
            "contact_sheet": "",
        }

    raster_command = (
        [
            pdftocairo,
            "-png",
            "-r",
            "60",
            str(pdf_path),
            str(outdir / "page"),
        ]
        if pdftocairo
        else [
            pdftoppm,
            "-png",
            "-r",
            "60",
            str(pdf_path),
            str(outdir / "page"),
        ]
    )
    try:
        raster = run_command(raster_command, timeout_seconds=180)
    except subprocess.TimeoutExpired:
        return {
            "render_pass": False,
            "reason": "PDF rasterization timed out after 180 seconds",
            "page_count": 0,
            "contact_sheet": "",
        }
    pages = sorted(
        outdir.glob("page-*.png"),
        key=lambda path: int(re.search(r"(\d+)$", path.stem).group(1)),
    )
    if raster.returncode != 0 or not pages:
        return {
            "render_pass": False,
            "reason": raster.stderr.strip() or raster.stdout.strip(),
            "page_count": len(pages),
            "contact_sheet": "",
        }

    contact = outdir / "UTI_HostOmics_U27B3D21_render_contact_sheet.png"
    make_contact_sheet(pages, contact)
    return {
        "render_pass": True,
        "reason": "Rendered with LibreOffice and pdftoppm",
        "page_count": len(pages),
        "contact_sheet": str(contact),
        "pdf_path": str(pdf_path),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root", default="__UTI_HOSTOMICS_PROJECT_ROOT__"
    )
    parser.add_argument("--source", default=DEFAULT_SOURCE)
    args = parser.parse_args()

    project = Path(args.project_root).resolve()
    source = Path(args.source).resolve()
    if not source.exists():
        raise FileNotFoundError(f"Source manuscript not found: {source}")

    source_module, module_path = load_u27b3d2_module(project)

    outdoc = project / "09_manuscript_docx" / TAG
    outtables = project / "06_tables" / TAG
    outmeta = project / "03_metadata" / TAG
    outresults = project / "05_results" / TAG
    renderdir = outdoc / "render_qa"
    for directory in (outdoc, outtables, outmeta, outresults, renderdir):
        directory.mkdir(parents=True, exist_ok=True)

    source_hash_before = sha256(source)
    source_document = Document(source)
    frozen_results_before = section_text(source_document, "Results", "Discussion")
    frozen_captions_before = figure_caption_texts(source_document)
    media_hashes_before = embedded_media_hashes(source)

    document = Document(source)

    misplaced_removed = remove_misplaced_discussion_between_figures(document)
    update_front_matter_headers_footers(document, source_module)

    replacement_counts = {
        "abstract_removed_elements": replace_section(
            document,
            "Abstract",
            "Introduction",
            source_module.ABSTRACT_PARAGRAPHS,
        ),
        "introduction_removed_elements": replace_section(
            document,
            "Introduction",
            "Methods",
            source_module.INTRODUCTION_PARAGRAPHS,
        ),
        "methods_removed_elements": replace_section_with_subheadings(
            document,
            "Methods",
            "Results",
            source_module.METHODS_SECTIONS,
        ),
        "discussion_removed_elements": replace_discussion_block(
            document, source_module
        ),
        "data_availability_removed_elements": replace_section(
            document,
            "Data availability",
            "Code availability",
            [source_module.DATA_AVAILABILITY],
        ),
        "code_availability_removed_elements": replace_section(
            document,
            "Code availability",
            "Ethics statement",
            [source_module.CODE_AVAILABILITY],
        ),
        "funding_removed_elements": replace_section(
            document,
            "Funding",
            "Acknowledgements",
            [source_module.FUNDING],
        ),
        "supplementary_removed_elements": replace_supplementary_section(
            document, source_module
        ),
    }

    legacy_cleanup = remove_legacy_tracking_tables(document)
    shifted_cleanup = cleanup_shifted_section_duplicates(document, source_module)

    output = outdoc / OUTPUT_NAME
    if output.exists():
        output.unlink()
    document.save(output)

    source_hash_after = sha256(source)
    output_hash = sha256(output)

    out_document = Document(output)
    frozen_results_after = section_text(out_document, "Results", "Discussion")
    frozen_captions_after = figure_caption_texts(out_document)
    media_hashes_after = embedded_media_hashes(output)

    results_preserved = frozen_results_before == frozen_results_after
    captions_preserved = (
        frozen_captions_before == frozen_captions_after
        and len(frozen_captions_after) == 8
    )
    images_preserved = media_hashes_before == media_hashes_after
    source_unchanged = source_hash_before == source_hash_after

    section_metrics = {
        "abstract_words": section_word_count(out_document, "Abstract", "Introduction"),
        "introduction_words": section_word_count(out_document, "Introduction", "Methods"),
        "methods_words": section_word_count(out_document, "Methods", "Results"),
        "results_words": section_word_count(out_document, "Results", "Discussion"),
        "discussion_words": section_word_count(out_document, "Discussion", "Data availability"),
        "data_availability_words": section_word_count(out_document, "Data availability", "Code availability"),
        "code_availability_words": section_word_count(out_document, "Code availability", "Ethics statement"),
        "funding_words": section_word_count(out_document, "Funding", "Acknowledgements"),
        "supplementary_words": section_word_count(
            out_document,
            "Supplementary tables",
            "Remaining reference gaps after citation-key cleanup",
        ),
    }

    full_text = "\n".join(normalize(paragraph.text) for paragraph in out_document.paragraphs)
    output_paragraphs = out_document.paragraphs
    figure3_index = next(
        index for index, paragraph in enumerate(output_paragraphs)
        if normalize(paragraph.text).startswith("Figure 3.")
    )
    figure4_index = next(
        index for index, paragraph in enumerate(output_paragraphs)
        if normalize(paragraph.text).startswith("Figure 4.")
    )
    misplaced_discussion_remaining = any(
        normalize(paragraph.text) == "Principal findings"
        for paragraph in output_paragraphs[figure3_index + 1:figure4_index]
    )

    obsolete_terms = [
        "GSE186800",
        "GSE261018",
        "04_scripts/",
        "17 modules",
        "17 validated modules",
        "Draft manuscript v5.2",
        "Pre-Zotero manuscript v5.2",
    ]
    obsolete_audit = pd.DataFrame(
        [
            {
                "term": term,
                "occurrence_count": len(
                    re.findall(re.escape(term), full_text, flags=re.I)
                ),
                "absent": re.search(
                    re.escape(term), full_text, flags=re.I
                )
                is None,
            }
            for term in obsolete_terms
        ]
    )
    obsolete_audit.to_csv(
        outtables / "UTI_HostOmics_U27B3D21_obsolete_term_audit.tsv",
        sep="\t",
        index=False,
    )

    required_audits = pd.DataFrame(
        [
            {
                "audit_id": "abstract_nonempty",
                "value": section_metrics["abstract_words"],
                "pass": section_metrics["abstract_words"] >= 180,
            },
            {
                "audit_id": "introduction_nonempty",
                "value": section_metrics["introduction_words"],
                "pass": section_metrics["introduction_words"] >= 250,
            },
            {
                "audit_id": "methods_nonempty",
                "value": section_metrics["methods_words"],
                "pass": section_metrics["methods_words"] >= 500,
            },
            {
                "audit_id": "discussion_nonempty",
                "value": section_metrics["discussion_words"],
                "pass": section_metrics["discussion_words"] >= 650,
            },
            {
                "audit_id": "data_availability_nonempty",
                "value": section_metrics["data_availability_words"],
                "pass": section_metrics["data_availability_words"] >= 35,
            },
            {
                "audit_id": "code_availability_nonempty",
                "value": section_metrics["code_availability_words"],
                "pass": section_metrics["code_availability_words"] >= 25,
            },
            {
                "audit_id": "funding_nonempty",
                "value": section_metrics["funding_words"],
                "pass": section_metrics["funding_words"] >= 20,
            },
            {
                "audit_id": "supplementary_items_present",
                "value": section_metrics["supplementary_words"],
                "pass": all(
                    f"Table S{number}." in full_text
                    for number in range(1, 11)
                ),
            },
            {
                "audit_id": "misplaced_discussion_absent",
                "value": misplaced_discussion_remaining,
                "pass": not misplaced_discussion_remaining,
            },
            {
                "audit_id": "legacy_figure_selection_table_removed",
                "value": legacy_cleanup["legacy_figure_selection_tables_removed"],
                "pass": "primary_asset" not in full_text,
            },
            {
                "audit_id": "legacy_open_issue_table_removed",
                "value": legacy_cleanup["legacy_open_issue_elements_removed"],
                "pass": "Open issues retained for final submission" not in full_text,
            },
        ]
    )
    required_audits.to_csv(
        outtables / "UTI_HostOmics_U27B3D21_section_structure_audit.tsv",
        sep="\t",
        index=False,
    )

    preservation = pd.DataFrame(
        [
            {
                "source_path": str(source),
                "source_sha256_before": source_hash_before,
                "source_sha256_after": source_hash_after,
                "source_unchanged": source_unchanged,
                "output_path": str(output),
                "output_sha256": output_hash,
                "results_preserved": results_preserved,
                "figure_legends_preserved": captions_preserved,
                "embedded_image_hashes_preserved": images_preserved,
                "embedded_images_before": len(media_hashes_before),
                "embedded_images_after": len(media_hashes_after),
                "misplaced_discussion_elements_removed": misplaced_removed,
                **replacement_counts,
                **legacy_cleanup,
                **shifted_cleanup,
            }
        ]
    )
    preservation.to_csv(
        outtables / "UTI_HostOmics_U27B3D21_preservation_repair_audit.tsv",
        sep="\t",
        index=False,
    )

    render = render_docx(output, renderdir)
    pd.DataFrame([render]).to_csv(
        outtables / "UTI_HostOmics_U27B3D21_render_audit.tsv",
        sep="\t",
        index=False,
    )

    structural_pass = bool(
        source_unchanged
        and results_preserved
        and captions_preserved
        and images_preserved
        and required_audits["pass"].all()
        and obsolete_audit["absent"].all()
    )

    if structural_pass and render["render_pass"]:
        decision = "READY_FOR_U27B3D3_V6_MANUSCRIPT_VISUAL_AND_SCIENTIFIC_AUDIT"
    elif structural_pass:
        decision = "V6_1_STRUCTURAL_REPAIR_COMPLETE_RENDER_QA_PENDING"
    else:
        decision = "TARGETED_U27B3D21_STRUCTURE_OR_PRESERVATION_REPAIR_REQUIRED"

    pd.DataFrame(
        [
            {
                "phase": "U27B3D2.1",
                "decision": decision,
                "source_unchanged": source_unchanged,
                "results_preserved": results_preserved,
                "figures_and_legends_preserved": captions_preserved,
                "embedded_images_preserved": images_preserved,
                "section_structure_audits_pass": bool(required_audits["pass"].all()),
                "obsolete_terms_absent": bool(obsolete_audit["absent"].all()),
                "render_pass": render["render_pass"],
                "page_count": render["page_count"],
                "new_derivative_created": True,
                "source_modified": False,
                "next_phase": (
                    "U27B3D3 visually inspect the corrected v6.1 manuscript and "
                    "perform full scientific coherence audit"
                    if decision.startswith("READY_FOR_U27B3D3")
                    else "Inspect the structural and render audits"
                ),
            }
        ]
    ).to_csv(
        outtables / "UTI_HostOmics_U27B3D21_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    pd.DataFrame(
        [
            {"field": "source_path", "value": str(source)},
            {"field": "source_sha256", "value": source_hash_before},
            {"field": "output_path", "value": str(output)},
            {"field": "output_sha256", "value": output_hash},
            {"field": "scientific_text_source_script", "value": str(module_path)},
            {"field": "version", "value": VERSION},
        ]
    ).to_csv(
        outmeta / "UTI_HostOmics_U27B3D21_derivative_record.tsv",
        sep="\t",
        index=False,
    )

    report_path = (
        outresults
        / "UTI_HostOmics_U27B3D21_corrected_v6_reconstruction_report.md"
    )
    report_path.write_text(
        "# Phase U27B3D2.1 - Corrected v6.1 manuscript reconstruction\n\n"
        f"- Version: `{VERSION}`\n"
        f"- Decision: **{decision}**\n"
        f"- Source: `{source}`\n"
        f"- Output: `{output}`\n"
        f"- Results preserved: **{results_preserved}**.\n"
        f"- Figures and legends preserved: **{captions_preserved}**.\n"
        f"- Embedded image hashes preserved: **{images_preserved}**.\n"
        f"- Abstract words: **{section_metrics['abstract_words']}**.\n"
        f"- Introduction words: **{section_metrics['introduction_words']}**.\n"
        f"- Methods words: **{section_metrics['methods_words']}**.\n"
        f"- Discussion words: **{section_metrics['discussion_words']}**.\n"
        f"- Misplaced Discussion elements removed: **{misplaced_removed}**.\n"
        f"- Legacy selected-figure tables removed: "
        f"**{legacy_cleanup['legacy_figure_selection_tables_removed']}**.\n"
        f"- Legacy open-issue elements removed: "
        f"**{legacy_cleanup['legacy_open_issue_elements_removed']}**.\n"
        f"- Shifted Code paragraphs removed from Ethics: "
        f"**{shifted_cleanup['duplicate_code_paragraphs_removed_from_ethics']}**.\n"
        f"- Shifted Funding paragraphs removed from Acknowledgements: "
        f"**{shifted_cleanup['duplicate_funding_paragraphs_removed_from_acknowledgements']}**.\n"
        f"- Duplicate supplementary items removed: "
        f"**{shifted_cleanup['duplicate_supplementary_items_removed']}**.\n"
        f"- Redundant page-break paragraphs removed: "
        f"**{shifted_cleanup['redundant_page_break_paragraphs_removed']}**.\n"
        f"- Render pass: **{render['render_pass']}**.\n"
        f"- Render pages: **{render['page_count']}**.\n"
        f"- Contact sheet: `{render.get('contact_sheet', '')}`.\n\n"
        "## Correction boundary\n\n"
        "The v6.0 output contained correctly reconstructed scientific text but "
        "placed several sections at stale downstream paragraph indices. This "
        "phase reuses the same validated scientific text, repairs section "
        "placement with stable anchors, removes duplicate misplaced content and "
        "legacy tracking tables, and preserves the frozen Results and figure "
        "package byte-for-byte.\n",
        encoding="utf-8",
    )

    manifest = {
        "version": VERSION,
        "decision": decision,
        "source": str(source),
        "source_sha256": source_hash_before,
        "output": str(output),
        "output_sha256": output_hash,
        "results_preserved": results_preserved,
        "figures_legends_preserved": captions_preserved,
        "images_preserved": images_preserved,
        "section_structure_pass": bool(required_audits["pass"].all()),
        "obsolete_terms_absent": bool(obsolete_audit["absent"].all()),
        "render_pass": render["render_pass"],
        "page_count": render["page_count"],
        "contact_sheet": render.get("contact_sheet", ""),
        "source_modified": False,
    }
    (
        outresults / "UTI_HostOmics_U27B3D21_run_manifest.json"
    ).write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    log(f"Results preserved: {results_preserved}")
    log(f"Figures/legends preserved: {captions_preserved}")
    log(f"Embedded images preserved: {images_preserved}")
    log(f"Section structure pass: {bool(required_audits['pass'].all())}")
    log(f"Render pass: {render['render_pass']}")
    log(f"Decision: {decision}")
    log(f"Output: {output}")
    log(f"Report: {report_path}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3D2.1] ERROR: {exc}", file=sys.stderr)
        raise
