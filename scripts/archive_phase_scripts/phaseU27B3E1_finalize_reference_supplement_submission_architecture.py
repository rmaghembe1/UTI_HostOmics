#!/usr/bin/env python3
"""
Phase U27B3E1
Separate internal reference/Zotero tracking material from the scientifically
harmonized v6.1 manuscript and construct the submission/supplement architecture.

This phase is non-destructive. It creates:
1. a cleaned v6.2 submission-architecture manuscript derivative with the
   internal reference-gap register and Zotero working table removed;
2. an internal reference-tracking companion DOCX plus TSV extracts;
3. a Supplementary Tables S1-S10 manifest and reproducible source-candidate
   registry;
4. a submission-finalization blocker checklist;
5. render QA for the cleaned manuscript derivative.

The phase does not finalize authors, affiliations, repository URLs, competing
interests, acknowledgements, contributions or citations. Those remain explicit
U27B3E2 tasks. Results, Figures 1-8 and definitive legends are preserved.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple
from xml.etree import ElementTree as ET

import pandas as pd

try:
    from PIL import Image
except ImportError as exc:
    raise RuntimeError("Pillow is required for contact-sheet generation.") from exc

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as exc:
    raise RuntimeError("python-docx is required for U27B3E1.") from exc


VERSION = "U27B3E1_v1.0_2026-07-16"
TAG = "phaseU27B3E1_reference_supplement_submission_architecture"

DEFAULT_SOURCE = (
    "__UTI_HOSTOMICS_PROJECT_ROOT__/"
    "09_manuscript_docx/phaseU27B3D21_corrected_v6_reconstruction/"
    "UTI_HostOmics_preZotero_manuscript_v6_1_"
    "U27B3D21_scientifically_harmonized_corrected.docx"
)

OUTPUT_FILENAME = (
    "UTI_HostOmics_preZotero_manuscript_v6_2_"
    "U27B3E1_submission_architecture_cleaned.docx"
)

INTERNAL_COMPANION_FILENAME = (
    "UTI_HostOmics_U27B3E1_internal_reference_and_Zotero_tracking.docx"
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS = {"w": W_NS, "r": R_NS, "a": A_NS}

GAP_HEADING = "Remaining reference gaps after citation-key cleanup"
REFERENCE_TABLE_HEADING = "Reference table for Zotero finalization"
SUPPLEMENT_HEADING = "Supplementary tables"

FIGURE_CAPTION_RE = re.compile(r"^Figure\s+([1-8])\.\s+", re.I)
SUPPLEMENT_RE = re.compile(r"^Table\s+S(\d+)\.\s*(.+)$", re.I)
GAP_RE = re.compile(r"^(U\d+[A-Z]?_GAP_\d+)\s*\(([^)]+)\):\s*(.+)$", re.I)

SUBMISSION_BLOCKERS = [
    (
        "author_affiliation_placeholder",
        "Affiliations and co-author list to be finalized",
        "Confirm final author order, affiliations and corresponding-author details.",
    ),
    (
        "repository_placeholder",
        "A public repository archive will be finalized before journal submission",
        "Create/finalize the public repository and replace provisional wording with a persistent URL/DOI.",
    ),
    (
        "contribution_placeholder",
        "should be updated after co-author review",
        "Confirm CRediT contributions with all authors.",
    ),
    (
        "competing_interest_placeholder",
        "should be finalized before submission",
        "Insert the final competing-interest declaration.",
    ),
    (
        "acknowledgement_placeholder",
        "Acknowledgements should be added before submission",
        "Insert the final acknowledgements.",
    ),
]

SUPPLEMENT_KEYWORDS: Dict[int, List[str]] = {
    1: ["dataset", "architecture", "sample", "design", "inclusion", "metadata"],
    2: ["78", "submodule", "library", "axis", "expanded"],
    3: ["module", "effect", "factorial", "adjusted", "contrast"],
    4: ["cross-dataset", "recurrence", "concordance", "evidence", "class"],
    5: ["gse280297", "pregnancy", "tissue", "preterm", "term", "outcome"],
    6: ["gse252321", "quality", "cluster", "marker", "refined", "subtype"],
    7: ["broad", "refined", "pseudobulk", "localization", "cellular"],
    8: ["complement", "endocrine", "metabolic", "attribution", "cellular"],
    9: ["figure", "source", "value", "manifest", "panel", "provenance"],
    10: ["interpretation", "boundary", "sensitivity", "claim", "traceability", "caveat"],
}

SEARCH_EXTENSIONS = {".tsv", ".csv", ".xlsx", ".xls", ".json", ".md"}
SKIP_PARTS = {
    ".git",
    "08_logs",
    "09_archives",
    "09_manuscript_docx",
    "render_qa",
    "__pycache__",
}


def log(message: str) -> None:
    print(f"[U27B3E1] {message}", flush=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            block = handle.read(1024 * 1024)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def normalize(text: object) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def element_text(element) -> str:
    pieces: List[str] = []
    for node in element.iter():
        local = node.tag.rsplit("}", 1)[-1]
        if local == "t":
            pieces.append(node.text or "")
        elif local == "tab":
            pieces.append("\t")
        elif local in {"br", "cr"}:
            pieces.append("\n")
    return normalize("".join(pieces))


def has_drawing(element) -> bool:
    return any(
        node.tag.rsplit("}", 1)[-1] in {"drawing", "pict"}
        for node in element.iter()
    )


def has_page_break(element) -> bool:
    for node in element.iter():
        if node.tag == qn("w:br") and node.attrib.get(qn("w:type"), "") == "page":
            return True
    ppr = element.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:pageBreakBefore")) is not None


def find_body_index(body, exact_text: str) -> int:
    matches = [
        index
        for index, element in enumerate(list(body))
        if element.tag == qn("w:p")
        and element_text(element).lower() == exact_text.lower()
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph {exact_text!r}; observed indices {matches}."
        )
    return matches[0]


def find_first_figure_block_index(body, after_index: int) -> int:
    children = list(body)
    candidates: List[int] = []
    for index in range(after_index + 1, len(children)):
        element = children[index]
        if element.tag != qn("w:p"):
            continue
        text = element_text(element)
        if has_drawing(element) or FIGURE_CAPTION_RE.match(text):
            candidates.append(index)
            break
    if not candidates:
        raise RuntimeError("Could not resolve the first frozen figure block.")
    return candidates[0]


def trim_blank_edges(elements: List[object]) -> List[object]:
    items = list(elements)
    while items:
        text = element_text(items[0])
        if items[0].tag == qn("w:p") and not text:
            items.pop(0)
        else:
            break
    while items:
        text = element_text(items[-1])
        if items[-1].tag == qn("w:p") and not text:
            items.pop()
        else:
            break
    return items


def extract_results_text(document: Document) -> List[str]:
    paragraphs = [normalize(paragraph.text) for paragraph in document.paragraphs]

    def unique_heading(text: str) -> int:
        matches = [
            index
            for index, value in enumerate(paragraphs)
            if value.lower().rstrip(".") == text.lower()
        ]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one {text} heading; observed {matches}.")
        return matches[0]

    start = unique_heading("Results")
    end = unique_heading("Discussion")
    if start >= end:
        raise RuntimeError("Results does not precede Discussion.")
    return paragraphs[start + 1 : end]


def media_hashes(path: Path) -> List[str]:
    with zipfile.ZipFile(path) as archive:
        names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        )
        return [hashlib.sha256(archive.read(name)).hexdigest() for name in names]


def figure_caption_texts(document: Document) -> List[str]:
    return [
        normalize(paragraph.text)
        for paragraph in document.paragraphs
        if FIGURE_CAPTION_RE.match(normalize(paragraph.text))
    ]


def set_document_version(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = normalize(paragraph.text)
        if text.startswith("Pre-Zotero manuscript v6.1"):
            paragraph.text = (
                "Pre-Zotero manuscript v6.2 | Submission architecture cleaned; "
                "internal reference tracking separated"
            )
        elif text.startswith("Updated:"):
            paragraph.text = "Updated: 2026-07-16"

    for section in document.sections:
        for collection in (section.header.paragraphs, section.footer.paragraphs):
            for paragraph in collection:
                if "v6.1" in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace("v6.1", "v6.2")


def extract_gap_records(gap_texts: Sequence[str]) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for text in gap_texts:
        match = GAP_RE.match(normalize(text))
        if not match:
            continue
        detail = match.group(3)
        citation_match = re.search(r"Citation/key:\s*`?([^`]+)`?", detail, re.I)
        rows.append(
            {
                "gap_id": match.group(1),
                "priority": match.group(2),
                "description": detail,
                "citation_key": citation_match.group(1).strip() if citation_match else "",
                "status": "OPEN",
            }
        )
    return pd.DataFrame(rows)


def table_to_dataframe(table) -> pd.DataFrame:
    rows = [
        [normalize(cell.text) for cell in row.cells]
        for row in table.rows
    ]
    if not rows:
        return pd.DataFrame()
    headers = rows[0]
    deduped: List[str] = []
    seen: Dict[str, int] = {}
    for index, header in enumerate(headers):
        name = header or f"column_{index + 1}"
        seen[name] = seen.get(name, 0) + 1
        deduped.append(name if seen[name] == 1 else f"{name}_{seen[name]}")
    return pd.DataFrame(rows[1:], columns=deduped)


def build_internal_companion(
    output: Path,
    source_path: Path,
    source_hash: str,
    gap_texts: Sequence[str],
    reference_frame: pd.DataFrame,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("UTI HostOmics - Internal Reference and Zotero Tracking")
    run.bold = True
    run.font.size = Pt(15)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Source: {source_path.name}\nSHA256: {source_hash}").font.size = Pt(8)

    document.add_heading(GAP_HEADING, level=1)
    for text in gap_texts:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(4)

    document.add_heading(REFERENCE_TABLE_HEADING, level=1)
    if reference_frame.empty:
        document.add_paragraph("No reference table was present in the source manuscript.")
    else:
        table = document.add_table(
            rows=1,
            cols=len(reference_frame.columns),
            style="Table Grid",
        )
        for index, column in enumerate(reference_frame.columns):
            table.rows[0].cells[index].text = str(column)
        for _, row in reference_frame.iterrows():
            cells = table.add_row().cells
            for index, column in enumerate(reference_frame.columns):
                cells[index].text = str(row[column])
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    document.save(output)


def parse_supplement_descriptions(document: Document) -> pd.DataFrame:
    paragraphs = [normalize(paragraph.text) for paragraph in document.paragraphs]
    heading_matches = [
        index
        for index, text in enumerate(paragraphs)
        if text.lower() == SUPPLEMENT_HEADING.lower()
    ]
    gap_matches = [
        index
        for index, text in enumerate(paragraphs)
        if text.lower() == GAP_HEADING.lower()
    ]
    if len(heading_matches) != 1 or len(gap_matches) != 1:
        raise RuntimeError("Could not resolve Supplementary tables -> reference gaps boundary.")

    rows: List[Dict[str, object]] = []
    for text in paragraphs[heading_matches[0] + 1 : gap_matches[0]]:
        match = SUPPLEMENT_RE.match(text)
        if match:
            number = int(match.group(1))
            rows.append(
                {
                    "supplementary_table": f"Table S{number}",
                    "table_number": number,
                    "description": match.group(2),
                    "status": "SOURCE_CONFIRMATION_REQUIRED",
                }
            )
    frame = pd.DataFrame(rows).sort_values("table_number") if rows else pd.DataFrame()
    return frame


def iter_candidate_files(project: Path, output_roots: Sequence[Path]) -> Iterable[Path]:
    output_resolved = [path.resolve() for path in output_roots]
    for path in project.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in SEARCH_EXTENSIONS:
            continue
        parts = set(path.relative_to(project).parts)
        if parts & SKIP_PARTS:
            continue
        resolved = path.resolve()
        if any(root == resolved or root in resolved.parents for root in output_resolved):
            continue
        yield path


def candidate_score(path: Path, keywords: Sequence[str]) -> Tuple[int, str]:
    haystack = str(path).lower().replace("_", " ").replace("-", " ")
    matched = [keyword for keyword in keywords if keyword.lower() in haystack]
    score = len(matched) * 3
    filename = path.name.lower()
    score += sum(2 for keyword in matched if keyword.lower() in filename)
    if "final" in filename or "frozen" in filename:
        score += 2
    if "audit" in filename and "audit" not in keywords:
        score -= 1
    return score, "; ".join(matched)


def build_candidate_registry(
    project: Path,
    supplementary: pd.DataFrame,
    output_roots: Sequence[Path],
) -> pd.DataFrame:
    candidates = list(iter_candidate_files(project, output_roots))
    rows: List[Dict[str, object]] = []

    for _, supplement_row in supplementary.iterrows():
        number = int(supplement_row["table_number"])
        keywords = SUPPLEMENT_KEYWORDS[number]
        scored = []
        for path in candidates:
            score, matched = candidate_score(path, keywords)
            if score <= 0:
                continue
            scored.append((score, path.stat().st_mtime, path, matched))

        scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
        for rank, (score, mtime, path, matched) in enumerate(scored[:12], start=1):
            rows.append(
                {
                    "supplementary_table": f"Table S{number}",
                    "candidate_rank": rank,
                    "candidate_score": score,
                    "matched_keywords": matched,
                    "candidate_path": str(path),
                    "file_type": path.suffix.lower().lstrip("."),
                    "size_bytes": path.stat().st_size,
                    "modified_time": datetime.fromtimestamp(mtime).isoformat(timespec="seconds"),
                }
            )

        if not scored:
            rows.append(
                {
                    "supplementary_table": f"Table S{number}",
                    "candidate_rank": "",
                    "candidate_score": 0,
                    "matched_keywords": "",
                    "candidate_path": "",
                    "file_type": "",
                    "size_bytes": 0,
                    "modified_time": "",
                }
            )

    return pd.DataFrame(rows)


def run_command(command: Sequence[str], env: Optional[Dict[str, str]] = None) -> subprocess.CompletedProcess:
    return subprocess.run(
        list(command),
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
        env=env,
    )


def make_contact_sheet(paths: Sequence[Path], output: Path, columns: int = 3) -> None:
    if not paths:
        raise RuntimeError("No page images were available for contact-sheet generation.")
    cell_width = 600
    padding = 24
    images: List[Image.Image] = []
    for path in paths:
        image = Image.open(path).convert("RGB")
        ratio = cell_width / image.width
        images.append(image.resize((cell_width, max(1, int(image.height * ratio)))))
    rows = (len(images) + columns - 1) // columns
    row_heights = []
    for row_index in range(rows):
        subset = images[row_index * columns : (row_index + 1) * columns]
        row_heights.append(max(image.height for image in subset))
    canvas = Image.new(
        "RGB",
        (
            columns * cell_width + (columns + 1) * padding,
            sum(row_heights) + (rows + 1) * padding,
        ),
        "white",
    )
    y = padding
    for row_index in range(rows):
        x = padding
        subset = images[row_index * columns : (row_index + 1) * columns]
        for image in subset:
            canvas.paste(image, (x, y))
            x += cell_width + padding
        y += row_heights[row_index] + padding
    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output)


def render_docx(docx_path: Path, output_dir: Path) -> Dict[str, object]:
    output_dir.mkdir(parents=True, exist_ok=True)
    for stale in list(output_dir.glob("page-*.png")) + list(output_dir.glob("*.pdf")) + list(output_dir.glob("*contact_sheet*.png")):
        try:
            stale.unlink()
        except FileNotFoundError:
            pass
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    pdftoppm = shutil.which("pdftoppm")
    if not libreoffice or not pdftoppm:
        return {
            "render_pass": False,
            "reason": "LibreOffice/soffice or pdftoppm not found",
            "page_count": 0,
            "pdf_path": "",
            "contact_sheet": "",
        }

    with tempfile.TemporaryDirectory(prefix="u27b3e1_lo_") as tmp:
        env = os.environ.copy()
        env["HOME"] = tmp
        profile_uri = Path(tmp).resolve().as_uri()
        conversion = run_command(
            [
                libreoffice,
                "--headless",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            env=env,
        )

    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if conversion.returncode != 0 or not pdf_path.exists():
        return {
            "render_pass": False,
            "reason": conversion.stderr.strip() or conversion.stdout.strip(),
            "page_count": 0,
            "pdf_path": str(pdf_path),
            "contact_sheet": "",
        }

    raster = run_command(
        [
            pdftoppm,
            "-png",
            "-r",
            "100",
            str(pdf_path),
            str(output_dir / "page"),
        ]
    )
    page_paths = sorted(
        output_dir.glob("page-*.png"),
        key=lambda path: int(re.search(r"(\d+)$", path.stem).group(1)),
    )
    if raster.returncode != 0 or not page_paths:
        return {
            "render_pass": False,
            "reason": raster.stderr.strip() or raster.stdout.strip(),
            "page_count": len(page_paths),
            "pdf_path": str(pdf_path),
            "contact_sheet": "",
        }

    contact_sheet = output_dir / "UTI_HostOmics_U27B3E1_render_contact_sheet.png"
    make_contact_sheet(page_paths, contact_sheet)
    return {
        "render_pass": contact_sheet.exists() and contact_sheet.stat().st_size > 0,
        "reason": "Render files created",
        "page_count": len(page_paths),
        "pdf_path": str(pdf_path),
        "contact_sheet": str(contact_sheet),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-root", default="__UTI_HOSTOMICS_PROJECT_ROOT__")
    parser.add_argument("--source", default=DEFAULT_SOURCE)
    args = parser.parse_args()

    project = Path(args.project_root).resolve()
    source = Path(args.source).resolve()
    if not source.exists():
        raise FileNotFoundError(f"Scientifically harmonized v6.1 manuscript not found: {source}")

    outdocx = project / "09_manuscript_docx" / TAG
    outtables = project / "06_tables" / TAG
    outmetadata = project / "03_metadata" / TAG
    outresults = project / "05_results" / TAG
    render_dir = outdocx / "render_qa"
    internal_dir = outdocx / "internal_reference_tracking"
    supplementary_dir = outdocx / "supplementary_architecture"

    for directory in (
        outdocx,
        outtables,
        outmetadata,
        outresults,
        render_dir,
        internal_dir,
        supplementary_dir,
    ):
        directory.mkdir(parents=True, exist_ok=True)

    source_hash_before = sha256(source)
    source_document = Document(source)
    source_results = extract_results_text(source_document)
    source_captions = figure_caption_texts(source_document)
    source_media_hashes = media_hashes(source)
    supplementary_manifest = parse_supplement_descriptions(source_document)

    if len(supplementary_manifest) != 10:
        raise RuntimeError(
            f"Expected Supplementary Tables S1-S10; observed {len(supplementary_manifest)} entries."
        )

    if len(source_document.tables) != 1:
        raise RuntimeError(
            f"Expected one internal Zotero reference table; observed {len(source_document.tables)} tables."
        )
    reference_frame = table_to_dataframe(source_document.tables[0])

    document = Document(source)
    body = document._element.body
    children = list(body)

    gap_start = find_body_index(body, GAP_HEADING)
    figure_start = find_first_figure_block_index(body, gap_start)
    reference_start = find_body_index(body, REFERENCE_TABLE_HEADING)

    if not (gap_start < figure_start < reference_start):
        raise RuntimeError(
            f"Unexpected internal-section ordering: gap={gap_start}, figures={figure_start}, reference={reference_start}."
        )

    gap_elements = trim_blank_edges(children[gap_start:figure_start])
    gap_texts = [element_text(element) for element in gap_elements if element_text(element)]

    # Remove the internal Zotero heading and every following body element except sectPr.
    children_now = list(body)
    reference_start = find_body_index(body, REFERENCE_TABLE_HEADING)
    for element in children_now[reference_start:]:
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)

    # Remove the internal reference-gap block and adjacent empty page-break paragraphs.
    children_now = list(body)
    gap_start = find_body_index(body, GAP_HEADING)
    figure_start = find_first_figure_block_index(body, gap_start)
    remove_start = gap_start
    while remove_start > 0:
        previous = children_now[remove_start - 1]
        if previous.tag == qn("w:p") and not element_text(previous) and has_page_break(previous):
            remove_start -= 1
        else:
            break
    for element in children_now[remove_start:figure_start]:
        if element.tag != qn("w:sectPr"):
            body.remove(element)

    set_document_version(document)

    output_path = outdocx / OUTPUT_FILENAME
    if output_path.exists():
        output_path.unlink()
    document.save(output_path)

    source_hash_after = sha256(source)
    output_document = Document(output_path)
    output_results = extract_results_text(output_document)
    output_captions = figure_caption_texts(output_document)
    output_media_hashes = media_hashes(output_path)
    output_full_text = "\n".join(normalize(paragraph.text) for paragraph in output_document.paragraphs)

    gap_removed = GAP_HEADING.lower() not in output_full_text.lower()
    reference_removed = REFERENCE_TABLE_HEADING.lower() not in output_full_text.lower()
    results_preserved = source_results == output_results
    captions_preserved = source_captions == output_captions and len(output_captions) == 8
    images_preserved = source_media_hashes == output_media_hashes and len(output_media_hashes) == 8
    source_unchanged = source_hash_before == source_hash_after

    # Internal companion and machine-readable extracts.
    gap_frame = extract_gap_records(gap_texts)
    gap_frame.to_csv(
        outtables / "UTI_HostOmics_U27B3E1_reference_gap_register.tsv",
        sep="\t",
        index=False,
    )
    reference_frame.to_csv(
        outtables / "UTI_HostOmics_U27B3E1_Zotero_reference_table.tsv",
        sep="\t",
        index=False,
    )

    companion_path = internal_dir / INTERNAL_COMPANION_FILENAME
    build_internal_companion(
        companion_path,
        source,
        source_hash_before,
        gap_texts,
        reference_frame,
    )

    supplementary_manifest.to_csv(
        outtables / "UTI_HostOmics_U27B3E1_supplementary_table_manifest.tsv",
        sep="\t",
        index=False,
    )

    candidate_registry = build_candidate_registry(
        project,
        supplementary_manifest,
        [outdocx, outtables, outmetadata, outresults],
    )
    candidate_registry.to_csv(
        outtables / "UTI_HostOmics_U27B3E1_supplementary_source_candidate_registry.tsv",
        sep="\t",
        index=False,
    )

    top_candidate_summary = (
        candidate_registry[candidate_registry["candidate_rank"] == 1]
        .copy()
        .sort_values("supplementary_table")
    )
    top_candidate_summary.to_csv(
        outtables / "UTI_HostOmics_U27B3E1_supplementary_top_candidate_summary.tsv",
        sep="\t",
        index=False,
    )

    # Submission blockers after internal material is separated.
    blocker_rows = []
    for blocker_id, phrase, required_action in SUBMISSION_BLOCKERS:
        present = phrase.lower() in output_full_text.lower()
        blocker_rows.append(
            {
                "blocker_id": blocker_id,
                "trigger_phrase": phrase,
                "present": present,
                "status": "OPEN" if present else "RESOLVED",
                "required_action": required_action,
            }
        )
    blocker_frame = pd.DataFrame(blocker_rows)
    blocker_frame.to_csv(
        outtables / "UTI_HostOmics_U27B3E1_submission_finalization_checklist.tsv",
        sep="\t",
        index=False,
    )

    separation_audit = pd.DataFrame(
        [
            {
                "source_path": str(source),
                "source_sha256_before": source_hash_before,
                "source_sha256_after": source_hash_after,
                "source_unchanged": source_unchanged,
                "output_path": str(output_path),
                "output_sha256": sha256(output_path),
                "reference_gap_section_removed": gap_removed,
                "Zotero_reference_table_removed": reference_removed,
                "internal_companion_created": companion_path.exists(),
                "reference_gap_records_extracted": len(gap_frame),
                "reference_table_rows_extracted": len(reference_frame),
                "results_preserved": results_preserved,
                "figures_and_legends_preserved": captions_preserved,
                "embedded_images_preserved": images_preserved,
                "supplementary_manifest_entries": len(supplementary_manifest),
                "open_submission_blockers": int(blocker_frame["present"].sum()),
            }
        ]
    )
    separation_audit.to_csv(
        outtables / "UTI_HostOmics_U27B3E1_separation_preservation_audit.tsv",
        sep="\t",
        index=False,
    )

    render_info = render_docx(output_path, render_dir)
    pd.DataFrame([render_info]).to_csv(
        outtables / "UTI_HostOmics_U27B3E1_render_audit.tsv",
        sep="\t",
        index=False,
    )

    structural_pass = bool(
        source_unchanged
        and gap_removed
        and reference_removed
        and companion_path.exists()
        and results_preserved
        and captions_preserved
        and images_preserved
        and len(supplementary_manifest) == 10
        and len(gap_frame) >= 5
        and len(reference_frame) >= 1
    )

    if structural_pass and render_info["render_pass"]:
        decision = (
            "READY_FOR_U27B3E2_REFERENCE_FRONT_MATTER_"
            "AND_SUPPLEMENT_SOURCE_CONFIRMATION"
        )
    elif structural_pass:
        decision = "SUBMISSION_ARCHITECTURE_COMPLETE_RENDER_QA_PENDING"
    else:
        decision = "TARGETED_U27B3E1_SEPARATION_OR_PRESERVATION_REPAIR_REQUIRED"

    pd.DataFrame(
        [
            {
                "phase": "U27B3E1",
                "decision": decision,
                "source_unchanged": source_unchanged,
                "internal_reference_material_removed_from_main_manuscript": (
                    gap_removed and reference_removed
                ),
                "internal_reference_companion_created": companion_path.exists(),
                "results_preserved": results_preserved,
                "figures_1_to_8_and_legends_preserved": captions_preserved,
                "embedded_images_preserved": images_preserved,
                "supplementary_tables_manifested": len(supplementary_manifest),
                "supplementary_candidate_rows": len(candidate_registry),
                "open_submission_blockers": int(blocker_frame["present"].sum()),
                "render_pass": render_info["render_pass"],
                "render_pages": render_info["page_count"],
                "manuscript_source_modified": False,
                "scientific_values_recalculated": False,
                "figure_assets_modified": False,
                "source_locks_changed": False,
                "next_phase": (
                    "U27B3E2 confirm authors/affiliations and administrative statements, "
                    "resolve reference gaps/Zotero records, and confirm one source mapping "
                    "for each Supplementary Table S1-S10"
                    if decision.startswith("READY_FOR_U27B3E2")
                    else "Inspect U27B3E1 structural and render audits"
                ),
            }
        ]
    ).to_csv(
        outtables / "UTI_HostOmics_U27B3E1_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    pd.DataFrame(
        [
            {"field": "source_manuscript", "value": str(source)},
            {"field": "source_sha256", "value": source_hash_before},
            {"field": "cleaned_derivative", "value": str(output_path)},
            {"field": "cleaned_derivative_sha256", "value": sha256(output_path)},
            {"field": "internal_reference_companion", "value": str(companion_path)},
            {
                "field": "supplementary_source_candidate_registry",
                "value": str(
                    outtables
                    / "UTI_HostOmics_U27B3E1_supplementary_source_candidate_registry.tsv"
                ),
            },
        ]
    ).to_csv(
        outmetadata / "UTI_HostOmics_U27B3E1_architecture_record.tsv",
        sep="\t",
        index=False,
    )

    report_path = (
        outresults / "UTI_HostOmics_U27B3E1_submission_architecture_report.md"
    )
    with report_path.open("w", encoding="utf-8") as handle:
        handle.write("# Phase U27B3E1 - Reference, supplement and submission architecture\n\n")
        handle.write(f"- Version: `{VERSION}`\n")
        handle.write(f"- Decision: **{decision}**\n")
        handle.write(f"- Cleaned manuscript derivative: `{output_path}`\n")
        handle.write(f"- Internal reference companion: `{companion_path}`\n")
        handle.write(f"- Reference-gap records extracted: **{len(gap_frame)}**.\n")
        handle.write(f"- Zotero table rows extracted: **{len(reference_frame)}**.\n")
        handle.write(f"- Supplementary tables manifested: **{len(supplementary_manifest)}/10**.\n")
        handle.write(f"- Candidate source rows: **{len(candidate_registry)}**.\n")
        handle.write(f"- Open submission blockers: **{int(blocker_frame['present'].sum())}**.\n")
        handle.write(f"- Results preserved: **{results_preserved}**.\n")
        handle.write(f"- Figures/legends preserved: **{captions_preserved}**.\n")
        handle.write(f"- Embedded images preserved: **{images_preserved}**.\n")
        handle.write(f"- Render pass: **{render_info['render_pass']}**.\n")
        handle.write(f"- Render pages: **{render_info['page_count']}**.\n\n")
        handle.write("## Architecture decision\n\n")
        handle.write(
            "The internal reference-gap register and Zotero working table were removed "
            "from the main manuscript and preserved in a separate internal companion. "
            "The cleaned derivative remains a pre-submission architecture file because "
            "authors/affiliations, repository language, contributions, competing interests "
            "and acknowledgements require confirmation. Supplementary Tables S1-S10 are "
            "manifested and linked to ranked project-source candidates, but source selection "
            "and table materialization are deferred to U27B3E2/U27B3E3.\n"
        )
        handle.write("\n## Integrity boundary\n\n")
        handle.write(
            "The scientifically harmonized source was not overwritten. Results, Figures 1-8, "
            "definitive legends and embedded image bytes were preserved. No scientific value "
            "was recalculated and no source lock was changed.\n"
        )

    manifest = {
        "version": VERSION,
        "decision": decision,
        "source": str(source),
        "source_sha256": source_hash_before,
        "output": str(output_path),
        "output_sha256": sha256(output_path),
        "internal_companion": str(companion_path),
        "reference_gap_records": len(gap_frame),
        "reference_table_rows": len(reference_frame),
        "supplementary_tables": len(supplementary_manifest),
        "candidate_rows": len(candidate_registry),
        "open_submission_blockers": int(blocker_frame["present"].sum()),
        "results_preserved": results_preserved,
        "figures_legends_preserved": captions_preserved,
        "images_preserved": images_preserved,
        "render_pass": render_info["render_pass"],
        "render_pages": render_info["page_count"],
        "source_modified": False,
        "scientific_values_recalculated": False,
        "figure_assets_modified": False,
        "source_locks_changed": False,
    }
    (
        outresults / "UTI_HostOmics_U27B3E1_run_manifest.json"
    ).write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    log(f"Reference-gap records extracted: {len(gap_frame)}")
    log(f"Zotero reference rows extracted: {len(reference_frame)}")
    log(f"Supplementary tables manifested: {len(supplementary_manifest)}/10")
    log(f"Supplementary source candidates: {len(candidate_registry)}")
    log(f"Open submission blockers: {int(blocker_frame['present'].sum())}")
    log(f"Results preserved: {results_preserved}")
    log(f"Figures and legends preserved: {captions_preserved}")
    log(f"Embedded images preserved: {images_preserved}")
    log(f"Render pass: {render_info['render_pass']}")
    log(f"Decision: {decision}")
    log(f"Output: {output_path}")
    log(f"Report: {report_path}")

    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3E1] ERROR: {exc}", file=sys.stderr)
        raise
