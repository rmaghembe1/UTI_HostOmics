#!/usr/bin/env python3
"""
Phase U27B3C4.1
Clean the integrated Figure 1-8 section in the v5.1 manuscript derivative.

Repairs
-------
1. Remove the six obsolete legacy figure-title/caption pairs that remained
   after Acknowledgements.
2. Remove the orphaned "Figures" and "Figure 1" heading paragraphs that
   created a nearly blank page before the new figure block.
3. Retain the eight frozen figure images and definitive U27B3B legends.
4. Reformat definitive legends as black manuscript body text, with only the
   figure-number/title prefix bold.
5. Update derivative-version metadata in the title line, header and footer.
6. Preserve the reconstructed Results section, Discussion onward text,
   tables, embedded frozen figure bytes and the v5.1 source derivative.
7. Render the v5.2 derivative and create a page contact sheet.

This phase does not harmonize the legacy Abstract, Methods, Discussion,
Data availability or Code availability text. Those sections are audited and
carried forward to a later manuscript-harmonization phase.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple
from xml.etree import ElementTree as ET

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image, ImageDraw

VERSION = "U27B3C41_v1.0_2026-07-16"
TAG = "phaseU27B3C41_integrated_figure_section_cleanup"

DEFAULT_SOURCE = (
    "__UTI_HOSTOMICS_PROJECT_ROOT__/"
    "09_manuscript_docx/phaseU27B3C4_frozen_figures_legends_integration/"
    "UTI_HostOmics_preZotero_manuscript_v5_1_"
    "U27B3C4_figures_legends_integrated.docx"
)

OUTPUT_FILENAME = (
    "UTI_HostOmics_preZotero_manuscript_v5_2_"
    "U27B3C41_figure_section_cleaned.docx"
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W_NS, "a": A_NS, "r": R_NS}

LEGACY_FIGURE_TITLES = (
    "Figure 1. Multi-context host-omics framework for reconstructing urinary tract infection host-response architecture.",
    "Figure 2. Integrated cross-cohort module atlas.",
    "Figure 3. Cross-cohort atlas with single-cell UPEC validation.",
    "Figure 4. Prioritized single-cell-validated UTI host-response modules.",
    "Figure 5. Pregnancy-associated UTI module architecture.",
    "Figure 6. Single-cell UPEC validation layer.",
)

DEFINITIVE_FIGURE_TITLES = {
    1: "Study architecture, datasets, contrasts and evidence hierarchy",
    2: "Cross-dataset infection effects, recurrent cores and contextual comparators",
    3: "Pregnancy, tissue and outcome-associated endocrine-metabolic-complement remodeling",
    4: "Single-cell composition, immune states and cellular localization",
    5: "Steroid, cholesterol and lipid-remodeling architecture",
    6: "Adipokine, insulin and integrated immunometabolic remodeling",
    7: "Complement branch and cellular architecture",
    8: "Integrated endocrine-metabolic-immune model and evidence boundaries",
}


def log(message: str) -> None:
    print(f"[U27B3C4.1] {message}", flush=True)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            block = handle.read(1024 * 1024)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def paragraph_has_drawing(paragraph) -> bool:
    return bool(
        paragraph._p.findall(f".//{{{W_NS}}}drawing")
        or paragraph._p.findall(f".//{{{W_NS}}}pict")
    )


def clear_paragraph_content(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def style_definitive_caption(paragraph, figure_number: int) -> None:
    text = normalize(paragraph.text)
    marker = " (A)"
    marker_index = text.find(marker)
    if marker_index < 0:
        raise RuntimeError(
            f"Figure {figure_number} definitive legend lacks the (A) marker."
        )

    prefix = text[:marker_index].strip()
    remainder = text[marker_index + 1 :].strip()

    expected_prefix = (
        f"Figure {figure_number}. "
        f"{DEFINITIVE_FIGURE_TITLES[figure_number]}."
    )
    if prefix != expected_prefix:
        raise RuntimeError(
            f"Unexpected Figure {figure_number} legend prefix:\n"
            f"observed={prefix!r}\nexpected={expected_prefix!r}"
        )

    clear_paragraph_content(paragraph)
    paragraph.style = paragraph._parent.part.document.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.widow_control = True

    prefix_run = paragraph.add_run(prefix + " ")
    prefix_run.bold = True
    prefix_run.font.size = Pt(8.5)
    prefix_run.font.color.rgb = RGBColor(0, 0, 0)

    body_run = paragraph.add_run(remainder)
    body_run.bold = False
    body_run.font.size = Pt(8.5)
    body_run.font.color.rgb = RGBColor(0, 0, 0)


def section_text(document: Document, heading: str, next_heading: str) -> List[str]:
    paragraphs = [normalize(paragraph.text) for paragraph in document.paragraphs]

    def find_heading(value: str) -> int:
        matches = []
        for index, text in enumerate(paragraphs):
            cleaned = re.sub(
                r"^\d+(?:\.\d+)*\s*", "", text.lower()
            ).rstrip(":.")
            if cleaned == value.lower():
                matches.append(index)
        if len(matches) != 1:
            raise RuntimeError(
                f"Expected one {value!r} heading; observed {matches}."
            )
        return matches[0]

    start = find_heading(heading)
    end = find_heading(next_heading)
    if start >= end:
        raise RuntimeError(f"{heading} does not precede {next_heading}.")
    return paragraphs[start + 1 : end]


def table_text(document: Document) -> List[str]:
    values = []
    for table in document.tables:
        values.append(
            normalize(
                " ".join(
                    cell.text
                    for row in table.rows
                    for cell in row.cells
                )
            )
        )
    return values


def read_docx_structure(path: Path) -> Dict[str, object]:
    with zipfile.ZipFile(path) as archive:
        document_root = ET.fromstring(archive.read("word/document.xml"))
        rel_root = ET.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        rel_map = {
            relationship.attrib.get("Id", ""): relationship.attrib.get(
                "Target", ""
            )
            for relationship in rel_root
        }

        body = document_root.find(f"{{{W_NS}}}body")
        if body is None:
            raise RuntimeError("No word/body element in DOCX.")

        paragraph_texts: List[str] = []
        drawing_hashes: List[str] = []
        definitive_captions: List[str] = []
        legacy_captions: List[str] = []

        for element in list(body):
            if element.tag != f"{{{W_NS}}}p":
                continue
            text = normalize(
                "".join(
                    node.text or ""
                    for node in element.findall(".//w:t", NS)
                )
            )
            paragraph_texts.append(text)

            if any(text.startswith(title) for title in LEGACY_FIGURE_TITLES):
                legacy_captions.append(text)

            if re.match(r"^Figure\s+[1-8]\.\s+", text) and "(A)" in text:
                definitive_captions.append(text)

            for blip in element.findall(".//a:blip", NS):
                relationship_id = blip.attrib.get(f"{{{R_NS}}}embed", "")
                target = rel_map.get(relationship_id, "")
                archive_path = (
                    f"word/{target}"
                    if target and not target.startswith("/")
                    else target.lstrip("/")
                )
                if archive_path in archive.namelist():
                    drawing_hashes.append(
                        hashlib.sha256(archive.read(archive_path)).hexdigest()
                    )

        media_files = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]

    return {
        "paragraph_texts": paragraph_texts,
        "drawing_hashes": drawing_hashes,
        "definitive_captions": definitive_captions,
        "legacy_captions": legacy_captions,
        "media_files": media_files,
    }


def render_docx(docx_path: Path, render_dir: Path) -> Dict[str, object]:
    render_dir.mkdir(parents=True, exist_ok=True)
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    pdftoppm = shutil.which("pdftoppm")

    if not libreoffice:
        return {
            "render_attempted": False,
            "render_pass": False,
            "reason": "LibreOffice/soffice not found",
            "pdf_path": "",
            "page_count": 0,
            "contact_sheet": "",
        }
    if not pdftoppm:
        return {
            "render_attempted": False,
            "render_pass": False,
            "reason": "pdftoppm not found",
            "pdf_path": "",
            "page_count": 0,
            "contact_sheet": "",
        }

    with tempfile.TemporaryDirectory(prefix="u27b3c41_lo_") as tmp:
        env = os.environ.copy()
        env["HOME"] = tmp
        profile_uri = Path(tmp).resolve().as_uri()
        result = subprocess.run(
            [
                libreoffice,
                "--headless",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(render_dir),
                str(docx_path),
            ],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            env=env,
            check=False,
        )

    pdf_path = render_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not pdf_path.exists():
        return {
            "render_attempted": True,
            "render_pass": False,
            "reason": result.stderr.strip() or result.stdout.strip(),
            "pdf_path": str(pdf_path),
            "page_count": 0,
            "contact_sheet": "",
        }

    page_prefix = render_dir / "page"
    raster = subprocess.run(
        [
            pdftoppm,
            "-png",
            "-r",
            "140",
            str(pdf_path),
            str(page_prefix),
        ],
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )

    pages = sorted(
        render_dir.glob("page-*.png"),
        key=lambda path: int(re.search(r"(\d+)$", path.stem).group(1)),
    )
    if raster.returncode != 0 or not pages:
        return {
            "render_attempted": True,
            "render_pass": False,
            "reason": raster.stderr.strip() or raster.stdout.strip(),
            "pdf_path": str(pdf_path),
            "page_count": len(pages),
            "contact_sheet": "",
        }

    contact_sheet = render_dir / "UTI_HostOmics_U27B3C41_render_contact_sheet.png"
    make_contact_sheet(pages, contact_sheet)

    return {
        "render_attempted": True,
        "render_pass": bool(
            pdf_path.stat().st_size > 0
            and contact_sheet.exists()
            and contact_sheet.stat().st_size > 0
        ),
        "reason": "Render files created",
        "pdf_path": str(pdf_path),
        "page_count": len(pages),
        "contact_sheet": str(contact_sheet),
        "libreoffice_stdout": result.stdout.strip(),
        "libreoffice_stderr": result.stderr.strip(),
    }


def make_contact_sheet(
    paths: Sequence[Path],
    output: Path,
    columns: int = 3,
    cell_width: int = 600,
    padding: int = 24,
) -> None:
    images: List[Image.Image] = []
    for path in paths:
        image = Image.open(path).convert("RGB")
        ratio = cell_width / image.width
        resized = image.resize(
            (cell_width, max(1, int(image.height * ratio)))
        )
        canvas = Image.new("RGB", (cell_width, resized.height + 22), "white")
        canvas.paste(resized, (0, 22))
        draw = ImageDraw.Draw(canvas)
        draw.text((4, 3), path.stem, fill="black")
        images.append(canvas)

    rows = math.ceil(len(images) / columns)
    row_heights = []
    for row_index in range(rows):
        subset = images[row_index * columns : (row_index + 1) * columns]
        row_heights.append(max(image.height for image in subset))

    width = columns * cell_width + (columns + 1) * padding
    height = sum(row_heights) + (rows + 1) * padding
    sheet = Image.new("RGB", (width, height), "white")

    y = padding
    for row_index in range(rows):
        x = padding
        subset = images[row_index * columns : (row_index + 1) * columns]
        for image in subset:
            sheet.paste(image, (x, y))
            x += cell_width + padding
        y += row_heights[row_index] + padding

    output.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root",
        default="__UTI_HOSTOMICS_PROJECT_ROOT__",
    )
    parser.add_argument("--source", default=DEFAULT_SOURCE)
    args = parser.parse_args()

    project = Path(args.project_root).resolve()
    source = Path(args.source).resolve()
    if not source.exists():
        raise FileNotFoundError(f"Integrated v5.1 derivative not found: {source}")

    outdocx = project / "09_manuscript_docx" / TAG
    outtables = project / "06_tables" / TAG
    outmetadata = project / "03_metadata" / TAG
    outresults = project / "05_results" / TAG
    render_dir = outdocx / "render_qa"

    for directory in (outdocx, outtables, outmetadata, outresults, render_dir):
        directory.mkdir(parents=True, exist_ok=True)

    output_path = outdocx / OUTPUT_FILENAME
    source_hash_before = sha256(source)

    source_document = Document(source)
    source_results = section_text(source_document, "Results", "Discussion")
    source_tables = table_text(source_document)
    source_structure = read_docx_structure(source)

    document = Document(source)

    # Update visible derivative metadata without changing scientific sections.
    if document.paragraphs and document.paragraphs[0].text.startswith("Pre-Zotero manuscript"):
        document.paragraphs[0].text = (
            "Pre-Zotero manuscript v5.2 | Results reconstructed; "
            "frozen Figures 1-8 integrated"
        )
    if len(document.paragraphs) > 1 and document.paragraphs[1].text.startswith("Generated:"):
        document.paragraphs[1].text = "Updated: 2026-07-16"

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if "UTI HostOmics Project" in paragraph.text:
                paragraph.text = "UTI HostOmics Project - Draft manuscript v5.2"
        for paragraph in section.footer.paragraphs:
            if paragraph.text:
                paragraph.text = "Draft updated 2026-07-16 | Pre-Zotero working draft"

    # Remove legacy captions and old title-only paragraphs.
    removed_legacy: List[str] = []
    for paragraph in list(document.paragraphs):
        text = normalize(paragraph.text)
        if any(text.startswith(title) for title in LEGACY_FIGURE_TITLES):
            removed_legacy.append(text)
            remove_paragraph(paragraph)

    # Remove orphaned Figure-section headings that produced a blank page.
    removed_orphans: List[str] = []
    for paragraph in list(document.paragraphs):
        text = normalize(paragraph.text)
        if text in {"Figures", "Figure 1"}:
            removed_orphans.append(text)
            remove_paragraph(paragraph)

    # Normalize the definitive captions and image paragraphs.
    definitive_captions: Dict[int, object] = {}
    drawing_paragraphs: List[object] = []

    for paragraph in document.paragraphs:
        if paragraph_has_drawing(paragraph):
            drawing_paragraphs.append(paragraph)
            paragraph.style = document.styles["Normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(4)
            continue

        text = normalize(paragraph.text)
        match = re.match(r"^Figure\s+([1-8])\.\s+", text)
        if match and "(A)" in text:
            figure_number = int(match.group(1))
            definitive_captions[figure_number] = paragraph

    if len(drawing_paragraphs) != 8:
        raise RuntimeError(
            f"Expected eight frozen drawing paragraphs; observed {len(drawing_paragraphs)}."
        )
    if sorted(definitive_captions) != list(range(1, 9)):
        raise RuntimeError(
            "Definitive caption set is incomplete: "
            f"{sorted(definitive_captions)}"
        )

    for figure_number in range(1, 9):
        style_definitive_caption(
            definitive_captions[figure_number],
            figure_number,
        )

    # The paragraph after Figure 8 remains a page break before the tables.
    if output_path.exists():
        output_path.unlink()
    document.save(output_path)

    source_hash_after = sha256(source)
    output_hash = sha256(output_path)
    output_document = Document(output_path)
    output_results = section_text(output_document, "Results", "Discussion")
    output_tables = table_text(output_document)
    output_structure = read_docx_structure(output_path)

    source_unchanged = source_hash_before == source_hash_after
    results_preserved = source_results == output_results
    tables_preserved = source_tables == output_tables
    figure_hashes_preserved = (
        source_structure["drawing_hashes"]
        == output_structure["drawing_hashes"]
        and len(output_structure["drawing_hashes"]) == 8
    )

    legacy_absent = len(output_structure["legacy_captions"]) == 0
    definitive_count = len(output_structure["definitive_captions"])
    orphan_heading_absent = all(
        text not in {"Figures", "Figure 1"}
        for text in output_structure["paragraph_texts"]
    )

    pd.DataFrame(
        [
            {
                "legacy_caption_paragraphs_detected_in_source": len(
                    source_structure["legacy_captions"]
                ),
                "legacy_caption_paragraphs_removed": len(removed_legacy),
                "orphan_heading_paragraphs_removed": "; ".join(removed_orphans),
                "legacy_caption_paragraphs_remaining": len(
                    output_structure["legacy_captions"]
                ),
                "definitive_caption_paragraphs": definitive_count,
                "orphan_figure_heading_absent": orphan_heading_absent,
            }
        ]
    ).to_csv(
        outtables / "UTI_HostOmics_U27B3C41_legacy_caption_cleanup_audit.tsv",
        sep="\t",
        index=False,
    )

    caption_rows = []
    for figure_number in range(1, 9):
        paragraph = definitive_captions[figure_number]
        caption_rows.append(
            {
                "figure_number": figure_number,
                "style": paragraph.style.name,
                "alignment": str(paragraph.alignment),
                "run_count": len(paragraph.runs),
                "prefix_bold": bool(paragraph.runs[0].bold),
                "body_bold": bool(paragraph.runs[1].bold),
                "prefix_color": str(paragraph.runs[0].font.color.rgb),
                "body_color": str(paragraph.runs[1].font.color.rgb),
                "caption_character_count": len(paragraph.text),
            }
        )
    pd.DataFrame(caption_rows).to_csv(
        outtables / "UTI_HostOmics_U27B3C41_caption_format_audit.tsv",
        sep="\t",
        index=False,
    )

    preservation = pd.DataFrame(
        [
            {
                "source_path": str(source),
                "source_sha256_before": source_hash_before,
                "source_sha256_after": source_hash_after,
                "source_unchanged": source_unchanged,
                "output_path": str(output_path),
                "output_sha256": output_hash,
                "results_text_preserved": results_preserved,
                "tables_preserved": tables_preserved,
                "drawing_count_source": len(source_structure["drawing_hashes"]),
                "drawing_count_output": len(output_structure["drawing_hashes"]),
                "media_parts_output": len(output_structure["media_files"]),
                "embedded_figure_hashes_preserved": figure_hashes_preserved,
                "legacy_captions_absent": legacy_absent,
                "definitive_captions_present": definitive_count == 8,
            }
        ]
    )
    preservation.to_csv(
        outtables / "UTI_HostOmics_U27B3C41_docx_preservation_audit.tsv",
        sep="\t",
        index=False,
    )

    # Known nonblocking manuscript-wide harmonization issues.
    full_text = "\n".join(normalize(p.text) for p in output_document.paragraphs)
    non_results_text = "\n".join(
        normalize(p.text)
        for p in output_document.paragraphs
        if normalize(p.text) not in set(output_results)
    )
    legacy_section_audit = pd.DataFrame(
        [
            {
                "issue_id": "LEGACY_GSE186800_OUTSIDE_RESULTS",
                "present": "GSE186800" in non_results_text,
                "required_action": (
                    "Harmonize Abstract, Methods, Discussion and Data availability "
                    "to the frozen GSE168600 architecture."
                ),
                "blocking_for_figure_section_cleanup": False,
            },
            {
                "issue_id": "LEGACY_ABSTRACT_MODULE_FRAMEWORK",
                "present": "Cross-cohort integration identified 17 modules" in full_text,
                "required_action": (
                    "Rewrite the Abstract against the frozen 78-submodule and "
                    "eight-figure synthesis."
                ),
                "blocking_for_figure_section_cleanup": False,
            },
            {
                "issue_id": "LEGACY_CODE_DIRECTORY_REFERENCE",
                "present": "stored under `04_scripts/`" in full_text,
                "required_action": (
                    "Update Code availability to the current 10_scripts and package paths."
                ),
                "blocking_for_figure_section_cleanup": False,
            },
        ]
    )
    legacy_section_audit.to_csv(
        outtables / "UTI_HostOmics_U27B3C41_nonblocking_section_harmonization_audit.tsv",
        sep="\t",
        index=False,
    )

    render_info = render_docx(output_path, render_dir)
    pd.DataFrame([render_info]).to_csv(
        outtables / "UTI_HostOmics_U27B3C41_render_audit.tsv",
        sep="\t",
        index=False,
    )

    structural_pass = bool(
        source_unchanged
        and results_preserved
        and tables_preserved
        and figure_hashes_preserved
        and legacy_absent
        and definitive_count == 8
        and orphan_heading_absent
        and len(output_structure["media_files"]) == 8
    )

    if structural_pass and render_info["render_pass"]:
        decision = "READY_FOR_U27B3C5_FIGURE_SECTION_VISUAL_AUDIT"
    elif structural_pass:
        decision = "FIGURE_SECTION_CLEANUP_COMPLETE_RENDER_QA_PENDING"
    else:
        decision = "TARGETED_U27B3C41_INTEGRATION_CLEANUP_REPAIR_REQUIRED"

    pd.DataFrame(
        [
            {
                "phase": "U27B3C4.1",
                "decision": decision,
                "legacy_caption_paragraphs_removed": len(removed_legacy),
                "legacy_caption_paragraphs_remaining": len(
                    output_structure["legacy_captions"]
                ),
                "orphan_headings_removed": len(removed_orphans),
                "frozen_figures_preserved": len(output_structure["drawing_hashes"]),
                "definitive_legends_preserved": definitive_count,
                "embedded_figure_hashes_preserved": figure_hashes_preserved,
                "results_text_preserved": results_preserved,
                "tables_preserved": tables_preserved,
                "source_derivative_unchanged": source_unchanged,
                "render_pass": render_info["render_pass"],
                "known_nonblocking_section_harmonization_issues": int(
                    legacy_section_audit["present"].sum()
                ),
                "scientific_values_recalculated": False,
                "frozen_figure_assets_modified": False,
                "source_locks_changed": False,
                "source_derivative_modified": False,
                "new_derivative_created": True,
                "next_phase": (
                    "U27B3C5 visually inspect the cleaned Figure 1-8 section; "
                    "then U27B3D harmonize Abstract, Methods, Discussion and "
                    "availability statements."
                    if decision.startswith("READY_FOR_U27B3C5")
                    else "Inspect cleanup and render audits."
                ),
            }
        ]
    ).to_csv(
        outtables / "UTI_HostOmics_U27B3C41_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    pd.DataFrame(
        [
            {"field": "source_path", "value": str(source)},
            {"field": "source_sha256", "value": source_hash_before},
            {"field": "cleaned_derivative_path", "value": str(output_path)},
            {"field": "cleaned_derivative_sha256", "value": output_hash},
            {"field": "source_overwritten", "value": "False"},
        ]
    ).to_csv(
        outmetadata / "UTI_HostOmics_U27B3C41_derivative_record.tsv",
        sep="\t",
        index=False,
    )

    report_path = (
        outresults / "UTI_HostOmics_U27B3C41_figure_section_cleanup_report.md"
    )
    with report_path.open("w", encoding="utf-8") as handle:
        handle.write(
            "# Phase U27B3C4.1 - Integrated figure-section cleanup\n\n"
        )
        handle.write(f"- Version: `{VERSION}`\n")
        handle.write(f"- Decision: **{decision}**\n")
        handle.write(f"- Source derivative: `{source}`\n")
        handle.write(f"- New derivative: `{output_path}`\n")
        handle.write(
            f"- Legacy caption paragraphs removed: **{len(removed_legacy)}**.\n"
        )
        handle.write(
            f"- Orphaned heading paragraphs removed: **{len(removed_orphans)}**.\n"
        )
        handle.write(
            f"- Frozen figures preserved: **{len(output_structure['drawing_hashes'])}/8**.\n"
        )
        handle.write(
            f"- Definitive legends preserved: **{definitive_count}/8**.\n"
        )
        handle.write(
            f"- Embedded figure hashes preserved: **{figure_hashes_preserved}**.\n"
        )
        handle.write(f"- Results text preserved: **{results_preserved}**.\n")
        handle.write(f"- Tables preserved: **{tables_preserved}**.\n")
        handle.write(f"- Render pass: **{render_info['render_pass']}**.\n")
        handle.write(
            f"- Render contact sheet: `{render_info.get('contact_sheet', '')}`.\n\n"
        )
        handle.write("## Corrected defects\n\n")
        handle.write(
            "The six legacy figure-title/caption pairs and the orphaned "
            "`Figures`/`Figure 1` headings were removed. The eight definitive "
            "legends were restyled as black manuscript text with bold figure "
            "prefixes, while the embedded frozen figure bytes were retained "
            "unchanged.\n\n"
        )
        handle.write("## Remaining nonblocking manuscript work\n\n")
        handle.write(
            "The legacy Abstract, Methods, Discussion, Data availability and "
            "Code availability sections remain outside the scope of this "
            "figure-section cleanup. Their outdated GSE186800 and pre-U27 "
            "language is carried forward for a dedicated manuscript-wide "
            "harmonization phase after visual approval of this derivative.\n"
        )

    manifest = {
        "version": VERSION,
        "decision": decision,
        "source_path": str(source),
        "source_sha256": source_hash_before,
        "output_path": str(output_path),
        "output_sha256": output_hash,
        "legacy_caption_paragraphs_removed": len(removed_legacy),
        "orphan_headings_removed": len(removed_orphans),
        "frozen_figures_preserved": len(output_structure["drawing_hashes"]),
        "definitive_legends_preserved": definitive_count,
        "embedded_figure_hashes_preserved": figure_hashes_preserved,
        "results_text_preserved": results_preserved,
        "tables_preserved": tables_preserved,
        "source_unchanged": source_unchanged,
        "render_pass": render_info["render_pass"],
        "contact_sheet": render_info.get("contact_sheet", ""),
        "known_nonblocking_section_harmonization_issues": int(
            legacy_section_audit["present"].sum()
        ),
    }
    (
        outresults / "UTI_HostOmics_U27B3C41_run_manifest.json"
    ).write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    log(f"Legacy captions removed: {len(removed_legacy)}")
    log(f"Orphan headings removed: {len(removed_orphans)}")
    log(f"Frozen figures preserved: {len(output_structure['drawing_hashes'])}/8")
    log(f"Definitive legends preserved: {definitive_count}/8")
    log(f"Results preserved: {results_preserved}")
    log(f"Tables preserved: {tables_preserved}")
    log(f"Render pass: {render_info['render_pass']}")
    log(f"Decision: {decision}")
    log(f"Output: {output_path}")
    log(f"Report: {report_path}")

    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3C4.1] ERROR: {exc}", file=sys.stderr)
        raise
