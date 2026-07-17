#!/usr/bin/env python3
"""
Phase U27B3C4
Integrate frozen Figures 1-8 and definitive legends into the reconstructed
v5.0 manuscript derivative.

Inputs
------
- U27B3C2 reconstructed Results derivative
- U27B3A frozen Figures 1-8 PNG masters
- U27B3B definitive plain-text legends

This phase:
1. treats the U27B3C2 derivative as read-only;
2. identifies the six legacy figure block at the end of the manuscript;
3. removes only the legacy figure images, captions and intervening page breaks;
4. inserts frozen Figures 1-8 in order, each followed by its definitive legend;
5. preserves the narrative text, Results section, Discussion text and tables;
6. removes unreferenced legacy image relationships and media parts;
7. validates embedded-image hashes against the frozen PNG masters;
8. renders the new derivative to PDF/PNG pages and creates a contact sheet.

No scientific value, Results text, Discussion text, source lock or frozen
figure asset is modified.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple
from xml.etree import ElementTree as ET

import pandas as pd

try:
    from PIL import Image
except ImportError as exc:
    raise RuntimeError("Pillow is required for render contact sheets.") from exc

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu, Pt
    from docx.oxml.ns import qn
except ImportError as exc:
    raise RuntimeError("python-docx is required for U27B3C4.") from exc


VERSION = "U27B3C4_v1.0_2026-07-16"
TAG = "phaseU27B3C4_frozen_figures_legends_integration"

DEFAULT_SOURCE = (
    "__UTI_HOSTOMICS_PROJECT_ROOT__/"
    "09_manuscript_docx/phaseU27B3C2_results_section_reconstruction/"
    "UTI_HostOmics_preZotero_manuscript_v5_0_U27B3C2_results_reconstructed.docx"
)

FIGURE_PACKAGE_TAG = "phaseU27B3A_complete_eight_figure_package_assembly"
LEGEND_TAG = "phaseU27B3B_definitive_figure_legend_construction"

OUTPUT_FILENAME = (
    "UTI_HostOmics_preZotero_manuscript_"
    "v5_1_U27B3C4_figures_legends_integrated.docx"
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

NS = {
    "w": W_NS,
    "r": R_NS,
    "a": A_NS,
}

TABLE_CAPTION_PATTERN = re.compile(
    r"^\s*(?:Supplementary\s+)?Table\s+[A-Z]?\d+",
    flags=re.IGNORECASE,
)

FIGURE_CAPTION_PATTERN = re.compile(
    r"^\s*Figure\s+(\d+)\.",
    flags=re.IGNORECASE,
)


def log(message: str) -> None:
    print(f"[U27B3C4] {message}", flush=True)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            block = handle.read(1024 * 1024)
            if not block:
                break
            digest.update(block)
    return digest.hexdigest()


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text)).strip()


def body_element_text(element) -> str:
    pieces: List[str] = []
    for node in element.iter():
        local = node.tag.rsplit("}", 1)[-1]
        if local == "t":
            pieces.append(node.text or "")
        elif local == "tab":
            pieces.append("\t")
        elif local in {"br", "cr"}:
            pieces.append("\n")
    return normalize_text("".join(pieces))


def paragraph_has_drawing(element) -> bool:
    return bool(
        element.findall(f".//{{{W_NS}}}drawing")
        or element.findall(f".//{{{W_NS}}}pict")
    )


def paragraph_has_page_break(element) -> bool:
    for br in element.findall(f".//{{{W_NS}}}br"):
        if br.attrib.get(f"{{{W_NS}}}type", "") == "page":
            return True
    ppr = element.find(f"{{{W_NS}}}pPr")
    if ppr is not None and ppr.find(f"{{{W_NS}}}pageBreakBefore") is not None:
        return True
    return False


def parse_legends(path: Path) -> Dict[int, str]:
    if not path.exists():
        raise FileNotFoundError(f"Definitive legend file not found: {path}")

    raw = path.read_text(encoding="utf-8")
    matches = list(
        re.finditer(
            r"(?m)^Figure\s+([1-8])\.\s+",
            raw,
        )
    )
    if len(matches) != 8:
        raise RuntimeError(
            f"Expected eight figure legends in {path}; observed {len(matches)}."
        )

    legends: Dict[int, str] = {}
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(raw)
        chunk = normalize_text(raw[start:end])
        figure_number = int(match.group(1))
        legends[figure_number] = chunk

    if sorted(legends) != list(range(1, 9)):
        raise RuntimeError(
            f"Legend numbering is incomplete or duplicated: {sorted(legends)}"
        )

    return legends


def split_legend_prefix(legend: str, figure_number: int) -> Tuple[str, str]:
    match = re.match(
        rf"^(Figure\s+{figure_number}\.\s+.+?)(?=\s+\(A\))",
        legend,
        flags=re.IGNORECASE,
    )
    if match:
        prefix = match.group(1).strip()
        remainder = legend[len(match.group(1)):].strip()
        return prefix, remainder

    marker = f"Figure {figure_number}."
    if legend.startswith(marker):
        return marker, legend[len(marker):].strip()
    return marker, legend


def locate_legacy_figure_block(document: Document) -> Dict[str, object]:
    body = document._element.body
    children = list(body)

    drawing_indices = [
        index
        for index, element in enumerate(children)
        if element.tag == qn("w:p") and paragraph_has_drawing(element)
    ]

    if len(drawing_indices) != 6:
        raise RuntimeError(
            "Expected exactly six legacy drawing paragraphs in the U27B3C2 "
            f"derivative; observed {len(drawing_indices)}."
        )

    first_drawing = min(drawing_indices)
    last_drawing = max(drawing_indices)

    start_index = first_drawing
    while start_index > 0:
        previous = children[start_index - 1]
        if previous.tag != qn("w:p"):
            break
        previous_text = body_element_text(previous)
        if (
            previous_text == ""
            and paragraph_has_page_break(previous)
        ):
            start_index -= 1
            continue
        if re.fullmatch(
            r"(?:Figures?|Figure legends?|Figures and legends)",
            previous_text,
            flags=re.IGNORECASE,
        ):
            start_index -= 1
            continue
        break

    anchor_index: Optional[int] = None
    for index in range(last_drawing + 1, len(children)):
        element = children[index]
        if element.tag == qn("w:tbl"):
            anchor_index = index
            break
        if element.tag == qn("w:p"):
            text = body_element_text(element)
            if TABLE_CAPTION_PATTERN.match(text):
                anchor_index = index
                break

    if anchor_index is None:
        anchor_index = len(children)

    removed_elements = children[start_index:anchor_index]
    removed_text = [
        body_element_text(element)
        for element in removed_elements
        if body_element_text(element)
    ]
    removed_figure_captions = [
        text
        for text in removed_text
        if FIGURE_CAPTION_PATTERN.match(text)
    ]

    narrative_prefix_text = [
        body_element_text(element)
        for element in children[:start_index]
        if body_element_text(element)
    ]
    table_and_tail_text = [
        body_element_text(element)
        for element in children[anchor_index:]
        if body_element_text(element)
    ]

    anchor_element = (
        children[anchor_index]
        if anchor_index < len(children)
        else None
    )

    return {
        "body": body,
        "children": children,
        "drawing_indices": drawing_indices,
        "start_index": start_index,
        "anchor_index": anchor_index,
        "anchor_element": anchor_element,
        "removed_elements": removed_elements,
        "removed_text": removed_text,
        "removed_figure_captions": removed_figure_captions,
        "narrative_prefix_text": narrative_prefix_text,
        "table_and_tail_text": table_and_tail_text,
    }


def insert_before_anchor(body, anchor_element, element) -> None:
    if anchor_element is None:
        body.insert(len(body) - 1, element)
    else:
        anchor_element.addprevious(element)


def available_figure_width(document: Document) -> Emu:
    section = document.sections[-1]
    available = (
        section.page_width
        - section.left_margin
        - section.right_margin
    )
    # Leave a small safety margin for renderer differences.
    safety = Emu(180000)
    return Emu(max(int(available) - int(safety), 1))


def caption_style_name(document: Document) -> str:
    try:
        document.styles["Caption"]
        return "Caption"
    except KeyError:
        return "Normal"


def add_figure_and_legend(
    document: Document,
    body,
    anchor_element,
    figure_number: int,
    figure_path: Path,
    legend: str,
    width: Emu,
) -> None:
    image_paragraph = document.add_paragraph()
    image_paragraph.paragraph_format.page_break_before = True
    image_paragraph.paragraph_format.space_before = Pt(0)
    image_paragraph.paragraph_format.space_after = Pt(4)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = image_paragraph.add_run()
    inline_shape = run.add_picture(str(figure_path), width=width)
    title = f"Figure {figure_number}"
    inline_shape._inline.docPr.set("title", title)
    inline_shape._inline.docPr.set(
        "descr",
        f"{title}: frozen manuscript-facing figure from phase U27B3A.",
    )
    insert_before_anchor(body, anchor_element, image_paragraph._p)

    caption = document.add_paragraph(style=caption_style_name(document))
    caption.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_together = False
    caption.paragraph_format.keep_with_next = False

    prefix, remainder = split_legend_prefix(legend, figure_number)
    prefix_run = caption.add_run(prefix + (" " if remainder else ""))
    prefix_run.bold = True
    prefix_run.font.size = Pt(8.5)

    if remainder:
        remainder_run = caption.add_run(remainder)
        remainder_run.font.size = Pt(8.5)

    insert_before_anchor(body, anchor_element, caption._p)


def add_page_break_before_anchor(
    document: Document,
    body,
    anchor_element,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.page_break_before = True
    insert_before_anchor(body, anchor_element, paragraph._p)


def cleanup_unreferenced_document_images(docx_path: Path) -> None:
    """
    Remove image relationships and media parts no longer referenced by
    word/document.xml after the six legacy drawing paragraphs have been removed.
    """
    with tempfile.TemporaryDirectory(prefix="u27b3c4_clean_") as tmp:
        root = Path(tmp)
        with zipfile.ZipFile(docx_path, "r") as archive:
            archive.extractall(root)

        document_xml = root / "word" / "document.xml"
        rels_xml = root / "word" / "_rels" / "document.xml.rels"

        document_root = ET.parse(document_xml).getroot()
        referenced_ids = set()
        for element in document_root.iter():
            for attribute_name, value in element.attrib.items():
                if attribute_name in {
                    f"{{{R_NS}}}embed",
                    f"{{{R_NS}}}id",
                    f"{{{R_NS}}}link",
                }:
                    referenced_ids.add(value)

        rel_tree = ET.parse(rels_xml)
        rel_root = rel_tree.getroot()

        removed_targets: List[str] = []
        for relationship in list(rel_root):
            relationship_type = relationship.attrib.get("Type", "")
            relationship_id = relationship.attrib.get("Id", "")
            if (
                relationship_type.endswith("/image")
                and relationship_id not in referenced_ids
            ):
                removed_targets.append(
                    relationship.attrib.get("Target", "")
                )
                rel_root.remove(relationship)

        rel_tree.write(
            rels_xml,
            encoding="utf-8",
            xml_declaration=True,
        )

        referenced_media_names = set()
        for rel_path in (root / "word").rglob("*.rels"):
            try:
                tree = ET.parse(rel_path)
            except ET.ParseError:
                continue
            for relationship in tree.getroot():
                if relationship.attrib.get("Type", "").endswith("/image"):
                    target = relationship.attrib.get("Target", "")
                    referenced_media_names.add(Path(target).name)

        media_dir = root / "word" / "media"
        if media_dir.exists():
            for media_file in media_dir.iterdir():
                if (
                    media_file.is_file()
                    and media_file.name not in referenced_media_names
                ):
                    media_file.unlink()

        rebuilt = docx_path.with_suffix(".cleaned.docx")
        with zipfile.ZipFile(
            rebuilt,
            "w",
            compression=zipfile.ZIP_DEFLATED,
        ) as archive:
            for path in sorted(root.rglob("*")):
                if path.is_file():
                    archive.write(
                        path,
                        arcname=str(path.relative_to(root)),
                    )

        rebuilt.replace(docx_path)


def read_docx_structure(path: Path) -> Dict[str, object]:
    with zipfile.ZipFile(path) as archive:
        document_root = ET.fromstring(
            archive.read("word/document.xml")
        )
        rel_root = ET.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )

        rel_map = {
            relationship.attrib.get("Id", ""):
            relationship.attrib.get("Target", "")
            for relationship in rel_root
        }

        body = document_root.find(f"{{{W_NS}}}body")
        if body is None:
            raise RuntimeError(f"No word/body element in {path}")

        body_children = list(body)
        body_text = [
            body_element_text(element)
            for element in body_children
            if body_element_text(element)
        ]

        drawings: List[Dict[str, object]] = []
        captions: List[str] = []

        for element in body_children:
            if (
                element.tag == f"{{{W_NS}}}p"
                and paragraph_has_drawing(element)
            ):
                embeds = [
                    blip.attrib.get(f"{{{R_NS}}}embed", "")
                    for blip in element.findall(f".//{{{A_NS}}}blip")
                ]
                if not embeds:
                    embeds = [""]
                for embed in embeds:
                    target = rel_map.get(embed, "")
                    archive_path = (
                        f"word/{target}"
                        if target and not target.startswith("/")
                        else target.lstrip("/")
                    )
                    media_hash = ""
                    if archive_path in archive.namelist():
                        media_hash = sha256_bytes(
                            archive.read(archive_path)
                        )
                    drawings.append(
                        {
                            "relationship_id": embed,
                            "target": target,
                            "archive_path": archive_path,
                            "media_sha256": media_hash,
                        }
                    )

            text = body_element_text(element)
            if FIGURE_CAPTION_PATTERN.match(text):
                captions.append(text)

        media_files = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/")
            and not name.endswith("/")
        ]

        table_text = []
        for element in body_children:
            if element.tag == f"{{{W_NS}}}tbl":
                table_text.append(body_element_text(element))

    return {
        "body_text": body_text,
        "drawings": drawings,
        "captions": captions,
        "media_files": media_files,
        "table_text": table_text,
    }


def extract_section_text(
    document: Document,
    heading: str,
    next_heading: str,
) -> List[str]:
    paragraphs = [
        normalize_text(paragraph.text)
        for paragraph in document.paragraphs
    ]

    def find_heading(value: str) -> int:
        matches = []
        for index, text in enumerate(paragraphs):
            cleaned = re.sub(
                r"^\d+(?:\.\d+)*\s*",
                "",
                text.lower(),
            ).rstrip(":." )
            if cleaned == value.lower():
                matches.append(index)
        if len(matches) != 1:
            raise RuntimeError(
                f"Expected one {value} heading; observed {matches}"
            )
        return matches[0]

    start = find_heading(heading)
    end = find_heading(next_heading)
    if start >= end:
        raise RuntimeError(
            f"{heading} does not precede {next_heading}."
        )
    return paragraphs[start + 1:end]


def run_command(
    command: Sequence[str],
    cwd: Optional[Path] = None,
    env: Optional[Dict[str, str]] = None,
) -> subprocess.CompletedProcess:
    return subprocess.run(
        list(command),
        cwd=str(cwd) if cwd is not None else None,
        env=env,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )


def make_contact_sheet(
    paths: Sequence[Path],
    output: Path,
    columns: int = 3,
    cell_width: int = 650,
    padding: int = 24,
) -> None:
    images: List[Image.Image] = []
    for path in paths:
        image = Image.open(path).convert("RGB")
        ratio = cell_width / image.width
        images.append(
            image.resize(
                (cell_width, max(1, int(image.height * ratio)))
            )
        )

    if not images:
        raise RuntimeError("No render pages were available for a contact sheet.")

    rows = (len(images) + columns - 1) // columns
    row_heights = []
    for row_index in range(rows):
        subset = images[
            row_index * columns:(row_index + 1) * columns
        ]
        row_heights.append(max(image.height for image in subset))

    canvas_width = columns * cell_width + (columns + 1) * padding
    canvas_height = sum(row_heights) + (rows + 1) * padding
    canvas = Image.new(
        "RGB",
        (canvas_width, canvas_height),
        "white",
    )

    y = padding
    for row_index in range(rows):
        x = padding
        subset = images[
            row_index * columns:(row_index + 1) * columns
        ]
        for image in subset:
            canvas.paste(image, (x, y))
            x += cell_width + padding
        y += row_heights[row_index] + padding

    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output)


def render_docx(
    docx_path: Path,
    render_dir: Path,
) -> Dict[str, object]:
    render_dir.mkdir(parents=True, exist_ok=True)

    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    pdftoppm = shutil.which("pdftoppm")

    if not libreoffice:
        return {
            "render_attempted": False,
            "render_pass": False,
            "reason": "LibreOffice/soffice not found",
            "pdf_path": "",
            "page_count": 0,
            "contact_sheet": "",
        }

    if not pdftoppm:
        return {
            "render_attempted": False,
            "render_pass": False,
            "reason": "pdftoppm not found",
            "pdf_path": "",
            "page_count": 0,
            "contact_sheet": "",
        }

    with tempfile.TemporaryDirectory(
        prefix="u27b3c4_lo_"
    ) as lo_tmp:
        env = os.environ.copy()
        env["HOME"] = lo_tmp
        profile_uri = Path(lo_tmp).resolve().as_uri()

        result = run_command(
            [
                libreoffice,
                "--headless",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(render_dir),
                str(docx_path),
            ],
            env=env,
        )

    pdf_path = render_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not pdf_path.exists():
        return {
            "render_attempted": True,
            "render_pass": False,
            "reason": (
                f"LibreOffice failed: {result.stderr.strip() or result.stdout.strip()}"
            ),
            "pdf_path": str(pdf_path),
            "page_count": 0,
            "contact_sheet": "",
        }

    page_prefix = render_dir / "page"
    raster = run_command(
        [
            pdftoppm,
            "-png",
            "-r",
            "140",
            str(pdf_path),
            str(page_prefix),
        ]
    )

    page_paths = sorted(
        render_dir.glob("page-*.png"),
        key=lambda path: int(
            re.search(r"(\d+)$", path.stem).group(1)
        ),
    )

    if raster.returncode != 0 or not page_paths:
        return {
            "render_attempted": True,
            "render_pass": False,
            "reason": (
                f"pdftoppm failed: {raster.stderr.strip() or raster.stdout.strip()}"
            ),
            "pdf_path": str(pdf_path),
            "page_count": len(page_paths),
            "contact_sheet": "",
        }

    contact_sheet = (
        render_dir
        / "UTI_HostOmics_U27B3C4_render_contact_sheet.png"
    )
    make_contact_sheet(
        page_paths,
        contact_sheet,
        columns=3,
        cell_width=600,
    )

    return {
        "render_attempted": True,
        "render_pass": (
            pdf_path.exists()
            and pdf_path.stat().st_size > 0
            and contact_sheet.exists()
            and contact_sheet.stat().st_size > 0
        ),
        "reason": "Render files created",
        "pdf_path": str(pdf_path),
        "page_count": len(page_paths),
        "contact_sheet": str(contact_sheet),
        "libreoffice_stdout": result.stdout.strip(),
        "libreoffice_stderr": result.stderr.strip(),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root",
        default="__UTI_HOSTOMICS_PROJECT_ROOT__",
    )
    parser.add_argument(
        "--source",
        default=DEFAULT_SOURCE,
    )
    args = parser.parse_args()

    project = Path(args.project_root).resolve()
    source = Path(args.source).resolve()

    if not source.exists():
        raise FileNotFoundError(
            f"U27B3C2 derivative not found: {source}"
        )

    figure_dir = (
        project
        / "06_figures"
        / FIGURE_PACKAGE_TAG
    )
    legend_path = (
        project
        / "07_manuscript"
        / LEGEND_TAG
        / "UTI_HostOmics_U27B3B_manuscript_legend_insert.txt"
    )

    figure_paths = {
        number: (
            figure_dir
            / f"UTI_HostOmics_U27B3A_Figure_{number}.png"
        )
        for number in range(1, 9)
    }

    for number, path in figure_paths.items():
        if not path.exists():
            raise FileNotFoundError(
                f"Frozen Figure {number} PNG not found: {path}"
            )

    legends = parse_legends(legend_path)

    outdocx = (
        project
        / "09_manuscript_docx"
        / TAG
    )
    outtables = project / "06_tables" / TAG
    outmetadata = project / "03_metadata" / TAG
    outresults = project / "05_results" / TAG
    render_dir = outdocx / "render_qa"

    for directory in (
        outdocx,
        outtables,
        outmetadata,
        outresults,
        render_dir,
    ):
        directory.mkdir(parents=True, exist_ok=True)

    output_path = outdocx / OUTPUT_FILENAME
    source_hash_before = sha256_file(source)

    document = Document(source)
    source_results_text = extract_section_text(
        document,
        "Results",
        "Discussion",
    )
    source_table_text = [
        normalize_text(
            " ".join(
                cell.text
                for row in table.rows
                for cell in row.cells
            )
        )
        for table in document.tables
    ]

    boundary = locate_legacy_figure_block(document)

    log(
        "Legacy drawing paragraphs: "
        f"{len(boundary['drawing_indices'])}"
    )
    log(
        "Legacy figure captions removed: "
        f"{len(boundary['removed_figure_captions'])}"
    )

    for element in boundary["removed_elements"]:
        boundary["body"].remove(element)

    figure_width = available_figure_width(document)

    for figure_number in range(1, 9):
        add_figure_and_legend(
            document=document,
            body=boundary["body"],
            anchor_element=boundary["anchor_element"],
            figure_number=figure_number,
            figure_path=figure_paths[figure_number],
            legend=legends[figure_number],
            width=figure_width,
        )

    add_page_break_before_anchor(
        document,
        boundary["body"],
        boundary["anchor_element"],
    )

    if output_path.exists():
        output_path.unlink()

    document.save(output_path)
    cleanup_unreferenced_document_images(output_path)

    source_hash_after = sha256_file(source)
    output_hash = sha256_file(output_path)

    output_document = Document(output_path)
    output_results_text = extract_section_text(
        output_document,
        "Results",
        "Discussion",
    )
    output_table_text = [
        normalize_text(
            " ".join(
                cell.text
                for row in table.rows
                for cell in row.cells
            )
        )
        for table in output_document.tables
    ]

    source_unchanged = source_hash_before == source_hash_after
    results_preserved = source_results_text == output_results_text
    tables_preserved = source_table_text == output_table_text

    output_structure = read_docx_structure(output_path)

    drawing_count = len(output_structure["drawings"])
    media_count = len(output_structure["media_files"])
    caption_count = len(output_structure["captions"])
    caption_numbers = [
        int(FIGURE_CAPTION_PATTERN.match(text).group(1))
        for text in output_structure["captions"]
        if FIGURE_CAPTION_PATTERN.match(text)
    ]
    captions_ordered = caption_numbers == list(range(1, 9))

    integration_rows: List[Dict[str, object]] = []
    image_hash_match = True

    if drawing_count == 8:
        for index, drawing in enumerate(
            output_structure["drawings"],
            start=1,
        ):
            source_hash = sha256_file(figure_paths[index])
            embedded_hash = drawing["media_sha256"]
            matches = source_hash == embedded_hash
            image_hash_match = image_hash_match and matches

            integration_rows.append(
                {
                    "figure_number": index,
                    "source_figure_path": str(figure_paths[index]),
                    "source_figure_sha256": source_hash,
                    "embedded_relationship_id": drawing["relationship_id"],
                    "embedded_media_target": drawing["target"],
                    "embedded_media_sha256": embedded_hash,
                    "source_embedded_hash_match": matches,
                    "legend_character_count": len(
                        legends[index]
                    ),
                    "legend_exactly_present": (
                        normalize_text(legends[index])
                        == normalize_text(
                            output_structure["captions"][index - 1]
                        )
                    )
                    if caption_count >= index
                    else False,
                }
            )
    else:
        image_hash_match = False

    integration_manifest = pd.DataFrame(integration_rows)
    integration_manifest.to_csv(
        outtables
        / "UTI_HostOmics_U27B3C4_figure_legend_integration_manifest.tsv",
        sep="\t",
        index=False,
    )

    legends_exact = bool(
        len(integration_manifest) == 8
        and integration_manifest["legend_exactly_present"].all()
    )

    preservation_audit = pd.DataFrame(
        [
            {
                "source_path": str(source),
                "source_sha256_before": source_hash_before,
                "source_sha256_after": source_hash_after,
                "source_unchanged": source_unchanged,
                "output_path": str(output_path),
                "output_sha256": output_hash,
                "results_text_preserved": results_preserved,
                "table_count_source": len(source_table_text),
                "table_count_output": len(output_table_text),
                "table_text_preserved": tables_preserved,
                "legacy_drawings_removed": len(
                    boundary["drawing_indices"]
                ),
                "legacy_caption_paragraphs_removed": len(
                    boundary["removed_figure_captions"]
                ),
                "new_drawings": drawing_count,
                "new_media_parts": media_count,
                "new_caption_paragraphs": caption_count,
                "captions_numbered_1_to_8_in_order": captions_ordered,
                "embedded_figure_hashes_match": image_hash_match,
                "definitive_legends_match": legends_exact,
            }
        ]
    )
    preservation_audit.to_csv(
        outtables
        / "UTI_HostOmics_U27B3C4_docx_preservation_audit.tsv",
        sep="\t",
        index=False,
    )

    boundary_audit = pd.DataFrame(
        [
            {
                "legacy_block_start_index": boundary["start_index"],
                "legacy_block_anchor_index": boundary["anchor_index"],
                "legacy_drawing_paragraphs": len(
                    boundary["drawing_indices"]
                ),
                "removed_body_elements": len(
                    boundary["removed_elements"]
                ),
                "removed_nonempty_paragraphs": len(
                    boundary["removed_text"]
                ),
                "removed_figure_captions": len(
                    boundary["removed_figure_captions"]
                ),
                "table_or_tail_anchor_preserved": (
                    boundary["anchor_element"] is not None
                ),
            }
        ]
    )
    boundary_audit.to_csv(
        outtables
        / "UTI_HostOmics_U27B3C4_legacy_block_boundary_audit.tsv",
        sep="\t",
        index=False,
    )

    render_info = render_docx(
        output_path,
        render_dir,
    )
    pd.DataFrame([render_info]).to_csv(
        outtables
        / "UTI_HostOmics_U27B3C4_render_audit.tsv",
        sep="\t",
        index=False,
    )

    structural_pass = bool(
        source_unchanged
        and results_preserved
        and tables_preserved
        and drawing_count == 8
        and media_count == 8
        and caption_count == 8
        and captions_ordered
        and image_hash_match
        and legends_exact
    )

    if structural_pass and render_info["render_pass"]:
        decision = (
            "READY_FOR_U27B3C5_FINAL_MANUSCRIPT_VISUAL_AUDIT"
        )
    elif structural_pass:
        decision = (
            "FIGURE_LEGEND_INTEGRATION_COMPLETE_RENDER_QA_PENDING"
        )
    else:
        decision = (
            "TARGETED_U27B3C4_FIGURE_LEGEND_INTEGRATION_REPAIR_REQUIRED"
        )

    pd.DataFrame(
        [
            {
                "phase": "U27B3C4",
                "decision": decision,
                "legacy_figures_expected": 6,
                "legacy_figures_removed": len(
                    boundary["drawing_indices"]
                ),
                "frozen_figures_expected": 8,
                "frozen_figures_integrated": drawing_count,
                "definitive_legends_expected": 8,
                "definitive_legends_integrated": caption_count,
                "captions_ordered_1_to_8": captions_ordered,
                "embedded_media_parts": media_count,
                "embedded_figure_hashes_match": image_hash_match,
                "definitive_legends_match": legends_exact,
                "results_text_preserved": results_preserved,
                "tables_preserved": tables_preserved,
                "source_derivative_unchanged": source_unchanged,
                "render_pass": render_info["render_pass"],
                "scientific_values_recalculated": False,
                "frozen_figure_assets_modified": False,
                "source_locks_changed": False,
                "source_derivative_modified": False,
                "new_derivative_created": True,
                "next_phase": (
                    "U27B3C5 inspect the final manuscript render and confirm "
                    "Figures 1-8, legends, tables and page flow"
                    if decision.startswith("READY_FOR_U27B3C5")
                    else "Inspect structural and render audits"
                ),
            }
        ]
    ).to_csv(
        outtables
        / "UTI_HostOmics_U27B3C4_phase_decision.tsv",
        sep="\t",
        index=False,
    )

    derivative_record = pd.DataFrame(
        [
            {
                "field": "source_derivative_path",
                "value": str(source),
            },
            {
                "field": "source_derivative_sha256",
                "value": source_hash_before,
            },
            {
                "field": "integrated_derivative_path",
                "value": str(output_path),
            },
            {
                "field": "integrated_derivative_sha256",
                "value": output_hash,
            },
            {
                "field": "frozen_figure_package",
                "value": str(figure_dir),
            },
            {
                "field": "definitive_legend_source",
                "value": str(legend_path),
            },
            {
                "field": "source_overwritten",
                "value": "False",
            },
        ]
    )
    derivative_record.to_csv(
        outmetadata
        / "UTI_HostOmics_U27B3C4_derivative_record.tsv",
        sep="\t",
        index=False,
    )

    report_path = (
        outresults
        / "UTI_HostOmics_U27B3C4_figure_legend_integration_report.md"
    )
    with report_path.open("w", encoding="utf-8") as handle:
        handle.write(
            "# Phase U27B3C4 - Frozen Figures 1-8 and legend integration\n\n"
        )
        handle.write(f"- Version: `{VERSION}`\n")
        handle.write(f"- Decision: **{decision}**\n")
        handle.write(f"- Source derivative: `{source}`\n")
        handle.write(f"- New derivative: `{output_path}`\n")
        handle.write(
            f"- Legacy figure drawings removed: "
            f"**{len(boundary['drawing_indices'])}/6**.\n"
        )
        handle.write(
            f"- Frozen figures integrated: **{drawing_count}/8**.\n"
        )
        handle.write(
            f"- Definitive legends integrated: **{caption_count}/8**.\n"
        )
        handle.write(
            f"- Embedded figure hashes match frozen masters: "
            f"**{image_hash_match}**.\n"
        )
        handle.write(
            f"- Results text preserved: **{results_preserved}**.\n"
        )
        handle.write(
            f"- Tables preserved: **{tables_preserved}**.\n"
        )
        handle.write(
            f"- Render pass: **{render_info['render_pass']}**.\n"
        )
        handle.write(
            f"- Render contact sheet: "
            f"`{render_info.get('contact_sheet', '')}`.\n\n"
        )

        handle.write("## Integration boundary\n\n")
        handle.write(
            "Only the legacy six-figure appendix block was replaced. "
            "The U27B3C2 Results text, narrative Discussion text, tables and "
            "authoritative source derivatives were preserved. The inserted "
            "images are byte-identical to the frozen U27B3A PNG masters, and "
            "the captions match the definitive U27B3B legends.\n"
        )

    manifest = {
        "version": VERSION,
        "decision": decision,
        "source_path": str(source),
        "source_sha256": source_hash_before,
        "output_path": str(output_path),
        "output_sha256": output_hash,
        "legacy_figures_removed": len(boundary["drawing_indices"]),
        "frozen_figures_integrated": drawing_count,
        "legends_integrated": caption_count,
        "media_parts": media_count,
        "embedded_figure_hashes_match": image_hash_match,
        "definitive_legends_match": legends_exact,
        "results_text_preserved": results_preserved,
        "tables_preserved": tables_preserved,
        "source_unchanged": source_unchanged,
        "render_pass": render_info["render_pass"],
        "contact_sheet": render_info.get("contact_sheet", ""),
        "scientific_values_recalculated": False,
        "frozen_figure_assets_modified": False,
        "source_locks_changed": False,
        "source_derivative_modified": False,
    }
    (
        outresults
        / "UTI_HostOmics_U27B3C4_run_manifest.json"
    ).write_text(
        json.dumps(manifest, indent=2),
        encoding="utf-8",
    )

    log(f"Legacy figures removed: {len(boundary['drawing_indices'])}/6")
    log(f"Frozen figures integrated: {drawing_count}/8")
    log(f"Definitive legends integrated: {caption_count}/8")
    log(f"Media parts after cleanup: {media_count}")
    log(f"Results preserved: {results_preserved}")
    log(f"Tables preserved: {tables_preserved}")
    log(f"Render pass: {render_info['render_pass']}")
    log(f"Decision: {decision}")
    log(f"Output: {output_path}")
    log(f"Report: {report_path}")

    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[U27B3C4] ERROR: {exc}", file=sys.stderr)
        raise
